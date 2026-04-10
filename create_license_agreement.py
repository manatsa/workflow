#!/usr/bin/env python3
"""
Script to generate a Software License Agreement for Sonar Workflow System
Tailored for the Zimbabwean market
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_license_agreement():
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================
    # COVER PAGE
    # ============================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('SONAR WORKFLOW SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('SOFTWARE LICENSE AGREEMENT', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.add_run('Version 1.0').bold = True
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Parties on cover
    cover_info = doc.add_paragraph()
    cover_info.add_run('Between').italic = True
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    licensor = doc.add_paragraph()
    licensor.add_run('SONAR MICROSYSTEMS PVT. LTD.').bold = True
    licensor.alignment = WD_ALIGN_PARAGRAPH.CENTER

    and_text = doc.add_paragraph()
    and_text.add_run('and').italic = True
    and_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    licensee = doc.add_paragraph()
    licensee.add_run('[LICENSEE COMPANY NAME]').bold = True
    licensee.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    market = doc.add_paragraph()
    market.add_run('For Operations in the Republic of Zimbabwe').italic = True
    market.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Document info box
    doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc_info = doc.add_paragraph()
    doc_info.add_run('Agreement Reference: ').bold = True
    doc_info.add_run('SONAR-LIC-ZW-___________')
    doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_info = doc.add_paragraph()
    date_info.add_run('Effective Date: ').bold = True
    date_info.add_run('_______________, 20____')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.add_run('CONFIDENTIAL AND PROPRIETARY').bold = True
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER

    copyright_notice = doc.add_paragraph()
    copyright_notice.add_run(f'© {datetime.now().year} Sonar Microsystems Pvt. Ltd. All Rights Reserved.')
    copyright_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break for main content
    doc.add_page_break()

    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    doc.add_heading('TABLE OF CONTENTS', level=1)

    toc_items = [
        ('PART I: LICENSING MODEL AND PRICING', 3),
        ('PART II: LICENSE AGREEMENT', 5),
        ('Article 1: Definitions', 5),
        ('Article 2: Grant of License', 7),
        ('Article 3: License Restrictions', 8),
        ('Article 4: Fees and Payment', 9),
        ('Article 5: Delivery and Installation', 11),
        ('Article 6: Support and Maintenance', 12),
        ('Article 7: Training', 14),
        ('Article 8: Warranties', 15),
        ('Article 9: Limitation of Liability', 16),
        ('Article 10: Intellectual Property', 17),
        ('Article 11: Confidentiality', 18),
        ('Article 12: Data Protection and Privacy', 19),
        ('Article 13: Compliance and Audit', 20),
        ('Article 14: Term and Termination', 21),
        ('Article 15: Effects of Termination', 22),
        ('Article 16: Dispute Resolution', 23),
        ('Article 17: General Provisions', 24),
        ('SCHEDULES', 26),
        ('Schedule A: Licensed Software Specifications', 26),
        ('Schedule B: Pricing and Payment Schedule', 27),
        ('Schedule C: Service Level Agreement', 28),
        ('Schedule D: Authorized Users', 29),
        ('SIGNATURE PAGE', 30),
    ]

    for item, page in toc_items:
        toc_para = doc.add_paragraph()
        if item.startswith('PART') or item.startswith('SCHEDULES') or item.startswith('SIGNATURE'):
            toc_para.add_run(item).bold = True
        else:
            toc_para.add_run('    ' + item)
        # Add dots and page number
        toc_para.add_run('.' * 3 + f' {page}')

    doc.add_page_break()

    # ============================================
    # PART I: LICENSING MODEL AND PRICING
    # ============================================
    doc.add_heading('PART I: LICENSING MODEL AND PRICING', level=1)

    intro = doc.add_paragraph()
    intro.add_run('Sonar Microsystems Pvt. Ltd. offers the Sonar Workflow System under a flexible, tiered licensing model designed to meet the diverse needs of organizations in the Zimbabwean market. This section outlines the available licensing options, features, and pricing structure.')

    doc.add_paragraph()

    # Section 1: Licensing Tiers
    doc.add_heading('1. LICENSING TIERS', level=2)

    tier_intro = doc.add_paragraph()
    tier_intro.add_run('The Sonar Workflow System is available in three (3) licensing tiers, each designed to address specific organizational requirements:')

    doc.add_paragraph()

    # Create licensing tiers table
    tier_table = doc.add_table(rows=13, cols=4)
    tier_table.style = 'Table Grid'

    # Header row
    headers = ['Feature / Component', 'Standard Edition', 'Professional Edition', 'Enterprise Edition']
    for i, header in enumerate(headers):
        cell = tier_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # Tier data
    tier_data = [
        ['Maximum Named Users', 'Up to 25 users', 'Up to 100 users', 'Unlimited users'],
        ['Concurrent Users', 'Up to 10', 'Up to 50', 'Unlimited'],
        ['Workflow Definitions', 'Up to 10', 'Up to 50', 'Unlimited'],
        ['Form Builder', 'Basic', 'Advanced', 'Advanced + Custom'],
        ['Approval Workflows', 'Up to 3 levels', 'Up to 7 levels', 'Unlimited levels'],
        ['Reporting & Analytics', 'Standard Reports', 'Advanced Analytics', 'Custom BI Integration'],
        ['API Access', 'Limited', 'Full REST API', 'Full API + Webhooks'],
        ['Multi-SBU Support', 'Single SBU', 'Up to 5 SBUs', 'Unlimited SBUs'],
        ['Audit Trail', 'Basic (90 days)', 'Extended (2 years)', 'Comprehensive (Unlimited)'],
        ['Support Level', 'Email Support', 'Priority Support', 'Dedicated Account Manager'],
        ['Training', '5 hours included', '15 hours included', 'Unlimited training'],
        ['Customization', 'Not included', 'Limited', 'Full customization'],
    ]

    for row_idx, row_data in enumerate(tier_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = tier_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Section 2: Pricing Structure
    doc.add_heading('2. PRICING STRUCTURE', level=2)

    pricing_intro = doc.add_paragraph()
    pricing_intro.add_run('All prices are quoted in United States Dollars (USD) and are exclusive of applicable taxes. Payments may also be made in Zimbabwean Dollars (ZWL) at the prevailing interbank exchange rate on the date of invoice.')

    doc.add_paragraph()

    # 2.1 License Fees
    doc.add_heading('2.1 Initial License Fees', level=3)

    license_table = doc.add_table(rows=4, cols=3)
    license_table.style = 'Table Grid'

    license_headers = ['Edition', 'Perpetual License (One-Time)', 'Subscription (Annual)']
    for i, header in enumerate(license_headers):
        cell = license_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')

    license_data = [
        ['Standard Edition', 'USD 15,000', 'USD 6,000 / year'],
        ['Professional Edition', 'USD 35,000', 'USD 14,000 / year'],
        ['Enterprise Edition', 'USD 75,000', 'USD 30,000 / year'],
    ]

    for row_idx, row_data in enumerate(license_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = license_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # 2.2 Additional User Licenses
    doc.add_heading('2.2 Additional User Licenses', level=3)

    user_table = doc.add_table(rows=4, cols=3)
    user_table.style = 'Table Grid'

    user_headers = ['Edition', 'Per Named User (One-Time)', 'Per Named User (Annual)']
    for i, header in enumerate(user_headers):
        cell = user_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFF2CC')

    user_data = [
        ['Standard Edition', 'USD 200 / user', 'USD 80 / user / year'],
        ['Professional Edition', 'USD 300 / user', 'USD 120 / user / year'],
        ['Enterprise Edition', 'USD 400 / user', 'USD 160 / user / year'],
    ]

    for row_idx, row_data in enumerate(user_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = user_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # 2.3 Support and Maintenance
    doc.add_heading('2.3 Annual Support and Maintenance', level=3)

    support_para = doc.add_paragraph()
    support_para.add_run('For Perpetual License holders, annual support and maintenance is charged at ')
    support_para.add_run('20% of the current list price').bold = True
    support_para.add_run(' of the licensed edition. Support and maintenance includes:')

    support_items = [
        'Software updates and patches',
        'Security updates and vulnerability fixes',
        'Access to new minor version releases',
        'Technical support via designated channels',
        'Access to online knowledge base and documentation',
    ]

    for item in support_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # 2.4 Implementation Services
    doc.add_heading('2.4 Implementation and Professional Services', level=3)

    impl_table = doc.add_table(rows=7, cols=2)
    impl_table.style = 'Table Grid'

    impl_headers = ['Service', 'Rate (USD)']
    for i, header in enumerate(impl_headers):
        cell = impl_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FCE4D6')

    impl_data = [
        ['Standard Implementation Package', 'USD 5,000 (fixed)'],
        ['Professional Implementation Package', 'USD 12,000 (fixed)'],
        ['Enterprise Implementation Package', 'USD 25,000+ (scoped)'],
        ['On-site Consultant (per day)', 'USD 800 / day'],
        ['Remote Consultant (per hour)', 'USD 100 / hour'],
        ['Custom Development (per hour)', 'USD 150 / hour'],
    ]

    for row_idx, row_data in enumerate(impl_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = impl_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # Section 3: Licensing Models
    doc.add_heading('3. LICENSING MODELS', level=2)

    # 3.1 Perpetual License
    doc.add_heading('3.1 Perpetual License Model', level=3)

    perp_para = doc.add_paragraph()
    perp_para.add_run('The Perpetual License grants the Licensee a permanent, non-exclusive right to use the licensed version of the Sonar Workflow System, subject to the terms of this Agreement. Key features include:')

    perp_items = [
        'One-time license fee payment',
        'Perpetual right to use the licensed version',
        'Optional annual support and maintenance subscription',
        'Major version upgrades available at discounted rates (50% of new license fee)',
        'No ongoing license fees after initial purchase',
    ]

    for item in perp_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # 3.2 Subscription License
    doc.add_heading('3.2 Subscription License Model', level=3)

    sub_para = doc.add_paragraph()
    sub_para.add_run('The Subscription License grants the Licensee a time-limited, non-exclusive right to use the Sonar Workflow System for the subscription period. Key features include:')

    sub_items = [
        'Annual subscription fee, payable in advance',
        'Access to all updates, patches, and new versions during the subscription period',
        'Full support and maintenance included in subscription fee',
        'Flexible scaling: add or remove users as needed',
        'Lower initial investment compared to perpetual license',
        'License terminates upon expiration or non-renewal of subscription',
    ]

    for item in sub_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Section 4: Special Terms for Zimbabwe
    doc.add_heading('4. SPECIAL TERMS FOR ZIMBABWEAN MARKET', level=2)

    zw_terms = [
        ('Currency and Payment', 'Payments may be made in USD, ZWL, or South African Rand (ZAR). Exchange rates will be determined using the official Reserve Bank of Zimbabwe interbank rate on the invoice date. A 3% foreign exchange buffer may be applied to account for rate fluctuations.'),
        ('Local Partner', 'Sonar Microsystems may appoint an authorized local partner in Zimbabwe to provide implementation, training, and first-level support services. The local partner will act as an agent of Sonar Microsystems for the purposes of this Agreement.'),
        ('Tax Compliance', 'The Licensee is responsible for all applicable Zimbabwean taxes, including but not limited to Value Added Tax (VAT), withholding taxes, and any other levies. Sonar Microsystems will provide necessary documentation for tax compliance purposes.'),
        ('Local Data Hosting', 'For Licensees requiring local data residency, Sonar Microsystems offers an optional local hosting arrangement through certified Zimbabwean data center partners at additional cost.'),
        ('Government and Parastatal Discounts', 'Special pricing may be available for Government of Zimbabwe entities, parastatals, and registered non-profit organizations. Contact Sonar Microsystems for specific quotations.'),
    ]

    for i, (title, content) in enumerate(zw_terms):
        para = doc.add_paragraph()
        para.add_run(f'4.{i + 1} {title}. ').bold = True
        para.add_run(content)

    doc.add_page_break()

    # ============================================
    # PART II: LICENSE AGREEMENT
    # ============================================
    doc.add_heading('PART II: SOFTWARE LICENSE AGREEMENT', level=1)

    doc.add_paragraph()

    # Preamble
    preamble = doc.add_paragraph()
    preamble.add_run('This Software License Agreement ("Agreement") is entered into as of the Effective Date set forth above by and between:')

    doc.add_paragraph()

    # Licensor
    licensor_para = doc.add_paragraph()
    licensor_para.add_run('LICENSOR: ').bold = True
    licensor_para.add_run('SONAR MICROSYSTEMS PVT. LTD.')
    licensor_para.add_run(', a company incorporated under the laws of India, having its registered office at ___________________________________ (hereinafter referred to as "')
    licensor_para.add_run('Sonar').bold = True
    licensor_para.add_run('" or "')
    licensor_para.add_run('Licensor').bold = True
    licensor_para.add_run('");')

    doc.add_paragraph()

    # Licensee
    licensee_para = doc.add_paragraph()
    licensee_para.add_run('LICENSEE: ').bold = True
    licensee_para.add_run('___________________________________')
    licensee_para.add_run(', a company incorporated under the laws of the Republic of Zimbabwe, having its registered office at ___________________________________, Registration Number: _______________ (hereinafter referred to as "')
    licensee_para.add_run('Licensee').bold = True
    licensee_para.add_run('");')

    doc.add_paragraph()

    collective = doc.add_paragraph()
    collective.add_run('(Sonar and Licensee are hereinafter individually referred to as a "Party" and collectively as the "Parties")')
    collective.italic = True

    doc.add_paragraph()

    # ============================================
    # ARTICLE 1: DEFINITIONS
    # ============================================
    doc.add_heading('ARTICLE 1: DEFINITIONS', level=1)

    definitions = [
        ('"Authorized Users"', 'means the employees, contractors, and agents of the Licensee who are authorized to access and use the Licensed Software pursuant to this Agreement, as specified in Schedule D.'),

        ('"Confidential Information"', 'means any non-public information disclosed by one Party to the other, including but not limited to trade secrets, business information, technical data, source code, and the terms of this Agreement.'),

        ('"Documentation"', 'means the user manuals, technical manuals, installation guides, training materials, and other documentation provided by Sonar in connection with the Licensed Software.'),

        ('"Effective Date"', 'means the date on which this Agreement is signed by both Parties or, if signed on different dates, the date of the last signature.'),

        ('"Error"', 'means any failure of the Licensed Software to conform in any material respect to the specifications set forth in the Documentation.'),

        ('"Initial Term"', 'means the initial period of this Agreement, as specified in Article 14.'),

        ('"Intellectual Property Rights"', 'means all patents, copyrights, trademarks, trade secrets, and other intellectual property rights, whether registered or unregistered.'),

        ('"License Fee"', 'means the fees payable by the Licensee for the license granted under this Agreement, as set forth in Schedule B.'),

        ('"Licensed Software"', 'means the Sonar Workflow System software licensed under this Agreement, including all modules, components, and features specified in Schedule A, together with any Updates and Upgrades provided during the Term.'),

        ('"Maintenance Services"', 'means the support and maintenance services provided by Sonar as described in Article 6 and Schedule C.'),

        ('"Named User"', 'means an individual Authorized User who is assigned a unique user credential to access the Licensed Software.'),

        ('"Concurrent User"', 'means an Authorized User who is actively logged into and using the Licensed Software at any given time.'),

        ('"Site"', 'means the physical location(s) in Zimbabwe where the Licensed Software is installed and used, as specified in Schedule A.'),

        ('"Territory"', 'means the Republic of Zimbabwe.'),

        ('"Update"', 'means a minor release of the Licensed Software that includes bug fixes, security patches, and minor enhancements, designated by a change in the version number after the first decimal point (e.g., 1.1 to 1.2).'),

        ('"Upgrade"', 'means a major release of the Licensed Software that includes significant new features or functionality, designated by a change in the version number before the first decimal point (e.g., 1.x to 2.0).'),
    ]

    for i, (term, definition) in enumerate(definitions):
        para = doc.add_paragraph()
        para.add_run(f'1.{i + 1} ').bold = True
        para.add_run(term).bold = True
        para.add_run(f' {definition}')

    doc.add_paragraph()

    # ============================================
    # ARTICLE 2: GRANT OF LICENSE
    # ============================================
    doc.add_heading('ARTICLE 2: GRANT OF LICENSE', level=1)

    license_sections = [
        ('License Grant', '''Subject to the terms and conditions of this Agreement and payment of the applicable License Fees, Sonar hereby grants to the Licensee a non-exclusive, non-transferable, limited license to:

(a) Install the Licensed Software on servers located at the Site(s) specified in Schedule A;
(b) Permit the number of Authorized Users specified in Schedule A to access and use the Licensed Software;
(c) Use the Licensed Software solely for the Licensee's internal business operations within the Territory;
(d) Make one (1) copy of the Licensed Software for backup and disaster recovery purposes;
(e) Use the Documentation in connection with the Licensee's authorized use of the Licensed Software.'''),

        ('License Type', 'The license granted hereunder is a [PERPETUAL / SUBSCRIPTION] license, as indicated in Schedule B. For Subscription licenses, the license is valid only for the subscription period and is contingent upon timely payment of subscription fees.'),

        ('User Limitations', 'The Licensee shall not permit more than the licensed number of Named Users or Concurrent Users (as applicable to the licensed edition) to access the Licensed Software at any time. The Licensee shall implement appropriate access controls to ensure compliance with these limitations.'),

        ('Edition-Specific Rights', 'The specific features, functionalities, and limitations applicable to the licensed edition (Standard, Professional, or Enterprise) are set forth in Part I of this Agreement and Schedule A. The Licensee shall not attempt to access or use features not included in the licensed edition.'),

        ('Territory Restriction', 'The license granted herein is limited to use within the Territory (Republic of Zimbabwe). The Licensee shall not use the Licensed Software, or permit the Licensed Software to be used, outside the Territory without Sonar\'s prior written consent.'),

        ('Reservation of Rights', 'All rights not expressly granted to the Licensee in this Agreement are reserved by Sonar. Nothing in this Agreement shall be construed as granting any rights to the Licensee by implication, estoppel, or otherwise.'),
    ]

    for i, (title, content) in enumerate(license_sections):
        para = doc.add_paragraph()
        para.add_run(f'2.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 3: LICENSE RESTRICTIONS
    # ============================================
    doc.add_heading('ARTICLE 3: LICENSE RESTRICTIONS', level=1)

    restrictions_intro = doc.add_paragraph()
    restrictions_intro.add_run('3.1 ').bold = True
    restrictions_intro.add_run('Prohibited Activities. ').bold = True
    restrictions_intro.add_run('The Licensee shall NOT:')

    restrictions = [
        'Copy, reproduce, or duplicate the Licensed Software except as expressly permitted in this Agreement;',
        'Modify, adapt, translate, or create derivative works based on the Licensed Software or Documentation;',
        'Reverse engineer, disassemble, decompile, or otherwise attempt to derive the source code of the Licensed Software;',
        'Remove, alter, or obscure any proprietary notices, labels, or trademarks on the Licensed Software or Documentation;',
        'Sublicense, rent, lease, loan, sell, distribute, or otherwise transfer the Licensed Software or any rights therein to any third party;',
        'Use the Licensed Software to provide services to third parties, including but not limited to software-as-a-service, time-sharing, or service bureau arrangements, without Sonar\'s prior written consent;',
        'Use the Licensed Software in any manner that violates applicable laws or regulations of Zimbabwe or any other jurisdiction;',
        'Use the Licensed Software for any unlawful, fraudulent, or malicious purpose;',
        'Attempt to circumvent or disable any license management, copy protection, or access control mechanisms in the Licensed Software;',
        'Benchmark, test, or evaluate the Licensed Software for the purpose of comparing it to competing products or publishing performance comparisons;',
        'Use the Licensed Software to develop any product or service that competes with the Sonar Workflow System;',
        'Permit any third party to access or use the Licensed Software, except as expressly authorized in this Agreement;',
        'Use the Licensed Software in a manner that exceeds the licensed user count, concurrent user limit, or other usage restrictions;',
        'Transfer the Licensed Software to any hardware or location not authorized under this Agreement without Sonar\'s consent.',
    ]

    for i, restriction in enumerate(restrictions):
        para = doc.add_paragraph(f'({chr(97 + i)}) {restriction}')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()

    para_32 = doc.add_paragraph()
    para_32.add_run('3.2 ').bold = True
    para_32.add_run('Third-Party Components. ').bold = True
    para_32.add_run('The Licensed Software may include third-party software components that are subject to separate license terms. The Licensee agrees to comply with all such third-party license terms. A list of third-party components and their applicable licenses is provided in the Documentation.')

    para_33 = doc.add_paragraph()
    para_33.add_run('3.3 ').bold = True
    para_33.add_run('Open Source Components. ').bold = True
    para_33.add_run('To the extent the Licensed Software includes open source components, such components are licensed under the applicable open source license terms, which shall prevail over this Agreement with respect to such components only.')

    doc.add_paragraph()

    # ============================================
    # ARTICLE 4: FEES AND PAYMENT
    # ============================================
    doc.add_heading('ARTICLE 4: FEES AND PAYMENT', level=1)

    payment_sections = [
        ('License Fees', 'The Licensee shall pay the License Fees set forth in Schedule B. For Perpetual licenses, the License Fee is due upon execution of this Agreement. For Subscription licenses, the annual subscription fee is due in advance at the beginning of each subscription year.'),

        ('Payment Terms', '''All invoices shall be paid within thirty (30) days of the invoice date. Payments shall be made in the currency specified in Schedule B via wire transfer to the bank account designated by Sonar. The Licensee shall:

(a) Pay all amounts due without any deduction or set-off;
(b) Not withhold any amounts for taxes unless required by law; if withholding is required, the Licensee shall gross up the payment so that Sonar receives the full invoiced amount;
(c) Provide Sonar with valid tax receipts for any withheld amounts within thirty (30) days of payment.'''),

        ('Late Payment', 'Any amounts not paid when due shall bear interest at the rate of 1.5% per month (or the maximum rate permitted by law, if lower), calculated from the due date until paid. Sonar may also suspend the license and/or Maintenance Services if payment is more than thirty (30) days overdue.'),

        ('Taxes', 'All fees are exclusive of applicable taxes. The Licensee shall be responsible for all taxes, duties, levies, and assessments imposed by any government authority in connection with this Agreement, including but not limited to Zimbabwe Value Added Tax (VAT), withholding taxes, and customs duties. Sonar shall be responsible only for taxes based on its net income.'),

        ('Price Adjustments', 'For Subscription licenses and Maintenance Services, Sonar may adjust prices annually upon sixty (60) days\' prior written notice. Price increases shall not exceed the greater of: (a) 10% of the prior year\'s fee; or (b) the annual inflation rate as published by the Reserve Bank of Zimbabwe.'),

        ('Currency Fluctuation', 'If payment is made in a currency other than USD, the exchange rate shall be the official interbank rate published by the Reserve Bank of Zimbabwe on the invoice date. Sonar reserves the right to invoice in USD if exchange rate volatility exceeds 20% in any quarter.'),

        ('Payment Disputes', 'If the Licensee disputes any invoice in good faith, the Licensee shall: (a) pay all undisputed amounts when due; (b) provide written notice of the dispute with supporting details within fifteen (15) days of the invoice date; and (c) work with Sonar in good faith to resolve the dispute. Disputed amounts determined to be owed shall be paid with interest from the original due date.'),
    ]

    for i, (title, content) in enumerate(payment_sections):
        para = doc.add_paragraph()
        para.add_run(f'4.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 5: DELIVERY AND INSTALLATION
    # ============================================
    doc.add_heading('ARTICLE 5: DELIVERY AND INSTALLATION', level=1)

    delivery_sections = [
        ('Delivery', 'Sonar shall deliver the Licensed Software to the Licensee electronically within five (5) business days of receipt of the initial License Fee payment. Delivery shall include the software installation files, license keys, and Documentation.'),

        ('Installation', 'Unless otherwise agreed in Schedule B, installation of the Licensed Software shall be performed by Sonar or its authorized representative. The Licensee shall provide Sonar with reasonable access to the Site and necessary technical infrastructure for installation.'),

        ('System Requirements', 'The Licensee shall ensure that its hardware, operating systems, and network infrastructure meet the minimum system requirements specified in the Documentation. Sonar shall not be responsible for any performance issues arising from failure to meet these requirements.'),

        ('Acceptance Testing', '''Upon installation, the Licensee shall have fifteen (15) business days to conduct acceptance testing. The Licensed Software shall be deemed accepted upon the earlier of: (a) written acceptance by the Licensee; (b) expiration of the acceptance period without written rejection; or (c) the Licensee's productive use of the Licensed Software.

If the Licensee rejects the Licensed Software due to material non-conformance with the Documentation, Sonar shall have thirty (30) days to correct the non-conformance. If Sonar fails to correct the non-conformance, the Licensee may terminate this Agreement and receive a full refund of the License Fee.'''),

        ('Configuration', 'Initial configuration of the Licensed Software for the Licensee\'s specific business requirements is included in the implementation package selected in Schedule B. Additional configuration beyond the scope of the implementation package shall be provided at Sonar\'s then-current professional services rates.'),
    ]

    for i, (title, content) in enumerate(delivery_sections):
        para = doc.add_paragraph()
        para.add_run(f'5.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 6: SUPPORT AND MAINTENANCE
    # ============================================
    doc.add_heading('ARTICLE 6: SUPPORT AND MAINTENANCE', level=1)

    support_sections = [
        ('Maintenance Services', '''During the Term, and subject to payment of applicable Maintenance Fees (for Perpetual licenses) or as included in the Subscription Fee (for Subscription licenses), Sonar shall provide the following Maintenance Services:

(a) Updates: Provision of Updates as they become generally available;
(b) Error Correction: Correction of Errors reported by the Licensee in accordance with this Article;
(c) Technical Support: Access to Sonar's technical support team via the designated support channels;
(d) Documentation Updates: Updated Documentation reflecting changes in the Licensed Software;
(e) Knowledge Base: Access to Sonar's online knowledge base and self-help resources.'''),

        ('Support Levels', '''Support levels vary by edition as follows:

Standard Edition:
- Email support during business hours (Monday-Friday, 8:00 AM - 5:00 PM CAT)
- Response time: Within 2 business days
- No phone support

Professional Edition:
- Email and phone support during extended hours (Monday-Friday, 7:00 AM - 7:00 PM CAT)
- Response time: Within 1 business day for Priority 1 and 2 issues
- Dedicated support queue

Enterprise Edition:
- 24/7 email and phone support for Priority 1 issues
- Dedicated Account Manager
- Response time: Within 4 hours for Priority 1 issues
- On-site support available (additional charges may apply)'''),

        ('Error Classification', '''Errors shall be classified as follows:

Priority 1 (Critical): The Licensed Software is completely inoperable or a critical function is unavailable, causing severe business impact with no workaround available.

Priority 2 (High): A major function is impaired, causing significant business impact, but a workaround exists.

Priority 3 (Medium): A function is impaired but does not significantly impact business operations.

Priority 4 (Low): Minor issues, cosmetic defects, or general questions about usage.'''),

        ('Response and Resolution', 'Target response and resolution times are set forth in Schedule C (Service Level Agreement). Response times are measured from the time a properly submitted support request is received by Sonar. Resolution times are targets and not guarantees, as resolution may depend on factors outside Sonar\'s control.'),

        ('Licensee Responsibilities', '''The Licensee shall:

(a) Designate up to two (2) trained technical contacts authorized to submit support requests;
(b) Provide Sonar with sufficient information to reproduce and diagnose reported Errors;
(c) Implement Updates and security patches in a timely manner;
(d) Maintain appropriate backups of data and configurations;
(e) Cooperate with Sonar's reasonable requests for information and access necessary to provide support.'''),

        ('Exclusions', '''Maintenance Services do not include support for:

(a) Errors caused by modifications to the Licensed Software not made by Sonar;
(b) Errors caused by use of the Licensed Software other than in accordance with the Documentation;
(c) Errors caused by third-party software, hardware, or network failures;
(d) Errors caused by operator error or failure to follow documented procedures;
(e) Consulting, training, or custom development services (available separately);
(f) Support for versions of the Licensed Software that have reached end-of-life.'''),

        ('Remote Access', 'To facilitate support, the Licensee agrees to provide Sonar with secure remote access to the Licensed Software environment when reasonably requested. Such access shall be subject to appropriate security protocols and the Licensee\'s reasonable security policies.'),
    ]

    for i, (title, content) in enumerate(support_sections):
        para = doc.add_paragraph()
        para.add_run(f'6.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 7: TRAINING
    # ============================================
    doc.add_heading('ARTICLE 7: TRAINING', level=1)

    training_sections = [
        ('Initial Training', 'Sonar shall provide initial training to the Licensee\'s Authorized Users as specified in Schedule B. Training may be delivered on-site, remotely, or through a combination of methods as agreed by the Parties.'),

        ('Training Content', '''Initial training shall cover:

(a) System administration and configuration;
(b) User management and access control;
(c) Workflow creation and management;
(d) Form builder usage;
(e) Reporting and analytics;
(f) Best practices and operational procedures.'''),

        ('Additional Training', 'Training beyond the hours included in the licensed edition is available at Sonar\'s then-current training rates. The Licensee may also purchase training packages for new users or refresher training.'),

        ('Training Materials', 'Training materials provided during training sessions remain the property of Sonar. The Licensee may use training materials for internal training purposes only and shall not distribute or reproduce them for any other purpose.'),
    ]

    for i, (title, content) in enumerate(training_sections):
        para = doc.add_paragraph()
        para.add_run(f'7.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 8: WARRANTIES
    # ============================================
    doc.add_heading('ARTICLE 8: WARRANTIES', level=1)

    warranty_sections = [
        ('Software Warranty', 'Sonar warrants that, for a period of ninety (90) days from the date of delivery (the "Warranty Period"), the Licensed Software will perform substantially in accordance with the Documentation when used in compliance with this Agreement. Sonar\'s sole obligation under this warranty is to correct any Error that causes the Licensed Software to fail to conform to this warranty, or, at Sonar\'s option, to refund the License Fee paid for the non-conforming component.'),

        ('Media Warranty', 'Sonar warrants that any physical media on which the Licensed Software is delivered will be free from defects in materials and workmanship for a period of thirty (30) days from delivery. Sonar\'s sole obligation under this warranty is to replace defective media.'),

        ('Services Warranty', 'Sonar warrants that all professional services, including installation, implementation, and training, will be performed in a professional and workmanlike manner consistent with industry standards.'),

        ('Non-Infringement', 'Sonar warrants that, to the best of its knowledge, the Licensed Software does not infringe any valid patent, copyright, or trade secret of any third party in the Territory. This warranty is subject to the limitations and indemnification provisions in Articles 9 and 10.'),

        ('Disclaimer', 'EXCEPT FOR THE EXPRESS WARRANTIES SET FORTH IN THIS ARTICLE 8, THE LICENSED SOFTWARE AND DOCUMENTATION ARE PROVIDED "AS IS" WITHOUT WARRANTY OF ANY KIND. SONAR DISCLAIMS ALL OTHER WARRANTIES, EXPRESS, IMPLIED, OR STATUTORY, INCLUDING WITHOUT LIMITATION ANY WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, NON-INFRINGEMENT, ACCURACY, QUIET ENJOYMENT, OR ARISING FROM COURSE OF DEALING OR USAGE OF TRADE. SONAR DOES NOT WARRANT THAT THE LICENSED SOFTWARE WILL BE ERROR-FREE, UNINTERRUPTED, OR MEET THE LICENSEE\'S SPECIFIC REQUIREMENTS.'),
    ]

    for i, (title, content) in enumerate(warranty_sections):
        para = doc.add_paragraph()
        para.add_run(f'8.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 9: LIMITATION OF LIABILITY
    # ============================================
    doc.add_heading('ARTICLE 9: LIMITATION OF LIABILITY', level=1)

    liability_sections = [
        ('Exclusion of Consequential Damages', 'IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR EXEMPLARY DAMAGES, INCLUDING BUT NOT LIMITED TO DAMAGES FOR LOSS OF PROFITS, REVENUE, GOODWILL, USE, DATA, OR OTHER INTANGIBLE LOSSES, ARISING OUT OF OR RELATING TO THIS AGREEMENT, EVEN IF SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.'),

        ('Limitation of Direct Damages', 'EXCEPT FOR LIABILITY ARISING FROM (A) BREACH OF CONFIDENTIALITY OBLIGATIONS, (B) INFRINGEMENT OF INTELLECTUAL PROPERTY RIGHTS, (C) THE LICENSEE\'S BREACH OF LICENSE RESTRICTIONS, OR (D) GROSS NEGLIGENCE OR WILLFUL MISCONDUCT, THE TOTAL CUMULATIVE LIABILITY OF EITHER PARTY FOR ALL CLAIMS ARISING OUT OF OR RELATING TO THIS AGREEMENT SHALL NOT EXCEED THE GREATER OF: (I) THE TOTAL FEES PAID OR PAYABLE BY THE LICENSEE UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTHS PRECEDING THE CLAIM; OR (II) USD 50,000.'),

        ('Essential Basis', 'THE PARTIES ACKNOWLEDGE THAT THE LIMITATIONS SET FORTH IN THIS ARTICLE 9 ARE AN ESSENTIAL ELEMENT OF THE BARGAIN AND REFLECT A REASONABLE ALLOCATION OF RISK. THE LIMITATIONS SHALL APPLY REGARDLESS OF THE FORM OF ACTION, WHETHER IN CONTRACT, TORT, STRICT LIABILITY, OR OTHERWISE, AND EVEN IF ANY LIMITED REMEDY FAILS OF ITS ESSENTIAL PURPOSE.'),

        ('Exceptions', 'Nothing in this Agreement shall limit liability for: (a) death or personal injury caused by negligence; (b) fraud or fraudulent misrepresentation; (c) any liability that cannot be limited or excluded by applicable law.'),
    ]

    for i, (title, content) in enumerate(liability_sections):
        para = doc.add_paragraph()
        para.add_run(f'9.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 10: INTELLECTUAL PROPERTY
    # ============================================
    doc.add_heading('ARTICLE 10: INTELLECTUAL PROPERTY', level=1)

    ip_sections = [
        ('Ownership', 'The Licensed Software, Documentation, and all related Intellectual Property Rights are and shall remain the sole and exclusive property of Sonar. The Licensee acquires only the limited license rights expressly granted in this Agreement. Nothing in this Agreement constitutes a sale or assignment of the Licensed Software or any Intellectual Property Rights therein.'),

        ('No Implied Rights', 'No license or right is granted by implication, estoppel, or otherwise to any Intellectual Property Rights other than those expressly set forth in this Agreement.'),

        ('Feedback', 'If the Licensee provides any suggestions, ideas, or feedback regarding the Licensed Software ("Feedback"), Sonar shall own all rights in such Feedback and may use, implement, and commercialize it without obligation to the Licensee.'),

        ('Infringement Indemnification', '''Sonar shall indemnify, defend, and hold harmless the Licensee from any third-party claim alleging that the Licensed Software infringes any valid patent, copyright, or trade secret enforceable in the Territory, provided that:

(a) The Licensee promptly notifies Sonar in writing of any such claim;
(b) Sonar has sole control of the defense and any settlement negotiations;
(c) The Licensee cooperates with Sonar and provides reasonable assistance;
(d) The Licensee does not make any admission or compromise without Sonar's consent.'''),

        ('Remedies for Infringement', '''If the Licensed Software is held to infringe or Sonar believes it may be held to infringe, Sonar may, at its option and expense:

(a) Obtain the right for the Licensee to continue using the Licensed Software;
(b) Modify or replace the Licensed Software to make it non-infringing; or
(c) If neither (a) nor (b) is commercially practicable, terminate this Agreement and refund a pro-rata portion of the License Fee based on a five-year depreciation from the Effective Date.'''),

        ('Exclusions', '''Sonar shall have no liability for infringement claims arising from:

(a) Modification of the Licensed Software by parties other than Sonar;
(b) Combination of the Licensed Software with non-Sonar products, data, or processes;
(c) Use of the Licensed Software other than in accordance with this Agreement and the Documentation;
(d) Failure to use Updates that would have avoided the infringement;
(e) Third-party software or content provided by the Licensee.'''),

        ('Licensee Indemnification', 'The Licensee shall indemnify, defend, and hold harmless Sonar from any claims arising from: (a) the Licensee\'s breach of this Agreement; (b) the Licensee\'s data or content processed using the Licensed Software; (c) the Licensee\'s combination of the Licensed Software with third-party products or services; or (d) the Licensee\'s violation of applicable laws.'),
    ]

    for i, (title, content) in enumerate(ip_sections):
        para = doc.add_paragraph()
        para.add_run(f'10.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 11: CONFIDENTIALITY
    # ============================================
    doc.add_heading('ARTICLE 11: CONFIDENTIALITY', level=1)

    conf_sections = [
        ('Confidentiality Obligations', 'Each Party agrees to maintain the confidentiality of the other Party\'s Confidential Information using at least the same degree of care it uses to protect its own confidential information, but in no event less than reasonable care. Neither Party shall disclose the other Party\'s Confidential Information to any third party without the prior written consent of the disclosing Party.'),

        ('Permitted Disclosures', 'A Party may disclose Confidential Information to its employees, contractors, and advisors who have a need to know such information and who are bound by confidentiality obligations at least as protective as those in this Agreement.'),

        ('Exclusions', '''Confidential Information does not include information that:

(a) Is or becomes publicly available through no fault of the receiving Party;
(b) Was rightfully in the receiving Party's possession prior to disclosure;
(c) Is independently developed by the receiving Party without use of Confidential Information;
(d) Is rightfully obtained from a third party without breach of any confidentiality obligation;
(e) Is required to be disclosed by law, provided that the receiving Party gives prompt notice and cooperates with efforts to limit disclosure.'''),

        ('Return of Confidential Information', 'Upon termination of this Agreement, each Party shall return or destroy all Confidential Information of the other Party, except that a Party may retain one archival copy for legal compliance purposes.'),

        ('Survival', 'The confidentiality obligations in this Article 11 shall survive termination of this Agreement for a period of five (5) years.'),
    ]

    for i, (title, content) in enumerate(conf_sections):
        para = doc.add_paragraph()
        para.add_run(f'11.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 12: DATA PROTECTION AND PRIVACY
    # ============================================
    doc.add_heading('ARTICLE 12: DATA PROTECTION AND PRIVACY', level=1)

    data_sections = [
        ('Data Ownership', 'All data entered into or generated by the Licensed Software by the Licensee or its Authorized Users ("Licensee Data") shall remain the property of the Licensee. Sonar claims no ownership rights in Licensee Data.'),

        ('Data Processing', 'To the extent Sonar processes any personal data on behalf of the Licensee in connection with this Agreement, Sonar shall process such data only in accordance with the Licensee\'s instructions and applicable data protection laws, including the Zimbabwe Data Protection Act and any regulations thereunder.'),

        ('Data Protection Compliance', '''The Licensee acknowledges that it is responsible for:

(a) Ensuring that its collection and use of personal data complies with applicable laws;
(b) Obtaining all necessary consents and authorizations for data processing;
(c) Configuring the Licensed Software in accordance with its data protection obligations;
(d) Implementing appropriate technical and organizational measures to protect personal data.'''),

        ('Security Measures', 'Sonar shall maintain appropriate technical and organizational security measures to protect Licensee Data against unauthorized access, disclosure, or destruction, as described in Schedule C.'),

        ('Data Breach Notification', 'In the event of a security incident affecting Licensee Data, Sonar shall notify the Licensee without undue delay and provide reasonable cooperation in the Licensee\'s investigation and response.'),

        ('Cross-Border Data Transfers', 'If any Licensee Data is transferred outside Zimbabwe, such transfer shall be conducted in compliance with applicable data protection laws and subject to appropriate safeguards.'),
    ]

    for i, (title, content) in enumerate(data_sections):
        para = doc.add_paragraph()
        para.add_run(f'12.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 13: COMPLIANCE AND AUDIT
    # ============================================
    doc.add_heading('ARTICLE 13: COMPLIANCE AND AUDIT', level=1)

    audit_sections = [
        ('Compliance', 'The Licensee shall use the Licensed Software in compliance with all applicable laws and regulations, including but not limited to laws relating to data protection, export control, anti-corruption, and financial services (if applicable).'),

        ('Records', 'The Licensee shall maintain accurate records of: (a) the number of Named Users and Concurrent Users accessing the Licensed Software; (b) the locations where the Licensed Software is installed; and (c) any copies of the Licensed Software made for backup or disaster recovery purposes.'),

        ('Audit Rights', 'Sonar or its authorized representative shall have the right, upon thirty (30) days\' prior written notice, to audit the Licensee\'s use of the Licensed Software to verify compliance with this Agreement. Audits shall be conducted during normal business hours and shall not unreasonably interfere with the Licensee\'s operations. The Licensee shall provide reasonable cooperation and access to relevant records and systems.'),

        ('Audit Frequency', 'Sonar shall not conduct more than one (1) audit in any twelve (12) month period, unless a prior audit reveals a material breach, in which case Sonar may conduct additional audits as reasonably necessary.'),

        ('Audit Costs', 'Sonar shall bear the costs of any audit. However, if an audit reveals that the Licensee has exceeded the licensed usage by more than 5% or has otherwise materially breached this Agreement, the Licensee shall reimburse Sonar for the reasonable costs of the audit.'),

        ('Under-Licensing Remediation', 'If an audit reveals that the Licensee has exceeded the licensed usage, the Licensee shall promptly pay the applicable fees for the excess usage at Sonar\'s then-current list prices, plus interest from the date the excess usage began. Repeated or intentional under-licensing may result in termination of this Agreement.'),
    ]

    for i, (title, content) in enumerate(audit_sections):
        para = doc.add_paragraph()
        para.add_run(f'13.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 14: TERM AND TERMINATION
    # ============================================
    doc.add_heading('ARTICLE 14: TERM AND TERMINATION', level=1)

    term_sections = [
        ('Initial Term', 'This Agreement shall commence on the Effective Date and shall continue for the Initial Term specified in Schedule B, unless earlier terminated in accordance with this Article.'),

        ('Renewal', '''
For Perpetual Licenses: This Agreement shall continue indefinitely, subject to termination rights set forth herein. Annual Maintenance Services shall automatically renew for successive one (1) year periods unless either Party provides written notice of non-renewal at least sixty (60) days prior to the end of the then-current term.

For Subscription Licenses: The subscription shall automatically renew for successive periods equal to the Initial Term unless either Party provides written notice of non-renewal at least sixty (60) days prior to the end of the then-current term.'''),

        ('Termination for Convenience', 'Either Party may terminate this Agreement for convenience upon ninety (90) days\' prior written notice. In case of termination for convenience by Sonar, Sonar shall refund a pro-rata portion of any prepaid fees for the period following termination.'),

        ('Termination for Breach', '''Either Party may terminate this Agreement immediately upon written notice if:

(a) The other Party materially breaches this Agreement and fails to cure such breach within thirty (30) days of receiving written notice thereof;
(b) The other Party becomes insolvent, makes an assignment for the benefit of creditors, or becomes subject to bankruptcy, receivership, or similar proceedings;
(c) The other Party ceases to conduct business in the ordinary course.'''),

        ('Termination by Sonar', '''Sonar may terminate this Agreement immediately upon written notice if:

(a) The Licensee fails to pay any fees when due and does not cure such failure within fifteen (15) days of written notice;
(b) The Licensee breaches any provision of Article 3 (License Restrictions);
(c) The Licensee breaches any provision of Article 11 (Confidentiality);
(d) Sonar reasonably believes that the Licensed Software is being used in violation of applicable laws.'''),

        ('Suspension', 'Sonar may suspend the Licensee\'s access to the Licensed Software or Maintenance Services if the Licensee fails to pay any fees when due, pending resolution of the payment issue.'),
    ]

    for i, (title, content) in enumerate(term_sections):
        para = doc.add_paragraph()
        para.add_run(f'14.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 15: EFFECTS OF TERMINATION
    # ============================================
    doc.add_heading('ARTICLE 15: EFFECTS OF TERMINATION', level=1)

    effects_sections = [
        ('Cessation of Use', 'Upon termination or expiration of this Agreement for any reason, the Licensee shall immediately: (a) cease all use of the Licensed Software; (b) uninstall and delete all copies of the Licensed Software from all systems; and (c) return or destroy all Documentation and other materials provided by Sonar.'),

        ('Certification', 'Within thirty (30) days of termination, the Licensee shall provide Sonar with a written certification, signed by an authorized officer, confirming that all copies of the Licensed Software have been destroyed and all Documentation has been returned or destroyed.'),

        ('Data Export', 'Prior to termination, the Licensee may export its data from the Licensed Software. Sonar shall provide reasonable assistance with data export upon request. After termination, Sonar has no obligation to retain or provide access to Licensee Data.'),

        ('Accrued Obligations', 'Termination shall not relieve either Party of obligations that accrued prior to termination, including payment obligations for fees accrued through the termination date.'),

        ('Survival', 'The following provisions shall survive termination of this Agreement: Article 1 (Definitions), Article 9 (Limitation of Liability), Article 10 (Intellectual Property), Article 11 (Confidentiality), Article 15 (Effects of Termination), Article 16 (Dispute Resolution), and Article 17 (General Provisions).'),
    ]

    for i, (title, content) in enumerate(effects_sections):
        para = doc.add_paragraph()
        para.add_run(f'15.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 16: DISPUTE RESOLUTION
    # ============================================
    doc.add_heading('ARTICLE 16: DISPUTE RESOLUTION', level=1)

    dispute_sections = [
        ('Governing Law', 'This Agreement shall be governed by and construed in accordance with the laws of the Republic of Zimbabwe, without regard to its conflict of laws principles, except that matters relating to intellectual property rights may be governed by applicable international treaties and conventions.'),

        ('Negotiation', 'In the event of any dispute arising out of or relating to this Agreement, the Parties shall first attempt to resolve the dispute through good faith negotiations between senior representatives with authority to settle the dispute. Such negotiations shall commence within ten (10) business days of written notice of the dispute.'),

        ('Mediation', 'If the dispute is not resolved through negotiation within thirty (30) days, either Party may refer the dispute to mediation administered by an agreed mediator or, failing agreement, by a mediator appointed by the Law Society of Zimbabwe. The mediation shall be conducted in Harare, Zimbabwe, in the English language.'),

        ('Arbitration', '''If the dispute is not resolved through mediation within sixty (60) days of the mediation request, either Party may refer the dispute to binding arbitration as follows:

(a) The arbitration shall be administered in accordance with the Arbitration Act [Chapter 7:15] of Zimbabwe;
(b) The arbitration shall be conducted in Harare, Zimbabwe, in the English language;
(c) The arbitral tribunal shall consist of one (1) arbitrator agreed upon by the Parties or, failing agreement, appointed by the Arbitration Centre of Zimbabwe;
(d) The award of the arbitrator shall be final and binding and may be entered as a judgment in any court of competent jurisdiction.'''),

        ('Injunctive Relief', 'Notwithstanding the foregoing, either Party may seek injunctive or other equitable relief from any court of competent jurisdiction to protect its intellectual property rights or confidential information without first resorting to negotiation, mediation, or arbitration.'),

        ('Costs', 'Each Party shall bear its own costs and attorney\'s fees in connection with any dispute resolution proceedings. The costs of mediation and arbitration shall be shared equally by the Parties, unless the arbitrator determines otherwise.'),

        ('Continued Performance', 'Unless otherwise agreed, the Parties shall continue to perform their obligations under this Agreement during the pendency of any dispute, except for obligations that are the subject of the dispute.'),
    ]

    for i, (title, content) in enumerate(dispute_sections):
        para = doc.add_paragraph()
        para.add_run(f'16.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 17: GENERAL PROVISIONS
    # ============================================
    doc.add_heading('ARTICLE 17: GENERAL PROVISIONS', level=1)

    general_sections = [
        ('Entire Agreement', 'This Agreement, together with all Schedules attached hereto, constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, understandings, and communications, whether written or oral.'),

        ('Amendment', 'This Agreement may not be amended or modified except by a written instrument signed by authorized representatives of both Parties. No purchase order, acknowledgment, or other business form shall modify or supplement this Agreement.'),

        ('Waiver', 'The failure of either Party to enforce any provision of this Agreement shall not constitute a waiver of that Party\'s right to enforce that provision or any other provision. A waiver of any breach shall not be deemed a waiver of any subsequent breach.'),

        ('Severability', 'If any provision of this Agreement is held to be invalid or unenforceable, the remaining provisions shall continue in full force and effect. The invalid provision shall be modified to the minimum extent necessary to make it valid and enforceable while preserving the Parties\' intent.'),

        ('Assignment', 'The Licensee may not assign or transfer this Agreement or any rights hereunder without Sonar\'s prior written consent. Any attempted assignment without consent shall be void. Sonar may assign this Agreement to any successor in interest or affiliate upon notice to the Licensee.'),

        ('Notices', '''All notices under this Agreement shall be in writing and shall be deemed given when:

(a) Delivered personally;
(b) Sent by registered mail, return receipt requested, to the address set forth herein;
(c) Sent by reputable international courier service; or
(d) Sent by email with acknowledgment of receipt.

Notices shall be sent to the addresses specified on the signature page or such other addresses as may be designated in writing.'''),

        ('Force Majeure', 'Neither Party shall be liable for any failure or delay in performance due to causes beyond its reasonable control, including but not limited to acts of God, natural disasters, war, terrorism, civil unrest, government actions, power failures, internet disruptions, or pandemics. The affected Party shall provide prompt notice and use reasonable efforts to mitigate the impact.'),

        ('Independent Contractors', 'The Parties are independent contractors. Nothing in this Agreement creates a partnership, joint venture, agency, employment, or fiduciary relationship between the Parties.'),

        ('No Third-Party Beneficiaries', 'This Agreement is for the sole benefit of the Parties and their permitted successors and assigns. Nothing in this Agreement confers any rights or remedies upon any third party.'),

        ('Export Controls', 'The Licensee shall comply with all applicable export control laws and regulations. The Licensee shall not export, re-export, or transfer the Licensed Software in violation of any applicable export restrictions.'),

        ('Anti-Corruption', 'Each Party represents that it has not made and will not make any payments, gifts, or other inducements to any government official, political party, or other person in connection with this Agreement that would violate applicable anti-corruption laws.'),

        ('Counterparts', 'This Agreement may be executed in counterparts, each of which shall be deemed an original and all of which together shall constitute one and the same instrument. Electronic signatures shall be valid and binding.'),

        ('Language', 'This Agreement is executed in the English language. In the event of any translation, the English version shall prevail.'),

        ('Interpretation', 'The headings in this Agreement are for convenience only and shall not affect interpretation. The terms "include," "including," and similar terms shall be construed as if followed by "without limitation." References to Articles and Schedules are to those of this Agreement.'),
    ]

    for i, (title, content) in enumerate(general_sections):
        para = doc.add_paragraph()
        para.add_run(f'17.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_page_break()

    # ============================================
    # SCHEDULES
    # ============================================
    doc.add_heading('SCHEDULES', level=1)

    # Schedule A
    doc.add_heading('SCHEDULE A: LICENSED SOFTWARE SPECIFICATIONS', level=2)

    spec_table = doc.add_table(rows=12, cols=2)
    spec_table.style = 'Table Grid'

    spec_data = [
        ['Item', 'Specification'],
        ['Licensed Software', 'Sonar Workflow System'],
        ['Version', '[Version Number]'],
        ['Licensed Edition', '[Standard / Professional / Enterprise]'],
        ['Named User Limit', '[Number] Named Users'],
        ['Concurrent User Limit', '[Number] Concurrent Users'],
        ['Licensed Modules', '[List of Modules]'],
        ['Site(s)', '[Physical Address(es) in Zimbabwe]'],
        ['Server Environment', '[On-premise / Cloud / Hybrid]'],
        ['Operating System', '[Windows Server / Linux]'],
        ['Database', '[PostgreSQL / Oracle / SQL Server]'],
        ['Integration Requirements', '[List any required integrations]'],
    ]

    for row_idx, row_data in enumerate(spec_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = spec_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

    # Schedule B
    doc.add_heading('SCHEDULE B: PRICING AND PAYMENT SCHEDULE', level=2)

    price_table = doc.add_table(rows=9, cols=3)
    price_table.style = 'Table Grid'

    price_data = [
        ['Fee Type', 'Amount (USD)', 'Payment Due'],
        ['License Fee', '[Amount]', 'Upon execution'],
        ['Implementation Fee', '[Amount]', '50% upon execution, 50% upon acceptance'],
        ['Training Fee (if additional)', '[Amount]', 'Upon completion of training'],
        ['Year 1 Maintenance', 'Included / [Amount]', 'N/A / Upon execution'],
        ['Annual Maintenance (subsequent)', '[Amount]', 'Annually on anniversary'],
        ['Additional Users (if any)', '[Amount per user]', 'Upon addition'],
        ['Custom Development (if any)', '[Hourly rate]', 'Monthly in arrears'],
        ['TOTAL INITIAL INVESTMENT', '[Total Amount]', 'Per above schedule'],
    ]

    for row_idx, row_data in enumerate(price_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = price_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0 or row_idx == 8:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E2EFDA')

    doc.add_paragraph()

    payment_info = doc.add_paragraph()
    payment_info.add_run('Payment Instructions:').bold = True

    doc.add_paragraph('Bank Name: ___________________________________')
    doc.add_paragraph('Account Name: Sonar Microsystems Pvt. Ltd.')
    doc.add_paragraph('Account Number: ___________________________________')
    doc.add_paragraph('SWIFT Code: ___________________________________')
    doc.add_paragraph('Reference: SONAR-LIC-ZW-[Agreement Number]')

    doc.add_paragraph()

    # Schedule C
    doc.add_heading('SCHEDULE C: SERVICE LEVEL AGREEMENT', level=2)

    sla_table = doc.add_table(rows=5, cols=4)
    sla_table.style = 'Table Grid'

    sla_data = [
        ['Priority', 'Description', 'Response Time', 'Resolution Target'],
        ['Priority 1 (Critical)', 'System down, no workaround', '4 hours', '24 hours'],
        ['Priority 2 (High)', 'Major function impaired', '8 business hours', '48 hours'],
        ['Priority 3 (Medium)', 'Function impaired, workaround exists', '2 business days', '5 business days'],
        ['Priority 4 (Low)', 'Minor issues, questions', '3 business days', '10 business days'],
    ]

    for row_idx, row_data in enumerate(sla_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = sla_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'FCE4D6')

    doc.add_paragraph()

    support_contact = doc.add_paragraph()
    support_contact.add_run('Support Contact Information:').bold = True

    doc.add_paragraph('Email: support@sonarmicrosystems.com')
    doc.add_paragraph('Phone: [Support Phone Number]')
    doc.add_paragraph('Portal: [Support Portal URL]')
    doc.add_paragraph('Business Hours (CAT): Monday - Friday, 8:00 AM - 5:00 PM')

    doc.add_paragraph()

    # Schedule D
    doc.add_heading('SCHEDULE D: AUTHORIZED USERS', level=2)

    auth_intro = doc.add_paragraph()
    auth_intro.add_run('Initial Authorized Users and Technical Contacts:')

    doc.add_paragraph()

    auth_table = doc.add_table(rows=8, cols=4)
    auth_table.style = 'Table Grid'

    auth_headers = ['Name', 'Role/Title', 'Email', 'User Type']
    for i, header in enumerate(auth_headers):
        cell = auth_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # Technical contacts row
    auth_table.rows[1].cells[0].text = '[Primary Technical Contact]'
    auth_table.rows[1].cells[3].text = 'Admin'
    auth_table.rows[2].cells[0].text = '[Secondary Technical Contact]'
    auth_table.rows[2].cells[3].text = 'Admin'

    doc.add_paragraph()

    # ============================================
    # SIGNATURE PAGE
    # ============================================
    doc.add_page_break()
    doc.add_heading('SIGNATURE PAGE', level=1)

    sig_intro = doc.add_paragraph()
    sig_intro.add_run('IN WITNESS WHEREOF, the Parties have executed this Software License Agreement as of the Effective Date.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature table
    sig_table = doc.add_table(rows=1, cols=2)

    # Licensor signature
    cell1 = sig_table.rows[0].cells[0]
    p1 = cell1.add_paragraph()
    p1.add_run('LICENSOR:').bold = True
    cell1.add_paragraph()
    p1a = cell1.add_paragraph()
    p1a.add_run('SONAR MICROSYSTEMS PVT. LTD.').bold = True
    cell1.add_paragraph()
    cell1.add_paragraph()
    cell1.add_paragraph('_________________________________')
    cell1.add_paragraph('Authorized Signatory')
    cell1.add_paragraph()
    cell1.add_paragraph('Name: _________________________')
    cell1.add_paragraph('Title: __________________________')
    cell1.add_paragraph('Date: __________________________')
    cell1.add_paragraph()
    cell1.add_paragraph('_________________________________')
    cell1.add_paragraph('Witness')
    cell1.add_paragraph('Name: _________________________')

    # Licensee signature
    cell2 = sig_table.rows[0].cells[1]
    p2 = cell2.add_paragraph()
    p2.add_run('LICENSEE:').bold = True
    cell2.add_paragraph()
    p2a = cell2.add_paragraph()
    p2a.add_run('[LICENSEE COMPANY NAME]').bold = True
    cell2.add_paragraph()
    cell2.add_paragraph()
    cell2.add_paragraph('_________________________________')
    cell2.add_paragraph('Authorized Signatory')
    cell2.add_paragraph()
    cell2.add_paragraph('Name: _________________________')
    cell2.add_paragraph('Title: __________________________')
    cell2.add_paragraph('Date: __________________________')
    cell2.add_paragraph()
    cell2.add_paragraph('_________________________________')
    cell2.add_paragraph('Witness')
    cell2.add_paragraph('Name: _________________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Company Seal area
    seal_para = doc.add_paragraph()
    seal_para.add_run('COMPANY SEALS:').bold = True
    seal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    seal_table = doc.add_table(rows=1, cols=2)
    seal_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    seal1 = seal_table.rows[0].cells[0]
    seal1.add_paragraph()
    seal1.add_paragraph()
    seal1.add_paragraph('[SONAR COMPANY SEAL]')
    seal1.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal1.add_paragraph()
    seal1.add_paragraph()

    seal2 = seal_table.rows[0].cells[1]
    seal2.add_paragraph()
    seal2.add_paragraph()
    seal2.add_paragraph('[LICENSEE COMPANY SEAL]')
    seal2.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal2.add_paragraph()
    seal2.add_paragraph()

    doc.add_paragraph()

    # Footer
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer1 = doc.add_paragraph()
    footer1.add_run('SONAR WORKFLOW SYSTEM - SOFTWARE LICENSE AGREEMENT').bold = True
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer2 = doc.add_paragraph()
    footer2.add_run('CONFIDENTIAL - FOR AUTHORIZED USE ONLY')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer3 = doc.add_paragraph()
    footer3.add_run(f'© {datetime.now().year} Sonar Microsystems Pvt. Ltd. All Rights Reserved.')
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save('Sonar Workflow System License.docx')
    print('License Agreement created successfully: Sonar Workflow System License.docx')

if __name__ == '__main__':
    create_license_agreement()
