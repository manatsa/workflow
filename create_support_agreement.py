#!/usr/bin/env python3
"""
Script to generate a Support and Maintenance Agreement for Sonar Workflow System
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_support_agreement():
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================
    # COVER PAGE
    # ============================================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('SONAR WORKFLOW SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('SUPPORT AND MAINTENANCE AGREEMENT', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.add_run('Version 1.0').bold = True
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Parties
    cover_info = doc.add_paragraph()
    cover_info.add_run('Between').italic = True
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    provider = doc.add_paragraph()
    provider.add_run('SONAR MICROSYSTEMS PVT. LTD.').bold = True
    provider.alignment = WD_ALIGN_PARAGRAPH.CENTER

    provider_label = doc.add_paragraph()
    provider_label.add_run('(Service Provider)').italic = True
    provider_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    and_text = doc.add_paragraph()
    and_text.add_run('and').italic = True
    and_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    customer = doc.add_paragraph()
    customer.add_run('[CUSTOMER NAME]').bold = True
    customer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    customer_label = doc.add_paragraph()
    customer_label.add_run('(Customer)').italic = True
    customer_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info
    doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc_info = doc.add_paragraph()
    doc_info.add_run('Agreement Reference: ').bold = True
    doc_info.add_run('SONAR-SMA-___________')
    doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_info = doc.add_paragraph()
    date_info.add_run('Effective Date: ').bold = True
    date_info.add_run('_______________, 20____')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    license_ref = doc.add_paragraph()
    license_ref.add_run('Related License Agreement: ').bold = True
    license_ref.add_run('SONAR-LIC-___________')
    license_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.add_run('CONFIDENTIAL AND PROPRIETARY').bold = True
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER

    copyright_notice = doc.add_paragraph()
    copyright_notice.add_run(f'© {datetime.now().year} Sonar Microsystems Pvt. Ltd. All Rights Reserved.')
    copyright_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    doc.add_heading('TABLE OF CONTENTS', level=1)

    toc_items = [
        ('Article 1: Definitions', 3),
        ('Article 2: Scope of Services', 5),
        ('Article 3: Support Services', 6),
        ('Article 4: Maintenance Services', 9),
        ('Article 5: Service Levels and Response Times', 11),
        ('Article 6: Escalation Procedures', 14),
        ('Article 7: Customer Responsibilities', 15),
        ('Article 8: Service Credits', 17),
        ('Article 9: Fees and Payment', 18),
        ('Article 10: Term and Renewal', 19),
        ('Article 11: Termination', 20),
        ('Article 12: Exclusions and Limitations', 21),
        ('Article 13: Confidentiality', 22),
        ('Article 14: Limitation of Liability', 23),
        ('Article 15: General Provisions', 24),
        ('EXHIBIT A: Service Level Agreement (SLA)', 26),
        ('EXHIBIT B: Support Tiers and Pricing', 28),
        ('EXHIBIT C: Escalation Matrix', 30),
        ('EXHIBIT D: Authorized Contacts', 31),
        ('EXHIBIT E: System Configuration', 32),
        ('SIGNATURE PAGE', 33),
    ]

    for item, page in toc_items:
        toc_para = doc.add_paragraph()
        if item.startswith('EXHIBIT') or item.startswith('SIGNATURE'):
            toc_para.add_run(item).bold = True
        else:
            toc_para.add_run(item)
        toc_para.add_run('.' * 3 + f' {page}')

    doc.add_page_break()

    # ============================================
    # PREAMBLE
    # ============================================
    doc.add_heading('SUPPORT AND MAINTENANCE AGREEMENT', level=1)

    preamble = doc.add_paragraph()
    preamble.add_run('This Support and Maintenance Agreement ("Agreement") is entered into as of the Effective Date set forth above by and between:')

    doc.add_paragraph()

    # Provider
    provider_para = doc.add_paragraph()
    provider_para.add_run('SERVICE PROVIDER: ').bold = True
    provider_para.add_run('SONAR MICROSYSTEMS PVT. LTD.')
    provider_para.add_run(', a company incorporated under the laws of India, having its registered office at ___________________________________ (hereinafter referred to as "')
    provider_para.add_run('Sonar').bold = True
    provider_para.add_run('" or "')
    provider_para.add_run('Provider').bold = True
    provider_para.add_run('");')

    doc.add_paragraph()

    # Customer
    customer_para = doc.add_paragraph()
    customer_para.add_run('CUSTOMER: ').bold = True
    customer_para.add_run('___________________________________')
    customer_para.add_run(', a company incorporated under the laws of _______________, having its registered office at ___________________________________ (hereinafter referred to as "')
    customer_para.add_run('Customer').bold = True
    customer_para.add_run('");')

    doc.add_paragraph()

    collective = doc.add_paragraph()
    collective.add_run('(Sonar and the Customer are hereinafter individually referred to as a "Party" and collectively as the "Parties")')
    collective.italic = True

    doc.add_paragraph()

    # Recitals
    doc.add_heading('RECITALS', level=2)

    recitals = [
        'WHEREAS, the Customer has licensed the Sonar Workflow System software from Sonar pursuant to a Software License Agreement (the "License Agreement");',
        'WHEREAS, the Customer desires to obtain ongoing support and maintenance services for the Sonar Workflow System to ensure optimal performance, security, and reliability;',
        'WHEREAS, Sonar is willing to provide such support and maintenance services subject to the terms and conditions set forth in this Agreement;',
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:'
    ]

    for recital in recitals:
        doc.add_paragraph(recital, style='List Number')

    doc.add_paragraph()

    # ============================================
    # ARTICLE 1: DEFINITIONS
    # ============================================
    doc.add_heading('ARTICLE 1: DEFINITIONS', level=1)

    definitions = [
        ('"Authorized Contact"', 'means an individual designated by the Customer who is authorized to submit Support Requests and communicate with Sonar regarding support matters, as listed in Exhibit D.'),

        ('"Business Day"', 'means Monday through Friday, excluding public holidays observed by Sonar at its principal place of business.'),

        ('"Business Hours"', 'means 8:00 AM to 6:00 PM in the Customer\'s local time zone on Business Days, unless otherwise specified in Exhibit A.'),

        ('"Critical Error"', 'means a Priority 1 issue where the Sonar Workflow System is completely inoperable or a critical business function is unavailable with no workaround, causing severe and immediate business impact.'),

        ('"Documentation"', 'means the user manuals, technical documentation, release notes, and other written materials provided by Sonar in connection with the Sonar Workflow System.'),

        ('"Effective Date"', 'means the date on which this Agreement becomes effective, as specified on the cover page.'),

        ('"Error"', 'means any failure of the Sonar Workflow System to conform in any material respect to the functional specifications set forth in the Documentation.'),

        ('"Extended Support Hours"', 'means support availability beyond standard Business Hours, as specified in the Customer\'s selected Support Tier.'),

        ('"Fix"', 'means a modification, workaround, or patch that resolves an Error.'),

        ('"Hotfix"', 'means an emergency patch released outside the normal update cycle to address a Critical Error or security vulnerability.'),

        ('"License Agreement"', 'means the Software License Agreement between Sonar and the Customer governing the use of the Sonar Workflow System.'),

        ('"Maintenance Fee"', 'means the annual fee payable by the Customer for the support and maintenance services provided under this Agreement.'),

        ('"Maintenance Window"', 'means the scheduled time period during which Sonar may perform system maintenance, typically occurring during off-peak hours with advance notice to the Customer.'),

        ('"Major Release"', 'means a new version of the Sonar Workflow System that includes significant new features, functionality, or architectural changes, designated by a change in the version number before the first decimal point (e.g., Version 1.x to Version 2.0).'),

        ('"Minor Release"', 'means a new version of the Sonar Workflow System that includes incremental enhancements, improvements, and bug fixes, designated by a change in the version number after the first decimal point (e.g., Version 1.1 to Version 1.2).'),

        ('"Patch"', 'means a software update that addresses specific Errors, security vulnerabilities, or minor issues without introducing new features.'),

        ('"Resolution"', 'means the successful correction of an Error or implementation of a permanent Fix.'),

        ('"Resolution Time"', 'means the elapsed time from acknowledgment of a Support Request to the delivery of a Resolution or acceptable workaround.'),

        ('"Response Time"', 'means the elapsed time from the submission of a Support Request to the initial acknowledgment by Sonar support personnel.'),

        ('"Service Credits"', 'means credits issued to the Customer in accordance with Article 8 for failure to meet Service Level commitments.'),

        ('"Service Level Agreement" or "SLA"', 'means the service level commitments set forth in Exhibit A, including Response Times, Resolution Times, and system availability targets.'),

        ('"Sonar Workflow System"', 'means the software licensed to the Customer under the License Agreement, including all Patches, Minor Releases, and Major Releases provided during the term of this Agreement.'),

        ('"Support Portal"', 'means Sonar\'s online customer support portal for submitting and tracking Support Requests.'),

        ('"Support Request"', 'means a request for technical assistance, Error resolution, or other support submitted by an Authorized Contact through the designated support channels.'),

        ('"Support Tier"', 'means the level of support services selected by the Customer, as described in Exhibit B.'),

        ('"Update"', 'means any Patch, Minor Release, or other software modification provided by Sonar to address Errors, improve performance, or enhance security.'),

        ('"Upgrade"', 'means a Major Release that includes significant new functionality or features.'),

        ('"Workaround"', 'means a temporary solution or alternative procedure that enables the Customer to continue using the Sonar Workflow System despite an Error, pending a permanent Fix.'),
    ]

    for i, (term, definition) in enumerate(definitions):
        para = doc.add_paragraph()
        para.add_run(f'1.{i + 1} ').bold = True
        para.add_run(term).bold = True
        para.add_run(f' {definition}')

    doc.add_paragraph()

    # ============================================
    # ARTICLE 2: SCOPE OF SERVICES
    # ============================================
    doc.add_heading('ARTICLE 2: SCOPE OF SERVICES', level=1)

    scope_sections = [
        ('Overview', '''Subject to the terms and conditions of this Agreement and payment of the applicable Maintenance Fee, Sonar shall provide the Customer with the following support and maintenance services for the Sonar Workflow System:

(a) Technical Support Services as described in Article 3;
(b) Maintenance Services as described in Article 4;
(c) Service Level commitments as described in Article 5 and Exhibit A;
(d) Access to the Support Portal and knowledge base resources;
(e) Communication regarding product updates, security advisories, and end-of-life notices.'''),

        ('Covered Software', 'This Agreement covers support and maintenance for the Sonar Workflow System software and modules licensed under the License Agreement, as detailed in Exhibit E. Only the current and immediately prior version of the Sonar Workflow System are eligible for full support. Older versions may receive limited support or may be subject to an upgrade requirement.'),

        ('Support Tier', 'The level of support services provided under this Agreement shall be determined by the Support Tier selected by the Customer, as specified in Exhibit B. Different Support Tiers offer varying levels of availability, response times, and additional services.'),

        ('Geographic Coverage', 'Support services shall be provided to the Customer at the locations specified in Exhibit E. Support for additional locations may require separate arrangements and additional fees.'),

        ('Relationship to License Agreement', 'This Agreement is supplemental to the License Agreement. In the event of any conflict between this Agreement and the License Agreement, the terms of this Agreement shall prevail with respect to support and maintenance matters. All other terms of the License Agreement remain in full force and effect.'),
    ]

    for i, (title, content) in enumerate(scope_sections):
        para = doc.add_paragraph()
        para.add_run(f'2.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 3: SUPPORT SERVICES
    # ============================================
    doc.add_heading('ARTICLE 3: SUPPORT SERVICES', level=1)

    support_sections = [
        ('Technical Support', '''Sonar shall provide technical support services to assist the Customer with:

(a) Diagnosis and resolution of Errors in the Sonar Workflow System;
(b) Troubleshooting performance issues and system problems;
(c) Guidance on software configuration and optimization;
(d) Clarification of Documentation and product functionality;
(e) Assistance with software installation and Updates;
(f) Resolution of integration issues with supported third-party systems;
(g) Guidance on best practices and recommended configurations.'''),

        ('Support Channels', '''The Customer may submit Support Requests through the following channels:

(a) Support Portal: The primary method for submitting and tracking Support Requests. Available 24/7 for request submission.

(b) Email Support: support@sonarmicrosystems.com - Monitored during Business Hours (or Extended Support Hours for applicable Support Tiers).

(c) Phone Support: Available for Priority 1 and Priority 2 issues during the hours specified in the Customer's Support Tier. Phone numbers are provided in Exhibit A.

(d) Live Chat: Available through the Support Portal during Business Hours for quick questions and status updates.'''),

        ('Support Request Submission', '''Each Support Request must include the following information:

(a) Customer name and Agreement reference number;
(b) Name and contact information of the Authorized Contact submitting the request;
(c) Priority level (as defined in Article 5);
(d) Detailed description of the issue, including steps to reproduce if applicable;
(e) Error messages, screenshots, or log files as relevant;
(f) Business impact and urgency;
(g) Environment information (version, configuration, etc.);
(h) Any troubleshooting steps already attempted.

Incomplete Support Requests may result in delayed response while Sonar gathers necessary information.'''),

        ('Support Request Handling', '''Upon receipt of a Support Request, Sonar shall:

(a) Acknowledge the request within the Response Time applicable to the Priority level;
(b) Assign a unique ticket number for tracking purposes;
(c) Assess the Priority level and adjust if necessary based on the information provided;
(d) Assign the request to an appropriate support engineer;
(e) Investigate the issue and communicate findings to the Customer;
(f) Work toward Resolution or Workaround within the applicable Resolution Time targets;
(g) Keep the Customer informed of progress through regular status updates;
(h) Document the resolution and close the ticket upon Customer confirmation.'''),

        ('Remote Support', '''To facilitate efficient support, the Customer agrees to:

(a) Provide Sonar with secure remote access to the Sonar Workflow System environment when reasonably required for diagnosis and resolution;
(b) Ensure that remote access complies with the Customer's security policies;
(c) Provide a designated contact to authorize and facilitate remote access sessions;
(d) Allow Sonar to collect diagnostic information necessary for troubleshooting.

All remote access sessions shall be conducted in accordance with Sonar's security protocols and the Customer's reasonable security requirements.'''),

        ('On-Site Support', 'On-site support is available for certain Support Tiers and may be requested for complex issues that cannot be resolved remotely. On-site support visits are subject to scheduling availability and may incur additional charges unless included in the Customer\'s Support Tier. Travel and accommodation expenses shall be borne by the Customer unless otherwise agreed.'),

        ('Knowledge Base and Self-Service', '''Sonar provides the following self-service resources:

(a) Online Knowledge Base: Searchable database of articles, FAQs, and troubleshooting guides;
(b) Documentation Library: Access to current and historical product documentation;
(c) Video Tutorials: Training videos on common tasks and features;
(d) Community Forum: Peer-to-peer discussion forum for sharing experiences and tips;
(e) Release Notes: Detailed information on each software release;
(f) Best Practices Guide: Recommendations for optimal configuration and usage.

The Customer is encouraged to consult self-service resources before submitting Support Requests for common issues.'''),

        ('Support Languages', 'Primary support is provided in English. Support in other languages may be available depending on the Support Tier and availability of multilingual support staff.'),
    ]

    for i, (title, content) in enumerate(support_sections):
        para = doc.add_paragraph()
        para.add_run(f'3.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 4: MAINTENANCE SERVICES
    # ============================================
    doc.add_heading('ARTICLE 4: MAINTENANCE SERVICES', level=1)

    maint_sections = [
        ('Software Updates', '''During the term of this Agreement, Sonar shall provide the Customer with:

(a) Patches: Bug fixes and minor corrections released to address specific Errors;
(b) Minor Releases: Incremental updates with enhancements, improvements, and accumulated fixes;
(c) Security Updates: Patches addressing security vulnerabilities and threats;
(d) Hotfixes: Emergency patches for Critical Errors or urgent security issues.

All Updates shall be made available through the Support Portal or delivered via the method specified by Sonar.'''),

        ('Major Releases (Upgrades)', '''Major Releases may include significant new features, architectural changes, or platform updates. For Customers with active support:

(a) Standard and Professional Tiers: Major Releases are available at a discounted upgrade fee (typically 50% of the current list price for new licenses);

(b) Enterprise Tier: Major Releases are included at no additional license fee, subject to applicable implementation and migration services fees;

(c) All Tiers: Sonar shall provide reasonable notice (typically 90 days) before releasing a Major Release and shall maintain support for the prior Major Release for at least twelve (12) months.'''),

        ('Update Schedule', '''Sonar maintains the following typical release schedule:

(a) Patches: As needed, typically monthly or when critical issues are identified;
(b) Minor Releases: Quarterly (approximately every 3 months);
(c) Major Releases: Annually or as determined by Sonar's product roadmap;
(d) Security Updates: As needed, with urgent security issues addressed within 48-72 hours of identification.

The Customer shall receive advance notification of scheduled releases and maintenance windows.'''),

        ('Update Installation', '''The responsibility for installing Updates depends on the deployment model:

(a) On-Premise Deployments: The Customer is responsible for installing Updates in their environment. Sonar shall provide installation instructions and Documentation. Installation assistance may be available as part of the Support Tier or as a professional service.

(b) Cloud/Hosted Deployments: Sonar shall install Updates during scheduled Maintenance Windows, with advance notice to the Customer. The Customer may request deferral of non-critical Updates for a reasonable period.

The Customer is strongly encouraged to install all Updates, particularly security Updates, in a timely manner. Failure to install Updates may affect Sonar's ability to provide effective support.'''),

        ('Version Support Policy', '''Sonar maintains the following version support policy:

(a) Current Version: Full support including all Patches, Updates, and enhancements;

(b) Current Version - 1 (Prior Version): Full support for twelve (12) months following the release of the Current Version, then limited support for an additional twelve (12) months;

(c) Older Versions: Limited or no support. Customers on older versions may be required to upgrade to receive support.

Limited support means Sonar will provide assistance with known issues and workarounds but will not develop new Fixes for the older version.'''),

        ('End of Life Policy', '''When a version of the Sonar Workflow System reaches end-of-life:

(a) Sonar shall provide at least six (6) months advance notice of the end-of-life date;
(b) Security Updates may continue for a limited period beyond end-of-life;
(c) The Customer shall be responsible for migrating to a supported version;
(d) Migration assistance may be available as a professional service.'''),

        ('Compatibility', 'Sonar shall use commercially reasonable efforts to maintain compatibility of Updates with the Customer\'s existing configuration, data, and integrations. However, some Updates may require changes to the Customer\'s environment. Sonar shall provide advance notice of any Updates that require significant changes.'),

        ('Documentation Updates', 'Sonar shall provide updated Documentation to reflect changes in each Update. Documentation updates shall be made available through the Support Portal concurrent with the release of each Update.'),
    ]

    for i, (title, content) in enumerate(maint_sections):
        para = doc.add_paragraph()
        para.add_run(f'4.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 5: SERVICE LEVELS AND RESPONSE TIMES
    # ============================================
    doc.add_heading('ARTICLE 5: SERVICE LEVELS AND RESPONSE TIMES', level=1)

    sla_sections = [
        ('Priority Classification', '''Support Requests shall be classified according to the following Priority levels:

PRIORITY 1 - CRITICAL
- The Sonar Workflow System is completely inoperable
- A critical business function is unavailable with no workaround
- Data integrity is at risk
- Security breach or vulnerability is actively being exploited
- Business Impact: Severe and immediate; operations halted

PRIORITY 2 - HIGH
- A major function is severely impaired
- Significant business impact but operations can continue at reduced capacity
- A workaround exists but is not sustainable long-term
- Business Impact: High; significant degradation of service

PRIORITY 3 - MEDIUM
- A function is impaired but a reasonable workaround exists
- Moderate business impact; operations can continue with inconvenience
- Non-critical feature is unavailable
- Business Impact: Moderate; acceptable short-term degradation

PRIORITY 4 - LOW
- Minor issues with minimal business impact
- Cosmetic defects or documentation errors
- General questions, how-to inquiries
- Enhancement requests or feature suggestions
- Business Impact: Minimal; normal operations continue'''),

        ('Response Time Commitments', 'Response Times vary by Support Tier and Priority level. The detailed Response Time commitments are set forth in Exhibit A. Response Time is measured from the time a properly submitted Support Request is received through the designated support channel to the time of initial acknowledgment by Sonar support personnel.'),

        ('Resolution Time Targets', 'Resolution Time targets vary by Support Tier and Priority level. The detailed Resolution Time targets are set forth in Exhibit A. Resolution Times are targets rather than guarantees, as actual resolution may depend on factors outside Sonar\'s control, including the complexity of the issue and Customer responsiveness.'),

        ('Priority Adjustment', 'Sonar reserves the right to adjust the Priority level of a Support Request based on its assessment of the actual business impact and technical severity. If Sonar adjusts the Priority level, it shall notify the Customer and provide an explanation. The Customer may dispute the adjustment through the escalation process.'),

        ('System Availability', '''For hosted/cloud deployments, Sonar commits to the following availability targets:

(a) Target Availability: 99.5% uptime measured monthly, excluding scheduled Maintenance Windows;
(b) Scheduled Maintenance: Planned maintenance shall occur during designated Maintenance Windows with at least 72 hours advance notice for routine maintenance and 24 hours for urgent maintenance;
(c) Emergency Maintenance: Emergency maintenance to address Critical Errors or security threats may occur with minimal notice.

Availability is calculated as: (Total Minutes in Month - Downtime Minutes) / Total Minutes in Month × 100%

Downtime excludes scheduled Maintenance Windows, Customer-caused outages, force majeure events, and periods where the Customer has not reported an issue.'''),

        ('Service Level Reporting', '''Sonar shall provide the Customer with periodic service level reports including:

(a) Summary of Support Requests by Priority and status;
(b) Response Time and Resolution Time performance against SLA targets;
(c) System availability metrics (for hosted deployments);
(d) Trend analysis and recommendations;
(e) Outstanding issues and action items.

Reports shall be provided monthly for Enterprise Tier customers and quarterly for other Support Tiers.'''),
    ]

    for i, (title, content) in enumerate(sla_sections):
        para = doc.add_paragraph()
        para.add_run(f'5.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Response Time Table
    doc.add_heading('Response Time Summary by Support Tier', level=2)

    rt_table = doc.add_table(rows=5, cols=4)
    rt_table.style = 'Table Grid'

    rt_headers = ['Priority', 'Standard Tier', 'Professional Tier', 'Enterprise Tier']
    for i, header in enumerate(rt_headers):
        cell = rt_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    rt_data = [
        ['Priority 1 (Critical)', '4 hours', '2 hours', '30 minutes'],
        ['Priority 2 (High)', '8 hours', '4 hours', '2 hours'],
        ['Priority 3 (Medium)', '2 Business Days', '1 Business Day', '4 hours'],
        ['Priority 4 (Low)', '3 Business Days', '2 Business Days', '1 Business Day'],
    ]

    for row_idx, row_data in enumerate(rt_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = rt_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # Resolution Time Table
    doc.add_heading('Resolution Time Targets by Support Tier', level=2)

    res_table = doc.add_table(rows=5, cols=4)
    res_table.style = 'Table Grid'

    res_headers = ['Priority', 'Standard Tier', 'Professional Tier', 'Enterprise Tier']
    for i, header in enumerate(res_headers):
        cell = res_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')

    res_data = [
        ['Priority 1 (Critical)', '24 hours', '8 hours', '4 hours'],
        ['Priority 2 (High)', '72 hours', '24 hours', '8 hours'],
        ['Priority 3 (Medium)', '10 Business Days', '5 Business Days', '3 Business Days'],
        ['Priority 4 (Low)', '30 Business Days', '15 Business Days', '10 Business Days'],
    ]

    for row_idx, row_data in enumerate(res_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = res_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # ============================================
    # ARTICLE 6: ESCALATION PROCEDURES
    # ============================================
    doc.add_heading('ARTICLE 6: ESCALATION PROCEDURES', level=1)

    esc_sections = [
        ('Escalation Triggers', '''A Support Request may be escalated under the following circumstances:

(a) Response Time or Resolution Time targets are at risk of being missed;
(b) The Customer believes the assigned Priority level is inappropriate;
(c) The issue is not making satisfactory progress toward resolution;
(d) The issue has broader business impact than initially assessed;
(e) The Customer requests management attention to the issue;
(f) Technical resources beyond standard support are required.'''),

        ('Escalation Levels', '''The escalation path consists of the following levels:

Level 1 - Support Team Lead
- Initial escalation point for delayed responses or stalled issues
- Authority to reassign resources and adjust priorities
- Target response: Within 2 hours

Level 2 - Support Manager
- Escalation for complex issues or customer dissatisfaction
- Authority to engage additional technical resources
- Target response: Within 4 hours

Level 3 - Director of Customer Success
- Escalation for critical business impact or service failures
- Authority to engage development and engineering teams
- Target response: Within 2 hours for Priority 1 issues

Level 4 - Executive Escalation
- Final escalation level for unresolved critical issues
- Engagement of senior leadership
- Target response: Immediate for business-critical situations

Escalation contact details are provided in Exhibit C.'''),

        ('Escalation Process', '''To escalate a Support Request:

(a) Contact the appropriate escalation level via the contact information in Exhibit C;
(b) Provide the Support Request ticket number and summary of the issue;
(c) Explain the reason for escalation and desired outcome;
(d) The escalation contact will acknowledge the escalation and take ownership;
(e) The escalation contact will provide a status update within the target response time;
(f) Escalation status will be documented in the Support Request ticket.'''),

        ('Automatic Escalation', 'Support Requests that exceed SLA thresholds may be automatically escalated according to the rules defined in Exhibit C. Automatic escalation ensures management visibility into service level performance.'),

        ('De-escalation', 'Once an escalated issue is resolved or the Customer is satisfied with the progress, the escalation may be closed. The Customer and Sonar escalation contact shall mutually agree that the escalation is resolved before closing.'),
    ]

    for i, (title, content) in enumerate(esc_sections):
        para = doc.add_paragraph()
        para.add_run(f'6.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 7: CUSTOMER RESPONSIBILITIES
    # ============================================
    doc.add_heading('ARTICLE 7: CUSTOMER RESPONSIBILITIES', level=1)

    cust_sections = [
        ('General Responsibilities', '''The Customer shall:

(a) Use the Sonar Workflow System in accordance with the Documentation and License Agreement;
(b) Maintain appropriate system administration and security practices;
(c) Install Updates in a timely manner, particularly security Updates;
(d) Maintain current backups of data and system configurations;
(e) Cooperate with Sonar in the diagnosis and resolution of issues;
(f) Provide accurate and complete information when submitting Support Requests;
(g) Respond promptly to Sonar's requests for information or access.'''),

        ('Authorized Contacts', '''The Customer shall:

(a) Designate a limited number of Authorized Contacts based on the Support Tier (Standard: 2, Professional: 4, Enterprise: 8);
(b) Ensure Authorized Contacts are appropriately trained and knowledgeable about the Sonar Workflow System;
(c) Keep the Authorized Contact list current and notify Sonar of changes;
(d) Ensure Support Requests are submitted only by Authorized Contacts;
(e) Provide 24/7 contact information for at least one Authorized Contact for Critical issues.'''),

        ('Environment Requirements', '''The Customer shall:

(a) Maintain hardware, operating systems, and network infrastructure that meet or exceed Sonar's published minimum requirements;
(b) Ensure adequate bandwidth and connectivity for remote support access;
(c) Maintain appropriate environmental controls (power, cooling, physical security);
(d) Keep supporting software (databases, web servers, etc.) at supported versions;
(e) Document and communicate any custom configurations or integrations.'''),

        ('Access for Support', '''The Customer shall:

(a) Provide Sonar with reasonable remote access to the Sonar Workflow System environment for support purposes;
(b) Ensure that remote access mechanisms are secure and compliant with Customer security policies;
(c) Designate personnel to facilitate and monitor remote access sessions;
(d) Provide access to relevant logs, configuration files, and diagnostic information;
(e) Arrange for on-site access when required for complex issues (with reasonable notice).'''),

        ('Issue Reporting', '''The Customer shall:

(a) Report Errors and issues promptly upon discovery;
(b) Provide detailed information to enable efficient diagnosis;
(c) Attempt basic troubleshooting before submitting Support Requests where practical;
(d) Accurately assess and assign Priority levels based on actual business impact;
(e) Notify Sonar of any changes in issue status or priority;
(f) Confirm resolution and provide feedback upon issue closure.'''),

        ('Testing Environment', 'The Customer is strongly encouraged to maintain a non-production testing environment for validating Updates before deployment to production. Sonar is not responsible for issues arising from deployment of Updates directly to production without prior testing.'),

        ('Data Protection', '''The Customer shall:

(a) Comply with all applicable data protection and privacy laws;
(b) Ensure appropriate consents are obtained for any personal data processed;
(c) Implement appropriate access controls and security measures;
(d) Notify Sonar of any data-related requirements or restrictions affecting support activities.'''),
    ]

    for i, (title, content) in enumerate(cust_sections):
        para = doc.add_paragraph()
        para.add_run(f'7.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 8: SERVICE CREDITS
    # ============================================
    doc.add_heading('ARTICLE 8: SERVICE CREDITS', level=1)

    credit_sections = [
        ('Eligibility', 'Service Credits are available to Customers on Professional and Enterprise Support Tiers when Sonar fails to meet the SLA commitments set forth in Exhibit A. Standard Tier customers are not eligible for Service Credits.'),

        ('Calculation', '''Service Credits shall be calculated as follows:

(a) Response Time Failure: If Sonar fails to meet the Response Time commitment for a Support Request, the Customer shall receive a credit equal to 2% of the monthly Maintenance Fee for each missed Response Time, up to a maximum of 10% per month.

(b) System Availability Failure (hosted deployments only): If system availability falls below the committed level:
   - 99.0% - 99.5% availability: 5% credit of monthly fee
   - 98.0% - 99.0% availability: 10% credit of monthly fee
   - 95.0% - 98.0% availability: 20% credit of monthly fee
   - Below 95.0% availability: 30% credit of monthly fee

(c) Maximum Credit: Total Service Credits in any month shall not exceed 30% of the monthly Maintenance Fee for that month.'''),

        ('Exclusions', '''Service Credits shall not apply to failures caused by:

(a) Actions or inactions of the Customer or its users;
(b) Failure of Customer equipment, software, or network connectivity;
(c) Force majeure events;
(d) Scheduled Maintenance Windows;
(e) Customer's failure to meet its responsibilities under this Agreement;
(f) Third-party services or software not provided by Sonar;
(g) Customer's use of the software in violation of the Documentation or License Agreement.'''),

        ('Claiming Credits', '''To claim Service Credits:

(a) The Customer must submit a credit request within thirty (30) days of the incident;
(b) The request must include the Support Request ticket number(s) and description of the SLA failure;
(c) Sonar shall review the request and respond within fifteen (15) Business Days;
(d) If approved, credits shall be applied to the next invoice or refunded if no further payments are due.'''),

        ('Sole Remedy', 'Service Credits are the Customer\'s sole and exclusive remedy for Sonar\'s failure to meet SLA commitments. Service Credits are not a penalty and do not limit Sonar\'s liability for other breaches of this Agreement.'),
    ]

    for i, (title, content) in enumerate(credit_sections):
        para = doc.add_paragraph()
        para.add_run(f'8.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 9: FEES AND PAYMENT
    # ============================================
    doc.add_heading('ARTICLE 9: FEES AND PAYMENT', level=1)

    fee_sections = [
        ('Maintenance Fee', 'The Customer shall pay the annual Maintenance Fee specified in Exhibit B for the selected Support Tier. The Maintenance Fee is calculated as a percentage of the current list price of the licensed software (typically 18-22% depending on Support Tier) or as a fixed annual amount as specified in Exhibit B.'),

        ('Payment Terms', '''
(a) The Maintenance Fee is due annually in advance on the anniversary of the Effective Date;
(b) Sonar shall issue an invoice at least thirty (30) days prior to the due date;
(c) Payment is due within thirty (30) days of the invoice date;
(d) All payments shall be made in the currency specified in Exhibit B.'''),

        ('Late Payment', 'If payment is not received within thirty (30) days of the due date, Sonar may: (a) charge interest at 1.5% per month on overdue amounts; (b) suspend support services until payment is received; (c) require prepayment for future periods. Suspension of services for non-payment shall not relieve the Customer of payment obligations.'),

        ('Price Adjustments', 'Sonar may adjust the Maintenance Fee annually upon sixty (60) days\' prior written notice. Price increases shall not exceed the greater of: (a) 8% of the prior year\'s fee; or (b) the applicable Consumer Price Index increase. The Customer may terminate this Agreement if it does not accept a price increase, provided notice of termination is given within thirty (30) days of receiving the price increase notice.'),

        ('Taxes', 'All fees are exclusive of applicable taxes, duties, and levies. The Customer is responsible for all such taxes except taxes based on Sonar\'s net income. If withholding is required, the Customer shall gross up the payment so that Sonar receives the full amount invoiced.'),

        ('Additional Services', 'Services not included in the Support Tier, such as on-site visits, custom development, or extended training, shall be provided at Sonar\'s then-current professional services rates and billed separately.'),
    ]

    for i, (title, content) in enumerate(fee_sections):
        para = doc.add_paragraph()
        para.add_run(f'9.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 10: TERM AND RENEWAL
    # ============================================
    doc.add_heading('ARTICLE 10: TERM AND RENEWAL', level=1)

    term_sections = [
        ('Initial Term', 'This Agreement shall commence on the Effective Date and continue for an initial term of one (1) year (the "Initial Term"), unless earlier terminated in accordance with Article 11.'),

        ('Renewal', 'Following the Initial Term, this Agreement shall automatically renew for successive one (1) year periods (each a "Renewal Term") unless either Party provides written notice of non-renewal at least sixty (60) days prior to the end of the then-current term.'),

        ('Continuous Coverage', 'Support coverage is continuous provided that: (a) the Customer maintains an active License Agreement; (b) the Maintenance Fee is paid when due; and (c) the Customer has not terminated or allowed this Agreement to lapse.'),

        ('Reinstatement', 'If this Agreement lapses due to non-payment or non-renewal, the Customer may request reinstatement. Reinstatement shall be at Sonar\'s discretion and may require: (a) payment of back fees for the lapsed period; (b) payment of a reinstatement fee (typically 50% of the annual Maintenance Fee); and (c) upgrade to a currently supported version of the software.'),
    ]

    for i, (title, content) in enumerate(term_sections):
        para = doc.add_paragraph()
        para.add_run(f'10.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 11: TERMINATION
    # ============================================
    doc.add_heading('ARTICLE 11: TERMINATION', level=1)

    term_sections = [
        ('Termination for Convenience', 'Either Party may terminate this Agreement for convenience upon sixty (60) days\' prior written notice. No refund of prepaid fees shall be provided for termination for convenience by the Customer.'),

        ('Termination for Breach', '''Either Party may terminate this Agreement immediately upon written notice if:

(a) The other Party materially breaches this Agreement and fails to cure such breach within thirty (30) days of receiving written notice;
(b) The other Party becomes insolvent, files for bankruptcy, or is subject to similar proceedings;
(c) The other Party ceases to conduct business in the ordinary course.'''),

        ('Termination of License', 'This Agreement shall automatically terminate upon termination of the License Agreement for the Sonar Workflow System.'),

        ('Effect of Termination', '''Upon termination of this Agreement:

(a) Sonar's obligation to provide support and maintenance services shall cease;
(b) The Customer's access to the Support Portal and support resources shall be terminated;
(c) The Customer may continue to use the licensed version of the software (for perpetual licenses) but shall not receive further Updates or support;
(d) Outstanding fees through the termination date shall become immediately due;
(e) Provisions of this Agreement that by their nature should survive shall continue in effect.'''),

        ('Refund on Termination for Breach', 'If Sonar terminates this Agreement due to the Customer\'s breach, no refund shall be provided. If the Customer terminates due to Sonar\'s material breach, Sonar shall refund a pro-rata portion of prepaid fees for the period following termination.'),
    ]

    for i, (title, content) in enumerate(term_sections):
        para = doc.add_paragraph()
        para.add_run(f'11.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 12: EXCLUSIONS AND LIMITATIONS
    # ============================================
    doc.add_heading('ARTICLE 12: EXCLUSIONS AND LIMITATIONS', level=1)

    excl_sections = [
        ('Excluded Services', '''This Agreement does not include:

(a) Support for software versions that have reached end-of-life;
(b) Support for issues caused by modifications not made or authorized by Sonar;
(c) Support for issues caused by use contrary to the Documentation;
(d) Support for third-party software, hardware, or network issues;
(e) Data recovery, migration, or conversion services;
(f) Custom development, customization, or enhancement services;
(g) Training services beyond those included in the Support Tier;
(h) Performance optimization or tuning beyond standard troubleshooting;
(i) Support for trial, evaluation, or beta software;
(j) Support for environments that do not meet minimum system requirements.'''),

        ('Unsupported Configurations', 'Sonar is not obligated to provide support for configurations or environments not certified or documented by Sonar. The Customer assumes responsibility for issues arising from unsupported configurations.'),

        ('Third-Party Integrations', 'Support for integrations with third-party systems is limited to the Sonar Workflow System components of the integration. Issues in third-party systems must be addressed by the respective vendors. Sonar may provide reasonable assistance in diagnosing integration issues.'),

        ('Customer-Caused Issues', 'Sonar may charge for support services provided to resolve issues caused by the Customer\'s actions, such as improper configuration, unauthorized modifications, or failure to follow documented procedures. Sonar shall notify the Customer before incurring such charges.'),

        ('Reasonable Efforts', 'Sonar shall use commercially reasonable efforts to resolve Errors but does not guarantee that all Errors will be resolved. Some issues may require workarounds rather than permanent fixes.'),
    ]

    for i, (title, content) in enumerate(excl_sections):
        para = doc.add_paragraph()
        para.add_run(f'12.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 13: CONFIDENTIALITY
    # ============================================
    doc.add_heading('ARTICLE 13: CONFIDENTIALITY', level=1)

    conf_sections = [
        ('Confidential Information', 'Each Party acknowledges that in the course of this Agreement, it may receive confidential information of the other Party, including technical information, business information, and the terms of this Agreement.'),

        ('Obligations', 'Each Party agrees to: (a) hold the other Party\'s confidential information in strict confidence; (b) not disclose such information to any third party without prior written consent; (c) use such information only for purposes of this Agreement; (d) protect such information with at least the same degree of care used for its own confidential information.'),

        ('Exceptions', 'Confidential information does not include information that: (a) is or becomes publicly available through no fault of the receiving Party; (b) was rightfully in the receiving Party\'s possession prior to disclosure; (c) is independently developed without use of the disclosing Party\'s confidential information; (d) is rightfully obtained from a third party; (e) is required to be disclosed by law.'),

        ('Survival', 'The confidentiality obligations shall survive termination of this Agreement for a period of five (5) years.'),
    ]

    for i, (title, content) in enumerate(conf_sections):
        para = doc.add_paragraph()
        para.add_run(f'13.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 14: LIMITATION OF LIABILITY
    # ============================================
    doc.add_heading('ARTICLE 14: LIMITATION OF LIABILITY', level=1)

    liab_sections = [
        ('Exclusion of Damages', 'IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR EXEMPLARY DAMAGES, INCLUDING DAMAGES FOR LOSS OF PROFITS, REVENUE, DATA, OR USE, ARISING OUT OF OR RELATING TO THIS AGREEMENT, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.'),

        ('Limitation of Liability', 'EXCEPT FOR LIABILITY ARISING FROM BREACH OF CONFIDENTIALITY OBLIGATIONS, GROSS NEGLIGENCE, OR WILLFUL MISCONDUCT, THE TOTAL CUMULATIVE LIABILITY OF EITHER PARTY FOR ALL CLAIMS ARISING OUT OF THIS AGREEMENT SHALL NOT EXCEED THE TOTAL MAINTENANCE FEES PAID OR PAYABLE BY THE CUSTOMER DURING THE TWELVE (12) MONTHS PRECEDING THE CLAIM.'),

        ('Essential Basis', 'THE PARTIES ACKNOWLEDGE THAT THESE LIMITATIONS ARE AN ESSENTIAL ELEMENT OF THE BARGAIN AND REFLECT A REASONABLE ALLOCATION OF RISK. THE LIMITATIONS SHALL APPLY REGARDLESS OF THE FORM OF ACTION AND EVEN IF ANY LIMITED REMEDY FAILS OF ITS ESSENTIAL PURPOSE.'),

        ('Service Credits', 'Service Credits, as provided in Article 8, are the Customer\'s sole and exclusive remedy for service level failures.'),
    ]

    for i, (title, content) in enumerate(liab_sections):
        para = doc.add_paragraph()
        para.add_run(f'14.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # ============================================
    # ARTICLE 15: GENERAL PROVISIONS
    # ============================================
    doc.add_heading('ARTICLE 15: GENERAL PROVISIONS', level=1)

    general_sections = [
        ('Entire Agreement', 'This Agreement, together with the License Agreement and all Exhibits attached hereto, constitutes the entire agreement between the Parties with respect to support and maintenance services and supersedes all prior agreements and understandings.'),

        ('Amendment', 'This Agreement may only be amended by a written instrument signed by authorized representatives of both Parties.'),

        ('Waiver', 'The failure of either Party to enforce any provision shall not constitute a waiver of that provision or the right to enforce it at a later time.'),

        ('Severability', 'If any provision is held to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.'),

        ('Assignment', 'The Customer may not assign this Agreement without Sonar\'s prior written consent. Sonar may assign this Agreement to any successor or affiliate upon notice.'),

        ('Notices', 'All notices shall be in writing and delivered by email with confirmation, registered mail, or courier to the addresses specified in the signature block or as updated in writing.'),

        ('Governing Law', 'This Agreement shall be governed by the laws of [Jurisdiction], without regard to conflict of laws principles.'),

        ('Dispute Resolution', 'Any disputes shall be resolved through good faith negotiation, followed by mediation, and if necessary, binding arbitration in accordance with the rules of [Arbitration Body].'),

        ('Force Majeure', 'Neither Party shall be liable for delays or failures due to causes beyond its reasonable control, including natural disasters, war, terrorism, government actions, or infrastructure failures.'),

        ('Independent Contractors', 'The Parties are independent contractors. Nothing in this Agreement creates a partnership, joint venture, or employment relationship.'),

        ('Counterparts', 'This Agreement may be executed in counterparts, each of which shall be deemed an original.'),

        ('Language', 'This Agreement is executed in English. In case of any translation, the English version shall prevail.'),
    ]

    for i, (title, content) in enumerate(general_sections):
        para = doc.add_paragraph()
        para.add_run(f'15.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_page_break()

    # ============================================
    # EXHIBIT A: SERVICE LEVEL AGREEMENT
    # ============================================
    doc.add_heading('EXHIBIT A: SERVICE LEVEL AGREEMENT (SLA)', level=1)

    doc.add_heading('A.1 Support Hours by Tier', level=2)

    hours_table = doc.add_table(rows=4, cols=4)
    hours_table.style = 'Table Grid'

    hours_headers = ['Support Tier', 'Standard Hours', 'Extended Hours', '24/7 Support']
    for i, header in enumerate(hours_headers):
        cell = hours_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    hours_data = [
        ['Standard', 'Mon-Fri 8:00-18:00', 'Not included', 'Not included'],
        ['Professional', 'Mon-Fri 7:00-19:00', 'Sat 9:00-15:00', 'P1 issues only'],
        ['Enterprise', '24/7', '24/7', 'All priorities'],
    ]

    for row_idx, row_data in enumerate(hours_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = hours_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    doc.add_heading('A.2 Response Time Commitments', level=2)

    resp_table = doc.add_table(rows=5, cols=4)
    resp_table.style = 'Table Grid'

    resp_headers = ['Priority', 'Standard', 'Professional', 'Enterprise']
    for i, header in enumerate(resp_headers):
        cell = resp_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')

    resp_data = [
        ['Priority 1', '4 hours', '2 hours', '30 minutes'],
        ['Priority 2', '8 hours', '4 hours', '2 hours'],
        ['Priority 3', '2 Business Days', '1 Business Day', '4 hours'],
        ['Priority 4', '3 Business Days', '2 Business Days', '1 Business Day'],
    ]

    for row_idx, row_data in enumerate(resp_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = resp_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    doc.add_heading('A.3 Resolution Time Targets', level=2)

    resol_table = doc.add_table(rows=5, cols=4)
    resol_table.style = 'Table Grid'

    resol_headers = ['Priority', 'Standard', 'Professional', 'Enterprise']
    for i, header in enumerate(resol_headers):
        cell = resol_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FCE4D6')

    resol_data = [
        ['Priority 1', '24 hours', '8 hours', '4 hours'],
        ['Priority 2', '72 hours', '24 hours', '8 hours'],
        ['Priority 3', '10 Business Days', '5 Business Days', '3 Business Days'],
        ['Priority 4', '30 Business Days', '15 Business Days', '10 Business Days'],
    ]

    for row_idx, row_data in enumerate(resol_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = resol_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    doc.add_heading('A.4 System Availability (Hosted Deployments)', level=2)

    avail_table = doc.add_table(rows=4, cols=2)
    avail_table.style = 'Table Grid'

    avail_data = [
        ['Support Tier', 'Availability Target'],
        ['Standard', '99.0%'],
        ['Professional', '99.5%'],
        ['Enterprise', '99.9%'],
    ]

    for row_idx, row_data in enumerate(avail_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = avail_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'FFF2CC')

    doc.add_paragraph()

    # ============================================
    # EXHIBIT B: SUPPORT TIERS AND PRICING
    # ============================================
    doc.add_page_break()
    doc.add_heading('EXHIBIT B: SUPPORT TIERS AND PRICING', level=1)

    doc.add_heading('B.1 Support Tier Comparison', level=2)

    tier_table = doc.add_table(rows=16, cols=4)
    tier_table.style = 'Table Grid'

    tier_headers = ['Feature', 'Standard', 'Professional', 'Enterprise']
    for i, header in enumerate(tier_headers):
        cell = tier_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    tier_data = [
        ['Annual Fee (% of License)', '18%', '20%', '22%'],
        ['Support Hours', 'Business Hours', 'Extended Hours', '24/7'],
        ['Phone Support', 'P1-P2 only', 'All priorities', 'All priorities'],
        ['Email Support', 'Yes', 'Yes', 'Yes'],
        ['Support Portal', 'Yes', 'Yes', 'Yes'],
        ['Live Chat', 'No', 'Yes', 'Yes'],
        ['Authorized Contacts', '2', '4', '8'],
        ['Dedicated Account Manager', 'No', 'No', 'Yes'],
        ['Quarterly Business Reviews', 'No', 'No', 'Yes'],
        ['On-site Support', 'Additional fee', '2 days/year incl.', '5 days/year incl.'],
        ['Major Releases', 'Discounted', 'Discounted', 'Included'],
        ['Service Credits', 'No', 'Yes', 'Yes'],
        ['Priority Routing', 'No', 'Yes', 'Yes'],
        ['Root Cause Analysis', 'No', 'P1 only', 'P1-P2'],
        ['Custom Reporting', 'No', 'Quarterly', 'Monthly'],
    ]

    for row_idx, row_data in enumerate(tier_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = tier_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    doc.add_heading('B.2 Pricing', level=2)

    pricing_para = doc.add_paragraph()
    pricing_para.add_run('Annual Maintenance Fee for this Agreement:').bold = True

    doc.add_paragraph()

    price_table = doc.add_table(rows=5, cols=2)
    price_table.style = 'Table Grid'

    price_data = [
        ['Selected Support Tier', '[Standard / Professional / Enterprise]'],
        ['Base License Value', 'USD _______________'],
        ['Maintenance Percentage', '_____ %'],
        ['Annual Maintenance Fee', 'USD _______________'],
        ['Payment Due Date', 'Annually on _______________'],
    ]

    for row_idx, row_data in enumerate(price_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = price_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ============================================
    # EXHIBIT C: ESCALATION MATRIX
    # ============================================
    doc.add_page_break()
    doc.add_heading('EXHIBIT C: ESCALATION MATRIX', level=1)

    doc.add_heading('C.1 Escalation Contacts', level=2)

    esc_table = doc.add_table(rows=5, cols=4)
    esc_table.style = 'Table Grid'

    esc_headers = ['Level', 'Role', 'Contact', 'Response Target']
    for i, header in enumerate(esc_headers):
        cell = esc_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E2EFDA')

    esc_data = [
        ['Level 1', 'Support Team Lead', 'escalation1@sonarmicrosystems.com', '2 hours'],
        ['Level 2', 'Support Manager', 'escalation2@sonarmicrosystems.com', '4 hours'],
        ['Level 3', 'Director, Customer Success', 'escalation3@sonarmicrosystems.com', '2 hours (P1)'],
        ['Level 4', 'VP, Operations', 'executive@sonarmicrosystems.com', 'Immediate'],
    ]

    for row_idx, row_data in enumerate(esc_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = esc_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    doc.add_heading('C.2 Automatic Escalation Rules', level=2)

    auto_table = doc.add_table(rows=5, cols=3)
    auto_table.style = 'Table Grid'

    auto_headers = ['Condition', 'Escalation Level', 'Action']
    for i, header in enumerate(auto_headers):
        cell = auto_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FCE4D6')

    auto_data = [
        ['P1 not responded in 1 hour', 'Level 1', 'Immediate notification'],
        ['P1 not resolved in 4 hours', 'Level 2', 'Management review'],
        ['P1 not resolved in 8 hours', 'Level 3', 'Executive notification'],
        ['Any SLA breach', 'Level 1', 'Logged for review'],
    ]

    for row_idx, row_data in enumerate(auto_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = auto_table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data

    doc.add_paragraph()

    # ============================================
    # EXHIBIT D: AUTHORIZED CONTACTS
    # ============================================
    doc.add_heading('EXHIBIT D: AUTHORIZED CONTACTS', level=1)

    doc.add_paragraph('The following individuals are authorized to submit Support Requests and communicate with Sonar regarding support matters:')

    doc.add_paragraph()

    contact_table = doc.add_table(rows=10, cols=5)
    contact_table.style = 'Table Grid'

    contact_headers = ['Name', 'Title', 'Email', 'Phone', 'Primary/Backup']
    for i, header in enumerate(contact_headers):
        cell = contact_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # Leave rows empty for filling
    for row in range(1, 10):
        for col in range(5):
            contact_table.rows[row].cells[col].text = ''

    doc.add_paragraph()

    doc.add_paragraph('Note: Changes to Authorized Contacts must be submitted in writing to Sonar. Updates will be effective within two (2) Business Days of receipt.')

    doc.add_paragraph()

    # ============================================
    # EXHIBIT E: SYSTEM CONFIGURATION
    # ============================================
    doc.add_heading('EXHIBIT E: SYSTEM CONFIGURATION', level=1)

    doc.add_paragraph('The following system configuration is covered under this Agreement:')

    doc.add_paragraph()

    config_table = doc.add_table(rows=12, cols=2)
    config_table.style = 'Table Grid'

    config_data = [
        ['Item', 'Details'],
        ['Software Product', 'Sonar Workflow System'],
        ['Licensed Version', ''],
        ['Licensed Edition', '[Standard / Professional / Enterprise]'],
        ['Licensed Users', ''],
        ['Deployment Type', '[On-Premise / Cloud / Hybrid]'],
        ['Production Environment', ''],
        ['Test/Staging Environment', ''],
        ['Database Platform', ''],
        ['Operating System', ''],
        ['Primary Site Location', ''],
        ['Disaster Recovery Site', ''],
    ]

    for row_idx, row_data in enumerate(config_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = config_table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            if row_idx == 0 or col_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
            if row_idx == 0:
                set_cell_shading(cell, 'E2EFDA')

    doc.add_paragraph()

    # ============================================
    # SIGNATURE PAGE
    # ============================================
    doc.add_page_break()
    doc.add_heading('SIGNATURE PAGE', level=1)

    sig_intro = doc.add_paragraph()
    sig_intro.add_run('IN WITNESS WHEREOF, the Parties have executed this Support and Maintenance Agreement as of the Effective Date.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature table
    sig_table = doc.add_table(rows=1, cols=2)

    # Provider signature
    cell1 = sig_table.rows[0].cells[0]
    p1 = cell1.add_paragraph()
    p1.add_run('SERVICE PROVIDER:').bold = True
    cell1.add_paragraph()
    p1a = cell1.add_paragraph()
    p1a.add_run('SONAR MICROSYSTEMS PVT. LTD.').bold = True
    cell1.add_paragraph()
    cell1.add_paragraph()
    cell1.add_paragraph('_________________________________')
    cell1.add_paragraph('Authorized Signatory')
    cell1.add_paragraph()
    cell1.add_paragraph('Name: _________________________')
    cell1.add_paragraph('Title: __________________________')
    cell1.add_paragraph('Date: __________________________')
    cell1.add_paragraph()
    cell1.add_paragraph('Email: _________________________')
    cell1.add_paragraph('Phone: _________________________')

    # Customer signature
    cell2 = sig_table.rows[0].cells[1]
    p2 = cell2.add_paragraph()
    p2.add_run('CUSTOMER:').bold = True
    cell2.add_paragraph()
    p2a = cell2.add_paragraph()
    p2a.add_run('[CUSTOMER NAME]').bold = True
    cell2.add_paragraph()
    cell2.add_paragraph()
    cell2.add_paragraph('_________________________________')
    cell2.add_paragraph('Authorized Signatory')
    cell2.add_paragraph()
    cell2.add_paragraph('Name: _________________________')
    cell2.add_paragraph('Title: __________________________')
    cell2.add_paragraph('Date: __________________________')
    cell2.add_paragraph()
    cell2.add_paragraph('Email: _________________________')
    cell2.add_paragraph('Phone: _________________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer1 = doc.add_paragraph()
    footer1.add_run('SONAR WORKFLOW SYSTEM - SUPPORT AND MAINTENANCE AGREEMENT').bold = True
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer2 = doc.add_paragraph()
    footer2.add_run('CONFIDENTIAL')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer3 = doc.add_paragraph()
    footer3.add_run(f'© {datetime.now().year} Sonar Microsystems Pvt. Ltd. All Rights Reserved.')
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save('Sonar Support and Maintenance Agreement.docx')
    print('Support and Maintenance Agreement created successfully: Sonar Support and Maintenance Agreement.docx')

if __name__ == '__main__':
    create_support_agreement()
