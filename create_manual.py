import re
import sys

# Read the function definitions file
with open('frontend/src/app/core/data/function-definitions.ts', 'r', encoding='utf-8') as f:
    content = f.read()

# Extract all function definitions
functions = {}
current_func = None
current_data = {}
in_params = False
in_examples = False
in_tips = False
in_related = False

lines = content.split('\n')
for line in lines:
    # Check for function start
    match = re.match(r"\s+'([A-Z_]+)':\s*\{", line)
    if match:
        if current_func:
            functions[current_func] = current_data
        current_func = match.group(1)
        current_data = {
            'name': current_func,
            'syntax': '',
            'category': '',
            'description': '',
            'explanation': '',
            'parameters': [],
            'examples': [],
            'tips': [],
            'relatedFunctions': []
        }
        in_params = False
        in_examples = False
        in_tips = False
        in_related = False
        continue

    if current_func:
        # Parse name
        match = re.search(r"name:\s*'([^']+)'", line)
        if match:
            current_data['name'] = match.group(1)

        # Parse syntax
        match = re.search(r"syntax:\s*'([^']+)'", line)
        if match:
            current_data['syntax'] = match.group(1)

        # Parse category
        match = re.search(r"category:\s*'([^']+)'", line)
        if match:
            current_data['category'] = match.group(1)

        # Parse description
        match = re.search(r"description:\s*'([^']+)'", line)
        if match:
            current_data['description'] = match.group(1)

        # Parse explanation
        match = re.search(r"explanation:\s*'([^']+)'", line)
        if match:
            current_data['explanation'] = match.group(1)

        # Check for sections
        if 'parameters:' in line:
            in_params = True
            in_examples = False
            in_tips = False
            in_related = False
            continue
        elif 'examples:' in line:
            in_params = False
            in_examples = True
            in_tips = False
            in_related = False
            continue
        elif 'tips:' in line:
            in_params = False
            in_examples = False
            in_tips = True
            in_related = False
            continue
        elif 'relatedFunctions:' in line:
            in_params = False
            in_examples = False
            in_tips = False
            in_related = True
            continue

        # Parse parameters
        if in_params:
            match = re.search(r"\{ name:\s*'([^']+)',\s*type:\s*'([^']+)',\s*description:\s*'([^']+)'(,\s*required:\s*(true|false))? \}", line)
            if match:
                current_data['parameters'].append({
                    'name': match.group(1),
                    'type': match.group(2),
                    'description': match.group(3),
                    'required': match.group(5) == 'true' if match.group(5) else False
                })

        # Parse examples
        if in_examples:
            match = re.search(r"\{ usage:\s*'([^']+)',\s*result:\s*'([^']+)'(,\s*description:\s*'([^']+)')? \}", line)
            if match:
                current_data['examples'].append({
                    'usage': match.group(1),
                    'result': match.group(2),
                    'description': match.group(4) if match.group(4) else ''
                })

        # Parse tips
        if in_tips:
            match = re.search(r"'([^']+)'", line)
            if match:
                current_data['tips'].append(match.group(1))

        # Parse related functions
        if in_related:
            match = re.search(r"\[([^\]]+)\]", line)
            if match:
                funcs = [f.strip().replace("'", "") for f in match.group(1).split(',')]
                current_data['relatedFunctions'] = funcs

# Add last function
if current_func:
    functions[current_func] = current_data

# Group by category
categories = {}
for name, func in functions.items():
    cat = func.get('category', 'Other')
    if cat not in categories:
        categories[cat] = []
    categories[cat].append(func)

# Create Word document
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Title
title = doc.add_heading('Sonar Workflow Functions Manual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Reference Guide for Workflow Expression Functions')
run.bold = True

doc.add_paragraph()

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'This manual provides a comprehensive reference for all expression functions available in the Sonar Workflow System. '
    'These functions can be used in computed fields, default values, validation rules, and workflow expressions to manipulate '
    'data, perform calculations, and implement complex business logic.'
)

doc.add_heading('How to Use This Manual', level=1)
doc.add_paragraph(
    'Functions are organized by category for easy reference. Each function entry includes:\n'
    '\n• Syntax - The proper function call format\n'
    '• Description - Brief overview of what the function does\n'
    '• Explanation - Detailed explanation of the function behavior\n'
    '• Parameters - Input parameters the function accepts\n'
    '• Examples - Practical usage examples with expected results\n'
    '• Tips - Helpful tips and best practices\n'
    '• Related Functions - Other functions that may be useful'
)

# Table of Contents
doc.add_heading('Table of Contents', level=1)
for cat in sorted(categories.keys()):
    doc.add_paragraph(f"{cat} Functions", style='List Bullet')

doc.add_page_break()

# Write functions by category
for category in sorted(categories.keys()):
    doc.add_heading(f'{category} Functions', level=1)

    funcs = sorted(categories[category], key=lambda x: x['name'])

    for func in funcs:
        # Function name as heading
        doc.add_heading(func['name'], level=2)

        # Syntax
        p = doc.add_paragraph()
        p.add_run('Syntax: ').bold = True
        code = p.add_run(func.get('syntax', func['name'] + '()'))
        code.font.name = 'Courier New'

        # Description
        if func.get('description'):
            p = doc.add_paragraph()
            p.add_run('Description: ').bold = True
            p.add_run(func['description'])

        # Explanation
        if func.get('explanation'):
            p = doc.add_paragraph()
            p.add_run('Explanation: ').bold = True
            p.add_run(func['explanation'])

        # Parameters
        if func.get('parameters'):
            p = doc.add_paragraph()
            p.add_run('Parameters:').bold = True
            for param in func['parameters']:
                req = ' (required)' if param['required'] else ' (optional)'
                p = doc.add_paragraph(f"  • {param['name']} ({param['type']}){req}: {param['description']}", style='List Bullet')

        # Examples
        if func.get('examples'):
            p = doc.add_paragraph()
            p.add_run('Examples:').bold = True
            for ex in func['examples']:
                p = doc.add_paragraph()
                p.add_run(f"  Usage: ").bold = True
                code = p.add_run(ex['usage'])
                code.font.name = 'Courier New'
                p.add_run(f"\n  Result: ")
                code = p.add_run(str(ex['result']))
                code.font.name = 'Courier New'
                if ex.get('description'):
                    p.add_run(f"\n  ({ex['description']})")

        # Tips
        if func.get('tips'):
            p = doc.add_paragraph()
            p.add_run('Tips:').bold = True
            for tip in func['tips']:
                doc.add_paragraph(f"  • {tip}", style='List Bullet')

        # Related Functions
        if func.get('relatedFunctions'):
            p = doc.add_paragraph()
            p.add_run('Related Functions: ').bold = True
            p.add_run(', '.join(func['relatedFunctions']))

        doc.add_paragraph()  # Empty line between functions

    doc.add_page_break()

# Quick Reference Appendix
doc.add_heading('Quick Reference', level=1)

# Create quick reference table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Function'
hdr_cells[1].text = 'Category'
hdr_cells[2].text = 'Syntax'

for cat in sorted(categories.keys()):
    for func in sorted(categories[cat], key=lambda x: x['name']):
        row_cells = table.add_row().cells
        row_cells[0].text = func['name']
        row_cells[1].text = cat
        row_cells[2].text = func.get('syntax', func['name'] + '()')

# Save document
doc.save('Sonar Workflow Functions Manual.docx')
print(f"Document created successfully!")
print(f"Total functions: {len(functions)}")
print(f"Categories: {len(categories)}")
