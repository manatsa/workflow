from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

def add_hdr_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1976D2')
        shd.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shd)

def add_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        run = row.cells[i].paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

# ====== COVER ======
for _ in range(3):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Sona Workflow System')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
r.bold = True

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Configuration File Manual')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('config.sona')
r.font.size = Pt(16)
r.font.name = 'Consolas'
r.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)

for _ in range(3):
    doc.add_paragraph()

v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('Version 1.5.0\n' + datetime.date.today().strftime('%B %d, %Y'))
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Acad Arch Solutions Pvt. Ltd.')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ====== TOC ======
doc.add_heading('Table of Contents', level=1)
toc = [
    ('1.', 'Introduction'), ('2.', 'File Location'), ('3.', 'File Format'),
    ('4.', 'Configuration Sections'), ('  4.1', 'Database Configuration'),
    ('  4.2', 'Server Configuration'), ('  4.3', 'Mail Configuration'),
    ('  4.4', 'Application Configuration'), ('5.', 'Complete Example'),
    ('6.', 'Common Scenarios'), ('  6.1', 'Remote Database'), ('  6.2', 'Change Port'),
    ('  6.3', 'Email Setup'), ('  6.4', 'Production Deployment'), ('  6.5', 'Minimal Config'),
    ('7.', 'Startup Verification'), ('8.', 'Troubleshooting'), ('9.', 'Security Considerations'),
]
for num, item in toc:
    p = doc.add_paragraph()
    r = p.add_run(num + '  ' + item)
    r.font.size = Pt(11)
    if not num.startswith(' '):
        r.bold = True

doc.add_page_break()

# ====== 1. INTRO ======
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('The config.sona file is the primary configuration file for the Sona Workflow System. It allows administrators to configure database connections, server settings, email integration, and application parameters without modifying internal configuration files.')
doc.add_paragraph('The file uses JSON format and is designed to be human-readable. When present, settings in config.sona override the application defaults.')

p = doc.add_paragraph()
p.add_run('Key Benefits:').bold = True
for b in ['Configure database connections for different environments', 'Change server port without recompiling',
           'Set up email notifications for workflow approvals', 'Customize file storage paths',
           'Secure sensitive credentials outside the JAR', 'Only include sections you need to override']:
    doc.add_paragraph(b, style='List Bullet')

# ====== 2. FILE LOCATION ======
doc.add_heading('2. File Location', level=1)
doc.add_paragraph('On startup, the system searches for config.sona in these locations (first found wins):')

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
add_hdr_row(t, ['Priority', 'Location', 'Description'])
add_row(t, ['1', 'Current working directory', 'Directory the application is launched from'])
add_row(t, ['2', 'Parent directory', 'One level up from working directory'])
add_row(t, ['3', 'JAR/EXE directory', 'Same folder as the application executable'])
add_row(t, ['4', 'User home directory', 'C:\\Users\\<username>\\ on Windows'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('For MSI installations: ').bold = True
p.add_run('Place in the installation directory, typically:')
code_block(doc, 'C:\\Program Files\\Sona\\Workflow System\\config.sona')

p = doc.add_paragraph()
p.add_run('For JAR deployments: ').bold = True
p.add_run('Place next to the JAR file:')
code_block(doc, 'C:\\Sona\\workflow-system-1.5.0.jar\nC:\\Sona\\config.sona')

# ====== 3. FILE FORMAT ======
doc.add_heading('3. File Format', level=1)
doc.add_paragraph('The file uses standard JSON format. It must be valid JSON.')

p = doc.add_paragraph()
p.add_run('Rules:').bold = True
for r in ['All string values in double quotes ("value")', 'Numbers without quotes (5432)',
           'No comments (// or /* */ are not allowed)', 'UTF-8 encoding', 'No trailing commas']:
    doc.add_paragraph(r, style='List Bullet')

# ====== 4. SECTIONS ======
doc.add_heading('4. Configuration Sections', level=1)
doc.add_paragraph('All sections are optional. Only include what you need to override.')

# 4.1 Database
doc.add_heading('4.1 Database Configuration', level=2)
doc.add_paragraph('Configures the PostgreSQL database connection.')
code_block(doc, '{\n  "database": {\n    "host": "localhost",\n    "port": 5432,\n    "name": "workflow",\n    "username": "sonar",\n    "password": "P@88345!"\n  }\n}')

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
add_hdr_row(t, ['Property', 'Type', 'Default', 'Description'])
add_row(t, ['host', 'String', 'localhost', 'Database server hostname or IP'])
add_row(t, ['port', 'Number', '5432', 'PostgreSQL port'])
add_row(t, ['name', 'String', 'workflow', 'Database name'])
add_row(t, ['username', 'String', 'sonar', 'Database user'])
add_row(t, ['password', 'String', 'P@88345!', 'Database password'])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Note: ')
r.bold = True
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
p.add_run('The database must exist before starting. Tables are created automatically.')

# 4.2 Server
doc.add_heading('4.2 Server Configuration', level=2)
code_block(doc, '{\n  "server": {\n    "port": 9500\n  }\n}')

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
add_hdr_row(t, ['Property', 'Type', 'Default', 'Description'])
add_row(t, ['port', 'Number', '9500', 'HTTP port the application listens on'])

# 4.3 Mail
doc.add_heading('4.3 Mail Configuration', level=2)
doc.add_paragraph('Configures SMTP email for notifications, approvals, and password resets.')
code_block(doc, '{\n  "mail": {\n    "host": "smtp.gmail.com",\n    "port": 587,\n    "username": "notifications@company.com",\n    "password": "app-specific-password"\n  }\n}')

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
add_hdr_row(t, ['Property', 'Type', 'Default', 'Description'])
add_row(t, ['host', 'String', 'smtp.gmail.com', 'SMTP server hostname'])
add_row(t, ['port', 'Number', '587', 'SMTP port (587=TLS, 465=SSL)'])
add_row(t, ['username', 'String', '(empty)', 'SMTP username/email'])
add_row(t, ['password', 'String', '(empty)', 'SMTP password or app password'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Common SMTP providers:').bold = True

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
add_hdr_row(t, ['Provider', 'Host', 'Port'])
add_row(t, ['Gmail', 'smtp.gmail.com', '587'])
add_row(t, ['Outlook/Office 365', 'smtp.office365.com', '587'])
add_row(t, ['Yahoo', 'smtp.mail.yahoo.com', '587'])
add_row(t, ['Amazon SES', 'email-smtp.<region>.amazonaws.com', '587'])

# 4.4 App
doc.add_heading('4.4 Application Configuration', level=2)
code_block(doc, '{\n  "app": {\n    "baseUrl": "http://your-server:9500",\n    "jwtSecret": "YourSecretKeyAtLeast256Bits",\n    "encryptionKey": "YourAESKey16Char",\n    "storagePath": "C:/Sonar Docs/"\n  }\n}')

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
add_hdr_row(t, ['Property', 'Type', 'Default', 'Description'])
add_row(t, ['baseUrl', 'String', '(auto)', 'Public URL for email links'])
add_row(t, ['jwtSecret', 'String', '(built-in)', 'JWT token signing key (min 256 bits)'])
add_row(t, ['encryptionKey', 'String', '(built-in)', 'AES encryption key (16 characters)'])
add_row(t, ['storagePath', 'String', 'C:/Sonar Docs/', 'Base path for file storage'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Auto-created subfolders: ').bold = True
p.add_run('attachments/, templates/, imports/, exports/, backups/')

# ====== 5. COMPLETE EXAMPLE ======
doc.add_page_break()
doc.add_heading('5. Complete Example', level=1)
doc.add_paragraph('A complete config.sona file with all sections:')
code_block(doc, '{\n  "database": {\n    "host": "192.168.1.100",\n    "port": 5432,\n    "name": "sona_production",\n    "username": "sona_app",\n    "password": "Str0ng!Pr0duct10n#Pass"\n  },\n  "server": {\n    "port": 8080\n  },\n  "mail": {\n    "host": "smtp.office365.com",\n    "port": 587,\n    "username": "workflow@company.com",\n    "password": "email-app-password"\n  },\n  "app": {\n    "baseUrl": "https://workflow.company.com",\n    "jwtSecret": "MyCompanySecureJWTKey2026ForProduction!",\n    "encryptionKey": "MyAESKey2026Prod",\n    "storagePath": "D:/SonaData/"\n  }\n}')

# ====== 6. SCENARIOS ======
doc.add_heading('6. Common Scenarios', level=1)

doc.add_heading('6.1 Remote Database', level=2)
code_block(doc, '{\n  "database": {\n    "host": "db-server.company.com",\n    "port": 5432,\n    "name": "sona_workflow",\n    "username": "sona_user",\n    "password": "SecureDbPassword!"\n  }\n}')

doc.add_heading('6.2 Change Port', level=2)
code_block(doc, '{\n  "server": {\n    "port": 8080\n  }\n}')

doc.add_heading('6.3 Email Setup', level=2)
code_block(doc, '{\n  "mail": {\n    "host": "smtp.gmail.com",\n    "port": 587,\n    "username": "sona.notifications@gmail.com",\n    "password": "abcd efgh ijkl mnop"\n  },\n  "app": {\n    "baseUrl": "http://192.168.1.50:9500"\n  }\n}')
doc.add_paragraph('The baseUrl is required for email links (approval buttons, view submission) to work correctly.')

doc.add_heading('6.4 Production Deployment', level=2)
code_block(doc, '{\n  "database": {\n    "host": "prod-db.internal",\n    "name": "sona_prod",\n    "username": "sona_prod_user",\n    "password": "V3ryStr0ng!Pr0d#2026"\n  },\n  "app": {\n    "baseUrl": "https://workflow.yourcompany.com",\n    "jwtSecret": "UniqueProductionJWTSecretAtLeast32Chars!!",\n    "encryptionKey": "ProdAES128Key!!!"\n  }\n}')

doc.add_heading('6.5 Minimal Config', level=2)
doc.add_paragraph('Only override what you need:')
code_block(doc, '{\n  "database": {\n    "host": "10.0.0.5",\n    "name": "my_workflow_db",\n    "username": "admin",\n    "password": "admin123"\n  }\n}')

# ====== 7. VERIFICATION ======
doc.add_heading('7. Startup Verification', level=1)
doc.add_paragraph('Check the console output on startup:')

p = doc.add_paragraph()
r = p.add_run('Successful load:')
r.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
code_block(doc, '[Sona Config] Loading configuration from: C:\\Sona\\config.sona\n[Sona Config] Database: postgresql://192.168.1.100:5432/sona_production (user: sona_app)\n[Sona Config] Server port: 8080\n[Sona Config] Applied 14 configuration properties.')

p = doc.add_paragraph()
r = p.add_run('File not found:')
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
code_block(doc, '[Sona Config] No config.sona file found. Using default application.yml settings.')

p = doc.add_paragraph()
r = p.add_run('Parse error:')
r.bold = True
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
code_block(doc, '[Sona Config] ERROR: Failed to parse config.sona: Unexpected character...\n[Sona Config] Falling back to default application.yml settings.')

# ====== 8. TROUBLESHOOTING ======
doc.add_heading('8. Troubleshooting', level=1)
issues = [
    ("Application won't start after editing config.sona", 'Validate JSON at jsonlint.com. Common: missing commas, trailing commas, single quotes instead of double.'),
    ('Database connection refused', 'Verify host is reachable, port correct, PostgreSQL running, database exists. Check firewall.'),
    ('Emails not sending', 'Verify SMTP credentials. For Gmail use App Password. Check host and port.'),
    ('config.sona is ignored', 'Ensure file is in a search location (Section 2). Name must be exactly "config.sona".'),
    ('Special characters in password', 'Escape backslash as \\\\ and double quote as \\" in JSON strings.'),
    ('Storage path not working', 'Use forward slashes: "C:/Sonar Docs/" not "C:\\\\Sonar Docs\\\\". Ensure write permissions.'),
]
for title, fix in issues:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(fix)

# ====== 9. SECURITY ======
doc.add_heading('9. Security Considerations', level=1)
for s in [
    'Restrict file permissions to app service account and admins only',
    'Never commit config.sona to version control',
    'Use strong, unique passwords for database and JWT in production',
    'encryptionKey must be exactly 16 characters for AES-128',
    'Change default JWT secret in production',
    'Environment variables override config.sona for most sensitive values',
    'Rotate passwords periodically and update config.sona',
    'Back up config.sona separately from the application',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\u00A9 2026 Acad Arch Solutions Pvt. Ltd. All rights reserved.')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('Sona Config File Manual.docx')
print('Created: Sona Config File Manual.docx')
