"""
Script to generate Sonar Workflow Quality Test and UAT Script documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_style(doc):
    """Add custom heading styles"""
    pass

def create_quality_test_document():
    """Create comprehensive Quality Test document"""
    doc = Document()

    # Title
    title = doc.add_heading('Sonar Workflow System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Quality Assurance Test Document', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document Info
    doc.add_paragraph()
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Document Version', '1.0'),
        ('Date', '2026-01-27'),
        ('Author', 'QA Team'),
        ('Status', 'Draft'),
        ('Classification', 'Internal')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        set_cell_shading(info_table.rows[i].cells[0], 'E3F2FD')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Test Objectives',
        '3. Test Scope',
        '4. Test Environment',
        '5. Test Cases - Authentication Module',
        '6. Test Cases - User Management Module',
        '7. Test Cases - Organization Management Module',
        '8. Test Cases - Workflow Builder Module',
        '9. Test Cases - Workflow Submission Module',
        '10. Test Cases - Approval Process Module',
        '11. Test Cases - Reports Module',
        '12. Test Cases - Audit Trail Module',
        '13. Test Cases - Settings Module',
        '14. Test Cases - API Endpoints',
        '15. Performance Testing',
        '16. Security Testing',
        '17. Test Summary'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        'This document outlines the comprehensive quality assurance testing plan for the Sonar Workflow System. '
        'The purpose of this document is to ensure that all system functionalities meet the specified requirements '
        'and perform as expected under various conditions.'
    )
    doc.add_paragraph(
        'The Sonar Workflow System is an enterprise workflow management application that enables organizations '
        'to create, manage, and automate business processes with multi-level approval workflows.'
    )

    # 2. Test Objectives
    doc.add_heading('2. Test Objectives', 1)
    objectives = [
        'Verify all functional requirements are implemented correctly',
        'Ensure system stability and reliability under normal and stress conditions',
        'Validate data integrity across all modules',
        'Confirm security controls are properly implemented',
        'Verify integration between different system components',
        'Ensure compliance with business rules and validation logic',
        'Test user interface responsiveness and usability',
        'Validate email notification functionality',
        'Test role-based access controls',
        'Verify audit trail accuracy and completeness'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # 3. Test Scope
    doc.add_heading('3. Test Scope', 1)
    doc.add_heading('3.1 In Scope', 2)
    in_scope = [
        'User authentication and authorization',
        'User and role management',
        'Organization hierarchy management (SBU, Corporate, Department, Branch, Category)',
        'Workflow template creation and configuration',
        'Dynamic form builder with field validations',
        'Workflow submission and approval processes',
        'Email notifications for approvals',
        'Reporting and analytics',
        'Audit trail logging',
        'System settings and configuration',
        'API endpoint testing',
        'Dark mode and theme support'
    ]
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Out of Scope', 2)
    out_scope = [
        'Third-party integrations not specified in requirements',
        'Mobile application testing',
        'Load testing beyond specified user limits',
        'Disaster recovery testing'
    ]
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Test Environment
    doc.add_heading('4. Test Environment', 1)
    env_table = doc.add_table(rows=7, cols=2)
    env_table.style = 'Table Grid'
    env_data = [
        ('Component', 'Specification'),
        ('Server OS', 'Windows Server 2019 / Linux Ubuntu 22.04'),
        ('Database', 'PostgreSQL 15+'),
        ('Application Server', 'Spring Boot 3.x with embedded Tomcat'),
        ('Frontend', 'Angular 19'),
        ('Browser Support', 'Chrome 120+, Firefox 120+, Edge 120+'),
        ('Java Version', 'Java 21 (Temurin)')
    ]
    for i, (col1, col2) in enumerate(env_data):
        env_table.rows[i].cells[0].text = col1
        env_table.rows[i].cells[1].text = col2
        if i == 0:
            set_cell_shading(env_table.rows[i].cells[0], '1976D2')
            set_cell_shading(env_table.rows[i].cells[1], '1976D2')
            for cell in env_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

    doc.add_page_break()

    # 5. Authentication Module Test Cases
    doc.add_heading('5. Test Cases - Authentication Module', 1)

    auth_tests = [
        ('TC-AUTH-001', 'Valid Login', 'Verify user can login with valid credentials',
         '1. Navigate to login page\n2. Enter valid username\n3. Enter valid password\n4. Click Login button',
         'User successfully logged in and redirected to dashboard', 'High'),
        ('TC-AUTH-002', 'Invalid Password', 'Verify error message for invalid password',
         '1. Navigate to login page\n2. Enter valid username\n3. Enter invalid password\n4. Click Login button',
         'Error message displayed: "Invalid credentials"', 'High'),
        ('TC-AUTH-003', 'Invalid Username', 'Verify error message for non-existent user',
         '1. Navigate to login page\n2. Enter non-existent username\n3. Enter any password\n4. Click Login button',
         'Error message displayed: "Invalid credentials"', 'High'),
        ('TC-AUTH-004', 'Empty Credentials', 'Verify validation for empty fields',
         '1. Navigate to login page\n2. Leave username empty\n3. Leave password empty\n4. Click Login button',
         'Validation messages displayed for required fields', 'Medium'),
        ('TC-AUTH-005', 'Session Timeout', 'Verify session expires after inactivity',
         '1. Login successfully\n2. Wait for session timeout period\n3. Try to access protected page',
         'User redirected to login page with session expired message', 'High'),
        ('TC-AUTH-006', 'Logout Functionality', 'Verify user can logout successfully',
         '1. Login successfully\n2. Click on user profile\n3. Click Logout button',
         'User logged out and redirected to login page', 'High'),
        ('TC-AUTH-007', 'Remember Me', 'Verify remember me functionality',
         '1. Login with Remember Me checked\n2. Close browser\n3. Reopen browser and navigate to app',
         'User session persisted', 'Medium'),
        ('TC-AUTH-008', 'Password Visibility Toggle', 'Verify password show/hide toggle',
         '1. Navigate to login page\n2. Enter password\n3. Click eye icon to show password\n4. Click again to hide',
         'Password visibility toggles correctly', 'Low'),
    ]

    create_test_table(doc, auth_tests)

    doc.add_page_break()

    # 6. User Management Module
    doc.add_heading('6. Test Cases - User Management Module', 1)

    user_tests = [
        ('TC-USER-001', 'Create New User', 'Verify admin can create a new user',
         '1. Login as admin\n2. Navigate to User Management\n3. Click Add User\n4. Fill all required fields\n5. Click Save',
         'New user created successfully with confirmation message', 'High'),
        ('TC-USER-002', 'Create User - Duplicate Username', 'Verify duplicate username prevention',
         '1. Login as admin\n2. Try to create user with existing username',
         'Error message: "Username already exists"', 'High'),
        ('TC-USER-003', 'Edit User Details', 'Verify user details can be edited',
         '1. Login as admin\n2. Navigate to User Management\n3. Click Edit on existing user\n4. Modify details\n5. Save',
         'User details updated successfully', 'High'),
        ('TC-USER-004', 'Deactivate User', 'Verify user can be deactivated',
         '1. Login as admin\n2. Navigate to User Management\n3. Click on active user\n4. Toggle active status to inactive',
         'User deactivated and cannot login', 'High'),
        ('TC-USER-005', 'Assign Role to User', 'Verify role assignment functionality',
         '1. Login as admin\n2. Navigate to User Management\n3. Edit user\n4. Assign different role\n5. Save',
         'Role assigned successfully, permissions updated', 'High'),
        ('TC-USER-006', 'Search Users', 'Verify user search functionality',
         '1. Login as admin\n2. Navigate to User Management\n3. Enter search term in search box',
         'Matching users displayed in results', 'Medium'),
        ('TC-USER-007', 'Filter Users by Role', 'Verify user filtering by role',
         '1. Login as admin\n2. Navigate to User Management\n3. Select role filter',
         'Only users with selected role displayed', 'Medium'),
        ('TC-USER-008', 'User Password Reset', 'Verify admin can reset user password',
         '1. Login as admin\n2. Navigate to User Management\n3. Select user\n4. Click Reset Password',
         'Password reset successfully, user can login with new password', 'High'),
        ('TC-USER-009', 'Bulk User Import', 'Verify bulk user import from Excel',
         '1. Login as admin\n2. Navigate to User Management\n3. Click Import\n4. Upload Excel file',
         'Users imported successfully with summary report', 'Medium'),
        ('TC-USER-010', 'Export Users', 'Verify user export functionality',
         '1. Login as admin\n2. Navigate to User Management\n3. Click Export',
         'Excel file downloaded with user data', 'Low'),
    ]

    create_test_table(doc, user_tests)

    doc.add_page_break()

    # 7. Organization Management
    doc.add_heading('7. Test Cases - Organization Management Module', 1)

    org_tests = [
        ('TC-ORG-001', 'Create SBU', 'Verify SBU creation',
         '1. Login as admin\n2. Navigate to Organization > SBU\n3. Click Add\n4. Enter SBU details\n5. Save',
         'SBU created successfully', 'High'),
        ('TC-ORG-002', 'Create Corporate', 'Verify Corporate creation under SBU',
         '1. Login as admin\n2. Navigate to Organization > Corporate\n3. Click Add\n4. Select parent SBU\n5. Enter details\n6. Save',
         'Corporate created under selected SBU', 'High'),
        ('TC-ORG-003', 'Create Department', 'Verify Department creation',
         '1. Login as admin\n2. Navigate to Organization > Department\n3. Click Add\n4. Select parent Corporate\n5. Enter details\n6. Save',
         'Department created under selected Corporate', 'High'),
        ('TC-ORG-004', 'Create Branch', 'Verify Branch creation',
         '1. Login as admin\n2. Navigate to Organization > Branch\n3. Click Add\n4. Enter branch details\n5. Save',
         'Branch created successfully', 'High'),
        ('TC-ORG-005', 'Create Category', 'Verify Category creation',
         '1. Login as admin\n2. Navigate to Organization > Category\n3. Click Add\n4. Enter category details\n5. Save',
         'Category created successfully', 'High'),
        ('TC-ORG-006', 'Edit Organization Unit', 'Verify organization unit editing',
         '1. Login as admin\n2. Navigate to any organization unit\n3. Click Edit\n4. Modify details\n5. Save',
         'Organization unit updated successfully', 'High'),
        ('TC-ORG-007', 'Delete Organization Unit', 'Verify organization unit deletion',
         '1. Login as admin\n2. Navigate to organization unit without children\n3. Click Delete\n4. Confirm deletion',
         'Organization unit deleted successfully', 'Medium'),
        ('TC-ORG-008', 'Delete Prevention - Has Children', 'Verify cannot delete unit with children',
         '1. Login as admin\n2. Try to delete organization unit with child units',
         'Error message: Cannot delete unit with child organizations', 'High'),
        ('TC-ORG-009', 'Organization Hierarchy View', 'Verify hierarchy tree view',
         '1. Login as admin\n2. Navigate to Organization\n3. View hierarchy tree',
         'Complete organization hierarchy displayed correctly', 'Medium'),
        ('TC-ORG-010', 'Assign User to Organization', 'Verify user-organization assignment',
         '1. Login as admin\n2. Edit user\n3. Assign to specific SBU/Corporate/Department',
         'User assigned to organization successfully', 'High'),
    ]

    create_test_table(doc, org_tests)

    doc.add_page_break()

    # 8. Workflow Builder Module
    doc.add_heading('8. Test Cases - Workflow Builder Module', 1)

    wf_builder_tests = [
        ('TC-WFB-001', 'Create New Workflow', 'Verify workflow creation',
         '1. Login as admin\n2. Navigate to Workflows\n3. Click Create Workflow\n4. Enter workflow name and description\n5. Save',
         'New workflow created successfully', 'High'),
        ('TC-WFB-002', 'Add Text Field', 'Verify adding text field to form',
         '1. Open workflow in builder\n2. Drag Text field from palette\n3. Configure field properties\n4. Save',
         'Text field added to form successfully', 'High'),
        ('TC-WFB-003', 'Add Number Field', 'Verify adding number field to form',
         '1. Open workflow in builder\n2. Drag Number field from palette\n3. Configure field properties\n4. Save',
         'Number field added with numeric validation', 'High'),
        ('TC-WFB-004', 'Add Date Field', 'Verify adding date field to form',
         '1. Open workflow in builder\n2. Drag Date field from palette\n3. Configure field properties\n4. Save',
         'Date field added with date picker', 'High'),
        ('TC-WFB-005', 'Add Dropdown Field', 'Verify adding dropdown field',
         '1. Open workflow in builder\n2. Drag Dropdown field\n3. Add options\n4. Save',
         'Dropdown field added with configured options', 'High'),
        ('TC-WFB-006', 'Add Checkbox Field', 'Verify adding checkbox field',
         '1. Open workflow in builder\n2. Drag Checkbox field\n3. Configure\n4. Save',
         'Checkbox field added successfully', 'High'),
        ('TC-WFB-007', 'Add File Upload Field', 'Verify adding file upload field',
         '1. Open workflow in builder\n2. Drag File Upload field\n3. Configure allowed types and size\n4. Save',
         'File upload field added with restrictions', 'High'),
        ('TC-WFB-008', 'Configure Field Validation - Required', 'Verify required field validation',
         '1. Open workflow builder\n2. Add text field\n3. Mark as required\n4. Save and test form',
         'Field shows required validation error when empty', 'High'),
        ('TC-WFB-009', 'Configure Field Validation - Min/Max Length', 'Verify length validation',
         '1. Add text field\n2. Set min length 5, max length 50\n3. Save and test',
         'Validation enforces length constraints', 'High'),
        ('TC-WFB-010', 'Configure Field Validation - Regex', 'Verify regex validation',
         '1. Add text field\n2. Set regex pattern for email\n3. Save and test',
         'Field validates against regex pattern', 'High'),
        ('TC-WFB-011', 'Configure Calculated Field', 'Verify calculated field formula',
         '1. Add number fields A and B\n2. Add calculated field C = A + B\n3. Save and test',
         'Calculated field shows correct result', 'High'),
        ('TC-WFB-012', 'Configure Conditional Visibility', 'Verify VisibleWhen function',
         '1. Add dropdown field\n2. Add text field with VisibleWhen(dropdown == "Other")\n3. Test',
         'Field visibility changes based on condition', 'High'),
        ('TC-WFB-013', 'Configure Conditional Mandatory', 'Verify MandatoryWhen function',
         '1. Add checkbox\n2. Add field with MandatoryWhen(checkbox == true)\n3. Test',
         'Field becomes required based on condition', 'High'),
        ('TC-WFB-014', 'Add Field Group', 'Verify field grouping',
         '1. Open workflow builder\n2. Create field group\n3. Add fields to group\n4. Save',
         'Fields grouped with collapsible section', 'Medium'),
        ('TC-WFB-015', 'Reorder Fields', 'Verify drag-and-drop field reordering',
         '1. Open workflow builder\n2. Drag field to new position\n3. Save',
         'Field order updated successfully', 'Medium'),
        ('TC-WFB-016', 'Delete Field', 'Verify field deletion',
         '1. Open workflow builder\n2. Select field\n3. Click Delete\n4. Confirm',
         'Field removed from form', 'High'),
        ('TC-WFB-017', 'Configure Approval Levels', 'Verify multi-level approval setup',
         '1. Open workflow\n2. Go to Approvers tab\n3. Add Level 1 approver\n4. Add Level 2 approver\n5. Save',
         'Multiple approval levels configured', 'High'),
        ('TC-WFB-018', 'Set Dynamic Approver', 'Verify dynamic approver based on field',
         '1. Configure approver based on department field\n2. Save and test submission',
         'Approver assigned based on field value', 'High'),
        ('TC-WFB-019', 'Clone Workflow', 'Verify workflow cloning',
         '1. Select existing workflow\n2. Click Clone\n3. Enter new name\n4. Save',
         'Workflow cloned with all configurations', 'Medium'),
        ('TC-WFB-020', 'Publish Workflow', 'Verify workflow publishing',
         '1. Complete workflow configuration\n2. Click Publish\n3. Confirm',
         'Workflow published and available for submission', 'High'),
    ]

    create_test_table(doc, wf_builder_tests)

    doc.add_page_break()

    # 9. Workflow Submission Module
    doc.add_heading('9. Test Cases - Workflow Submission Module', 1)

    submission_tests = [
        ('TC-SUB-001', 'Submit New Request', 'Verify workflow submission',
         '1. Login as user\n2. Navigate to Submit Request\n3. Select workflow\n4. Fill form\n5. Submit',
         'Request submitted successfully with reference number', 'High'),
        ('TC-SUB-002', 'Submit with Attachments', 'Verify file attachment on submission',
         '1. Fill workflow form\n2. Upload attachments\n3. Submit',
         'Request submitted with attachments stored', 'High'),
        ('TC-SUB-003', 'Save as Draft', 'Verify draft saving functionality',
         '1. Start filling workflow form\n2. Click Save as Draft',
         'Draft saved successfully, can be resumed later', 'High'),
        ('TC-SUB-004', 'Resume Draft', 'Verify resuming draft submission',
         '1. Navigate to My Submissions\n2. Find draft\n3. Click Continue',
         'Draft loaded with previously entered data', 'High'),
        ('TC-SUB-005', 'Submit with Validation Errors', 'Verify validation prevents submission',
         '1. Fill form with invalid data\n2. Try to submit',
         'Validation errors displayed, submission blocked', 'High'),
        ('TC-SUB-006', 'View My Submissions', 'Verify submission list view',
         '1. Login as user\n2. Navigate to My Submissions',
         'List of user submissions displayed with status', 'High'),
        ('TC-SUB-007', 'View Submission Details', 'Verify submission detail view',
         '1. Navigate to My Submissions\n2. Click on submission',
         'Full submission details displayed', 'High'),
        ('TC-SUB-008', 'Track Submission Status', 'Verify status tracking',
         '1. View submission details\n2. Check approval timeline',
         'Current status and approval history displayed', 'High'),
        ('TC-SUB-009', 'Cancel Submission', 'Verify submission cancellation',
         '1. View pending submission\n2. Click Cancel\n3. Provide reason\n4. Confirm',
         'Submission cancelled successfully', 'Medium'),
        ('TC-SUB-010', 'Resubmit Rejected Request', 'Verify resubmission after rejection',
         '1. View rejected submission\n2. Click Resubmit\n3. Make corrections\n4. Submit',
         'Request resubmitted for approval', 'High'),
    ]

    create_test_table(doc, submission_tests)

    doc.add_page_break()

    # 10. Approval Process Module
    doc.add_heading('10. Test Cases - Approval Process Module', 1)

    approval_tests = [
        ('TC-APR-001', 'View Pending Approvals', 'Verify pending approvals list',
         '1. Login as approver\n2. Navigate to Approvals',
         'List of pending approvals displayed', 'High'),
        ('TC-APR-002', 'Approve Request', 'Verify request approval',
         '1. View pending request\n2. Review details\n3. Click Approve\n4. Add comments\n5. Confirm',
         'Request approved and moved to next level or completed', 'High'),
        ('TC-APR-003', 'Reject Request', 'Verify request rejection',
         '1. View pending request\n2. Review details\n3. Click Reject\n4. Add reason\n5. Confirm',
         'Request rejected with notification to submitter', 'High'),
        ('TC-APR-004', 'Request More Information', 'Verify return for clarification',
         '1. View pending request\n2. Click Request Info\n3. Add questions\n4. Submit',
         'Request returned to submitter for additional info', 'High'),
        ('TC-APR-005', 'Delegate Approval', 'Verify approval delegation',
         '1. View pending request\n2. Click Delegate\n3. Select delegate\n4. Confirm',
         'Approval delegated to selected user', 'Medium'),
        ('TC-APR-006', 'Bulk Approval', 'Verify multiple request approval',
         '1. Select multiple pending requests\n2. Click Bulk Approve\n3. Confirm',
         'All selected requests approved', 'Medium'),
        ('TC-APR-007', 'Approval via Email', 'Verify email approval link',
         '1. Receive approval email\n2. Click Approve link\n3. Enter credentials if prompted',
         'Request approved via email link', 'High'),
        ('TC-APR-008', 'Rejection via Email', 'Verify email rejection link',
         '1. Receive approval email\n2. Click Reject link\n3. Enter reason',
         'Request rejected via email link', 'High'),
        ('TC-APR-009', 'Multi-Level Approval Flow', 'Verify sequential approval levels',
         '1. Submit request requiring 3 approvals\n2. Level 1 approves\n3. Level 2 approves\n4. Level 3 approves',
         'Request flows through all levels correctly', 'High'),
        ('TC-APR-010', 'Approval History', 'Verify approval audit trail',
         '1. View approved/rejected request\n2. Check approval history section',
         'Complete approval history with timestamps displayed', 'High'),
        ('TC-APR-011', 'Approval Notifications', 'Verify email notifications',
         '1. Submit request\n2. Check approver receives email\n3. Approve\n4. Check submitter receives email',
         'All parties receive appropriate notifications', 'High'),
        ('TC-APR-012', 'Approval Timeout', 'Verify approval reminder/escalation',
         '1. Submit request\n2. Wait for configured timeout period',
         'Reminder sent or escalation triggered as configured', 'Medium'),
    ]

    create_test_table(doc, approval_tests)

    doc.add_page_break()

    # 11. Reports Module
    doc.add_heading('11. Test Cases - Reports Module', 1)

    report_tests = [
        ('TC-RPT-001', 'View Dashboard Reports', 'Verify dashboard statistics',
         '1. Login\n2. View dashboard',
         'Statistics and charts displayed correctly', 'High'),
        ('TC-RPT-002', 'Generate Workflow Report', 'Verify workflow-based reporting',
         '1. Navigate to Reports\n2. Select workflow\n3. Set date range\n4. Generate',
         'Report generated with workflow submissions data', 'High'),
        ('TC-RPT-003', 'Export Report to Excel', 'Verify Excel export',
         '1. Generate report\n2. Click Export to Excel',
         'Excel file downloaded with report data', 'High'),
        ('TC-RPT-004', 'Export Report to PDF', 'Verify PDF export',
         '1. Generate report\n2. Click Export to PDF',
         'PDF file downloaded with report data', 'Medium'),
        ('TC-RPT-005', 'Filter Report by Status', 'Verify status filtering',
         '1. Generate report\n2. Apply status filter (Pending/Approved/Rejected)',
         'Report filtered by selected status', 'Medium'),
        ('TC-RPT-006', 'Filter Report by Date Range', 'Verify date range filtering',
         '1. Generate report\n2. Set start and end dates',
         'Report shows data within date range only', 'Medium'),
        ('TC-RPT-007', 'Filter Report by Department', 'Verify department filtering',
         '1. Generate report\n2. Select specific department',
         'Report shows department-specific data', 'Medium'),
        ('TC-RPT-008', 'Save Report Configuration', 'Verify saving report settings',
         '1. Configure report filters\n2. Click Save Configuration\n3. Enter name',
         'Report configuration saved for future use', 'Low'),
        ('TC-RPT-009', 'Schedule Report', 'Verify scheduled report generation',
         '1. Configure report\n2. Set schedule (daily/weekly/monthly)\n3. Save',
         'Report generated and emailed on schedule', 'Low'),
        ('TC-RPT-010', 'Custom Report Builder', 'Verify custom report creation',
         '1. Navigate to Custom Reports\n2. Select fields to include\n3. Set grouping\n4. Generate',
         'Custom report generated with selected fields', 'Medium'),
    ]

    create_test_table(doc, report_tests)

    doc.add_page_break()

    # 12. Audit Trail Module
    doc.add_heading('12. Test Cases - Audit Trail Module', 1)

    audit_tests = [
        ('TC-AUD-001', 'View Audit Logs', 'Verify audit log access',
         '1. Login as admin\n2. Navigate to Audit Logs',
         'Audit logs displayed with all system activities', 'High'),
        ('TC-AUD-002', 'Audit Log - User Login', 'Verify login events logged',
         '1. User logs in\n2. Check audit log',
         'Login event recorded with timestamp and IP', 'High'),
        ('TC-AUD-003', 'Audit Log - User Logout', 'Verify logout events logged',
         '1. User logs out\n2. Check audit log',
         'Logout event recorded', 'High'),
        ('TC-AUD-004', 'Audit Log - Data Creation', 'Verify creation events logged',
         '1. Create new user/workflow/submission\n2. Check audit log',
         'Creation event logged with details', 'High'),
        ('TC-AUD-005', 'Audit Log - Data Update', 'Verify update events logged',
         '1. Update existing record\n2. Check audit log',
         'Update event logged with before/after values', 'High'),
        ('TC-AUD-006', 'Audit Log - Data Deletion', 'Verify deletion events logged',
         '1. Delete record\n2. Check audit log',
         'Deletion event logged with deleted data', 'High'),
        ('TC-AUD-007', 'Filter Audit by User', 'Verify user-based filtering',
         '1. Access audit logs\n2. Filter by specific user',
         'Only selected user activities displayed', 'Medium'),
        ('TC-AUD-008', 'Filter Audit by Date', 'Verify date-based filtering',
         '1. Access audit logs\n2. Set date range filter',
         'Logs filtered by date range', 'Medium'),
        ('TC-AUD-009', 'Filter Audit by Action', 'Verify action-based filtering',
         '1. Access audit logs\n2. Filter by action type (CREATE/UPDATE/DELETE)',
         'Logs filtered by action type', 'Medium'),
        ('TC-AUD-010', 'Export Audit Logs', 'Verify audit log export',
         '1. Access audit logs\n2. Click Export',
         'Audit logs exported to Excel/CSV', 'Medium'),
    ]

    create_test_table(doc, audit_tests)

    doc.add_page_break()

    # 13. Settings Module
    doc.add_heading('13. Test Cases - Settings Module', 1)

    settings_tests = [
        ('TC-SET-001', 'Update Email Settings', 'Verify SMTP configuration',
         '1. Login as admin\n2. Navigate to Settings > Email\n3. Configure SMTP\n4. Test connection\n5. Save',
         'Email settings saved and test email sent successfully', 'High'),
        ('TC-SET-002', 'Update System Name', 'Verify system branding',
         '1. Navigate to Settings > General\n2. Update system name\n3. Save',
         'System name updated throughout application', 'Medium'),
        ('TC-SET-003', 'Configure Session Timeout', 'Verify session timeout setting',
         '1. Navigate to Settings > Security\n2. Set timeout value\n3. Save\n4. Test',
         'Session expires after configured time', 'High'),
        ('TC-SET-004', 'Configure Password Policy', 'Verify password requirements',
         '1. Navigate to Settings > Security\n2. Set password rules\n3. Save\n4. Test with new user',
         'Password policy enforced on new passwords', 'High'),
        ('TC-SET-005', 'Configure File Upload Limits', 'Verify upload restrictions',
         '1. Navigate to Settings > Files\n2. Set max file size\n3. Set allowed types\n4. Save',
         'Upload restrictions enforced', 'Medium'),
        ('TC-SET-006', 'Enable/Disable Features', 'Verify feature toggles',
         '1. Navigate to Settings > Features\n2. Toggle feature off\n3. Save\n4. Verify feature hidden',
         'Feature visibility controlled by settings', 'Medium'),
        ('TC-SET-007', 'Configure Notification Templates', 'Verify email template customization',
         '1. Navigate to Settings > Notifications\n2. Edit template\n3. Save\n4. Trigger notification',
         'Custom template used in notifications', 'Medium'),
        ('TC-SET-008', 'Backup Configuration', 'Verify configuration backup',
         '1. Navigate to Settings > Backup\n2. Click Export Configuration',
         'Configuration exported to file', 'Low'),
        ('TC-SET-009', 'Restore Configuration', 'Verify configuration restore',
         '1. Navigate to Settings > Backup\n2. Upload configuration file\n3. Apply',
         'Configuration restored from backup', 'Low'),
        ('TC-SET-010', 'Dark Mode Toggle', 'Verify theme switching',
         '1. Click theme toggle in header\n2. Switch between light/dark mode',
         'Theme changes throughout application', 'Low'),
    ]

    create_test_table(doc, settings_tests)

    doc.add_page_break()

    # 14. API Endpoints
    doc.add_heading('14. Test Cases - API Endpoints', 1)

    api_tests = [
        ('TC-API-001', 'Authentication API', 'Verify /api/auth/login endpoint',
         '1. POST to /api/auth/login\n2. Body: {"username": "test", "password": "test123"}',
         'JWT token returned on success, 401 on failure', 'High'),
        ('TC-API-002', 'Get Users API', 'Verify /api/users endpoint',
         '1. GET /api/users with valid JWT\n2. Check response',
         'List of users returned with pagination', 'High'),
        ('TC-API-003', 'Create User API', 'Verify POST /api/users endpoint',
         '1. POST /api/users with user data\n2. Check response',
         'User created, 201 status returned', 'High'),
        ('TC-API-004', 'Get Workflows API', 'Verify /api/workflows endpoint',
         '1. GET /api/workflows with valid JWT\n2. Check response',
         'List of workflows returned', 'High'),
        ('TC-API-005', 'Submit Workflow API', 'Verify POST /api/workflow-instances',
         '1. POST with form data\n2. Check response',
         'Submission created, reference number returned', 'High'),
        ('TC-API-006', 'Approve Request API', 'Verify PUT /api/workflow-instances/{id}/approve',
         '1. PUT with approval data\n2. Check response',
         'Request approved, status updated', 'High'),
        ('TC-API-007', 'API Rate Limiting', 'Verify rate limit enforcement',
         '1. Send requests exceeding rate limit\n2. Check response',
         '429 Too Many Requests returned', 'Medium'),
        ('TC-API-008', 'API Invalid Token', 'Verify unauthorized access handling',
         '1. Send request with invalid/expired JWT\n2. Check response',
         '401 Unauthorized returned', 'High'),
        ('TC-API-009', 'API Validation Errors', 'Verify input validation',
         '1. POST with invalid data\n2. Check response',
         '400 Bad Request with validation errors', 'High'),
        ('TC-API-010', 'API CORS Headers', 'Verify CORS configuration',
         '1. Send cross-origin request\n2. Check response headers',
         'Appropriate CORS headers present', 'Medium'),
    ]

    create_test_table(doc, api_tests)

    doc.add_page_break()

    # 15. Performance Testing
    doc.add_heading('15. Performance Testing', 1)

    doc.add_heading('15.1 Load Testing Scenarios', 2)
    perf_table = doc.add_table(rows=6, cols=4)
    perf_table.style = 'Table Grid'
    perf_headers = ['Scenario', 'Concurrent Users', 'Duration', 'Expected Response Time']
    for i, header in enumerate(perf_headers):
        perf_table.rows[0].cells[i].text = header
        set_cell_shading(perf_table.rows[0].cells[i], '1976D2')
        for paragraph in perf_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    perf_data = [
        ('Login', '100', '5 minutes', '< 2 seconds'),
        ('Dashboard Load', '50', '10 minutes', '< 3 seconds'),
        ('Form Submission', '50', '10 minutes', '< 5 seconds'),
        ('Report Generation', '20', '5 minutes', '< 10 seconds'),
        ('Search Operations', '100', '5 minutes', '< 2 seconds'),
    ]
    for i, (scenario, users, duration, response) in enumerate(perf_data, 1):
        perf_table.rows[i].cells[0].text = scenario
        perf_table.rows[i].cells[1].text = users
        perf_table.rows[i].cells[2].text = duration
        perf_table.rows[i].cells[3].text = response

    doc.add_heading('15.2 Stress Testing', 2)
    doc.add_paragraph('Stress testing will be performed to identify system breaking points:')
    stress_items = [
        'Gradually increase concurrent users until response time degrades',
        'Monitor memory usage under sustained load',
        'Test database connection pool exhaustion',
        'Verify graceful degradation under extreme load'
    ]
    for item in stress_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 16. Security Testing
    doc.add_heading('16. Security Testing', 1)

    security_tests = [
        ('TC-SEC-001', 'SQL Injection Prevention', 'Verify SQL injection protection',
         '1. Enter SQL injection strings in form fields\n2. Submit\n3. Check for errors',
         'Input sanitized, no SQL execution', 'High'),
        ('TC-SEC-002', 'XSS Prevention', 'Verify cross-site scripting protection',
         '1. Enter script tags in form fields\n2. Submit\n3. View data',
         'Scripts escaped/sanitized, not executed', 'High'),
        ('TC-SEC-003', 'CSRF Protection', 'Verify CSRF token validation',
         '1. Attempt form submission without CSRF token',
         'Request rejected', 'High'),
        ('TC-SEC-004', 'Password Hashing', 'Verify passwords are hashed',
         '1. Create user\n2. Check database',
         'Password stored as hash, not plaintext', 'High'),
        ('TC-SEC-005', 'Session Fixation', 'Verify session ID regeneration',
         '1. Note session ID before login\n2. Login\n3. Check session ID',
         'New session ID generated after login', 'High'),
        ('TC-SEC-006', 'Authorization Bypass', 'Verify access control',
         '1. Login as regular user\n2. Try to access admin endpoints',
         '403 Forbidden returned', 'High'),
        ('TC-SEC-007', 'File Upload Security', 'Verify malicious file prevention',
         '1. Try to upload executable file\n2. Try to upload file with double extension',
         'Dangerous files rejected', 'High'),
        ('TC-SEC-008', 'Sensitive Data Exposure', 'Verify data masking',
         '1. Check API responses for sensitive data\n2. Check browser network tab',
         'Passwords/tokens not exposed in responses', 'High'),
        ('TC-SEC-009', 'Brute Force Protection', 'Verify login attempt limiting',
         '1. Attempt multiple failed logins\n2. Check if account locked/rate limited',
         'Account locked or rate limited after failures', 'High'),
        ('TC-SEC-010', 'HTTPS Enforcement', 'Verify secure connection',
         '1. Try to access via HTTP\n2. Check for redirect',
         'HTTP redirected to HTTPS', 'High'),
    ]

    create_test_table(doc, security_tests)

    doc.add_page_break()

    # 17. Test Summary
    doc.add_heading('17. Test Summary', 1)

    doc.add_heading('17.1 Test Case Summary', 2)
    summary_table = doc.add_table(rows=15, cols=3)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Module', 'Total Test Cases', 'Priority'),
        ('Authentication', '8', 'High'),
        ('User Management', '10', 'High'),
        ('Organization Management', '10', 'High'),
        ('Workflow Builder', '20', 'High'),
        ('Workflow Submission', '10', 'High'),
        ('Approval Process', '12', 'High'),
        ('Reports', '10', 'Medium'),
        ('Audit Trail', '10', 'Medium'),
        ('Settings', '10', 'Medium'),
        ('API Endpoints', '10', 'High'),
        ('Performance', '5', 'Medium'),
        ('Security', '10', 'High'),
        ('TOTAL', '125', '-'),
    ]
    for i, (module, count, priority) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = module
        summary_table.rows[i].cells[1].text = count
        summary_table.rows[i].cells[2].text = priority
        if i == 0 or i == len(summary_data) - 1:
            for cell in summary_table.rows[i].cells:
                set_cell_shading(cell, '1976D2' if i == 0 else 'E3F2FD')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if i == 0:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

    doc.add_heading('17.2 Sign-Off', 2)
    signoff_table = doc.add_table(rows=4, cols=3)
    signoff_table.style = 'Table Grid'
    signoff_headers = ['Role', 'Name', 'Signature & Date']
    for i, header in enumerate(signoff_headers):
        signoff_table.rows[0].cells[i].text = header
        set_cell_shading(signoff_table.rows[0].cells[i], 'E3F2FD')

    roles = ['QA Lead', 'Development Lead', 'Project Manager']
    for i, role in enumerate(roles, 1):
        signoff_table.rows[i].cells[0].text = role
        signoff_table.rows[i].cells[1].text = ''
        signoff_table.rows[i].cells[2].text = ''

    # Save document
    doc.save('tests/Sonar_Workflow_Quality_Test_Document.docx')
    print('Quality Test Document created successfully!')


def create_test_table(doc, tests):
    """Create a formatted test case table"""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Headers
    headers = ['Test ID', 'Test Name', 'Description', 'Steps', 'Expected Result', 'Priority']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '1976D2')
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for test_id, name, desc, steps, expected, priority in tests:
        row = table.add_row()
        row.cells[0].text = test_id
        row.cells[1].text = name
        row.cells[2].text = desc
        row.cells[3].text = steps
        row.cells[4].text = expected
        row.cells[5].text = priority

        # Set font size for all cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

        # Color priority cell
        if priority == 'High':
            set_cell_shading(row.cells[5], 'FFCDD2')
        elif priority == 'Medium':
            set_cell_shading(row.cells[5], 'FFF9C4')
        else:
            set_cell_shading(row.cells[5], 'C8E6C9')


def create_uat_document():
    """Create comprehensive UAT Script document"""
    doc = Document()

    # Title
    title = doc.add_heading('Sonar Workflow System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('User Acceptance Testing (UAT) Script', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document Info
    doc.add_paragraph()
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Document Version', '1.0'),
        ('Date', '2026-01-27'),
        ('Prepared By', 'QA Team'),
        ('UAT Lead', '[To be assigned]'),
        ('Status', 'Ready for UAT'),
        ('Classification', 'Internal')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        set_cell_shading(info_table.rows[i].cells[0], 'E3F2FD')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc = [
        '1. Introduction',
        '2. UAT Objectives',
        '3. UAT Scope',
        '4. UAT Schedule',
        '5. UAT Team & Responsibilities',
        '6. Test Environment',
        '7. Entry & Exit Criteria',
        '8. UAT Scenarios - End User',
        '9. UAT Scenarios - Approver',
        '10. UAT Scenarios - Administrator',
        '11. UAT Scenarios - Report User',
        '12. Defect Management',
        '13. UAT Sign-Off',
        '14. Appendix'
    ]
    for item in toc:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        'This User Acceptance Testing (UAT) document provides a comprehensive script for business users '
        'to validate that the Sonar Workflow System meets their business requirements and is ready for '
        'production deployment.'
    )
    doc.add_paragraph(
        'UAT is the final phase of testing before the system goes live. It focuses on validating the '
        'system from an end-user perspective, ensuring that business processes work as expected and '
        'the system is fit for purpose.'
    )

    # 2. UAT Objectives
    doc.add_heading('2. UAT Objectives', 1)
    objectives = [
        'Validate that the system meets documented business requirements',
        'Verify that business workflows function correctly end-to-end',
        'Confirm that the user interface is intuitive and user-friendly',
        'Ensure data accuracy and integrity throughout the system',
        'Validate email notifications are received correctly',
        'Verify that reports provide accurate business information',
        'Confirm that the system integrates properly with existing processes',
        'Identify any gaps between requirements and implementation',
        'Obtain formal business sign-off for production deployment'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # 3. UAT Scope
    doc.add_heading('3. UAT Scope', 1)
    doc.add_paragraph('The following business functions will be tested during UAT:')

    scope_table = doc.add_table(rows=9, cols=3)
    scope_table.style = 'Table Grid'
    scope_data = [
        ('Business Function', 'Description', 'Test Priority'),
        ('User Authentication', 'Login, logout, password management', 'Critical'),
        ('Workflow Submission', 'Creating and submitting workflow requests', 'Critical'),
        ('Approval Process', 'Reviewing and approving/rejecting requests', 'Critical'),
        ('Email Notifications', 'Receiving approval requests and status updates', 'High'),
        ('My Submissions', 'Tracking submitted requests and their status', 'High'),
        ('Reports & Analytics', 'Generating and viewing business reports', 'Medium'),
        ('User Profile', 'Managing personal settings and preferences', 'Low'),
        ('Administrative Tasks', 'Managing users, roles, and system configuration', 'Medium'),
    ]
    for i, (func, desc, priority) in enumerate(scope_data):
        scope_table.rows[i].cells[0].text = func
        scope_table.rows[i].cells[1].text = desc
        scope_table.rows[i].cells[2].text = priority
        if i == 0:
            for cell in scope_table.rows[i].cells:
                set_cell_shading(cell, '1976D2')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

    doc.add_page_break()

    # 4. UAT Schedule
    doc.add_heading('4. UAT Schedule', 1)

    schedule_table = doc.add_table(rows=8, cols=4)
    schedule_table.style = 'Table Grid'
    schedule_data = [
        ('Phase', 'Activity', 'Duration', 'Participants'),
        ('Preparation', 'UAT environment setup, test data preparation', '2 days', 'IT Team'),
        ('Training', 'UAT training for business users', '1 day', 'All UAT participants'),
        ('Cycle 1', 'First round of UAT execution', '5 days', 'Business users'),
        ('Defect Fix', 'Development team fixes identified issues', '3 days', 'Development team'),
        ('Cycle 2', 'Regression testing and retest of fixes', '3 days', 'Business users'),
        ('Final Review', 'Final validation and documentation', '2 days', 'UAT Lead, Stakeholders'),
        ('Sign-Off', 'Formal UAT sign-off', '1 day', 'Business owners'),
    ]
    for i, row_data in enumerate(schedule_data):
        for j, cell_data in enumerate(row_data):
            schedule_table.rows[i].cells[j].text = cell_data
        if i == 0:
            for cell in schedule_table.rows[i].cells:
                set_cell_shading(cell, '1976D2')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

    # 5. UAT Team
    doc.add_heading('5. UAT Team & Responsibilities', 1)

    team_table = doc.add_table(rows=6, cols=3)
    team_table.style = 'Table Grid'
    team_data = [
        ('Role', 'Name', 'Responsibilities'),
        ('UAT Lead', '[Name]', 'Coordinate UAT activities, track progress, manage sign-off'),
        ('Business SME', '[Name]', 'Validate business rules, provide domain expertise'),
        ('End User Tester', '[Name]', 'Execute test scenarios for submission workflows'),
        ('Approver Tester', '[Name]', 'Execute test scenarios for approval processes'),
        ('IT Support', '[Name]', 'Environment support, technical issue resolution'),
    ]
    for i, row_data in enumerate(team_data):
        for j, cell_data in enumerate(row_data):
            team_table.rows[i].cells[j].text = cell_data
        if i == 0:
            for cell in team_table.rows[i].cells:
                set_cell_shading(cell, '1976D2')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

    doc.add_page_break()

    # 6. Test Environment
    doc.add_heading('6. Test Environment', 1)

    doc.add_heading('6.1 Environment Details', 2)
    env_table = doc.add_table(rows=6, cols=2)
    env_table.style = 'Table Grid'
    env_data = [
        ('Parameter', 'Value'),
        ('Environment URL', 'https://uat.sonarworkflow.com'),
        ('Database', 'UAT PostgreSQL Instance'),
        ('Test Data', 'Anonymized production data snapshot'),
        ('Email Server', 'UAT SMTP server (sandbox mode)'),
        ('Browser Requirements', 'Chrome 120+, Firefox 120+, or Edge 120+'),
    ]
    for i, (param, value) in enumerate(env_data):
        env_table.rows[i].cells[0].text = param
        env_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in env_table.rows[i].cells:
                set_cell_shading(cell, 'E3F2FD')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_heading('6.2 Test Accounts', 2)
    accounts_table = doc.add_table(rows=6, cols=4)
    accounts_table.style = 'Table Grid'
    accounts_data = [
        ('Username', 'Password', 'Role', 'Purpose'),
        ('uat_user1', 'UatTest123!', 'End User', 'Workflow submission testing'),
        ('uat_approver1', 'UatTest123!', 'Approver', 'Level 1 approval testing'),
        ('uat_approver2', 'UatTest123!', 'Approver', 'Level 2 approval testing'),
        ('uat_admin', 'UatAdmin123!', 'Administrator', 'Admin function testing'),
        ('uat_report', 'UatTest123!', 'Report Viewer', 'Report testing'),
    ]
    for i, row_data in enumerate(accounts_data):
        for j, cell_data in enumerate(row_data):
            accounts_table.rows[i].cells[j].text = cell_data
        if i == 0:
            for cell in accounts_table.rows[i].cells:
                set_cell_shading(cell, 'E3F2FD')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_page_break()

    # 7. Entry & Exit Criteria
    doc.add_heading('7. Entry & Exit Criteria', 1)

    doc.add_heading('7.1 Entry Criteria', 2)
    entry = [
        'System testing completed with no critical or high severity defects open',
        'UAT environment deployed and accessible',
        'Test data loaded and verified',
        'UAT training completed for all participants',
        'Test accounts created and access verified',
        'UAT script document reviewed and approved',
        'Business stakeholder availability confirmed'
    ]
    for item in entry:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 Exit Criteria', 2)
    exit_criteria = [
        'All critical and high priority test scenarios executed',
        '100% of critical test cases passed',
        '95% of high priority test cases passed',
        'No critical or high severity defects remain open',
        'All medium severity defects have approved workarounds or deferred',
        'Business stakeholders have signed off on UAT completion',
        'Go-live checklist completed'
    ]
    for item in exit_criteria:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 8. UAT Scenarios - End User
    doc.add_heading('8. UAT Scenarios - End User', 1)
    doc.add_paragraph('These scenarios are to be executed by users who will submit workflow requests.')

    end_user_scenarios = [
        {
            'id': 'UAT-EU-001',
            'title': 'Login to System',
            'objective': 'Verify end user can successfully login',
            'preconditions': 'Valid user account exists',
            'steps': [
                'Navigate to the login page',
                'Enter your username: uat_user1',
                'Enter your password: UatTest123!',
                'Click the Login button',
                'Verify you are redirected to the dashboard'
            ],
            'expected': 'Dashboard displays with welcome message and user name in header',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-002',
            'title': 'View Available Workflows',
            'objective': 'Verify user can see list of available workflows',
            'preconditions': 'User is logged in',
            'steps': [
                'Click on "Submit Request" or navigate to workflow list',
                'Verify available workflows are displayed',
                'Check that workflow names and descriptions are clear'
            ],
            'expected': 'List of published workflows visible with names and descriptions',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-003',
            'title': 'Submit Leave Request Workflow',
            'objective': 'Verify user can submit a complete leave request',
            'preconditions': 'Leave Request workflow is published and available',
            'steps': [
                'Select "Leave Request" workflow',
                'Fill in all required fields:\n   - Leave Type: Annual Leave\n   - Start Date: Select future date\n   - End Date: Select date after start date\n   - Number of Days: Enter calculated days\n   - Reason: Enter reason for leave',
                'Attach supporting document if required',
                'Review the form for accuracy',
                'Click Submit button',
                'Note the reference number provided'
            ],
            'expected': 'Success message with reference number displayed. Request appears in My Submissions with "Pending" status.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-004',
            'title': 'Form Validation - Required Fields',
            'objective': 'Verify form prevents submission with missing required fields',
            'preconditions': 'User is on workflow submission form',
            'steps': [
                'Leave one or more required fields empty',
                'Attempt to submit the form',
                'Observe validation messages'
            ],
            'expected': 'Form shows clear validation errors for each missing required field. Submit is blocked.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-005',
            'title': 'Save Draft Functionality',
            'objective': 'Verify user can save incomplete form as draft',
            'preconditions': 'User is filling a workflow form',
            'steps': [
                'Start filling a workflow form',
                'Fill approximately half the fields',
                'Click "Save as Draft" button',
                'Log out and log back in',
                'Navigate to My Submissions',
                'Find and open the draft',
                'Verify previously entered data is preserved'
            ],
            'expected': 'Draft saved successfully. Data retained when resuming draft.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-006',
            'title': 'File Attachment Upload',
            'objective': 'Verify user can attach files to submission',
            'preconditions': 'Workflow has file attachment field',
            'steps': [
                'Open workflow form',
                'Click on file attachment field',
                'Select a valid file (PDF, Word, or image)',
                'Verify file is uploaded successfully',
                'Try to upload an invalid file type',
                'Verify error message for invalid file'
            ],
            'expected': 'Valid files upload successfully. Invalid files show appropriate error message.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-007',
            'title': 'Track Submission Status',
            'objective': 'Verify user can track their submission status',
            'preconditions': 'User has submitted at least one request',
            'steps': [
                'Navigate to My Submissions',
                'Locate the previously submitted request',
                'Click to view details',
                'Review the status and approval timeline',
                'Verify all submitted data is displayed correctly'
            ],
            'expected': 'Submission details show current status, approval progress, and all submitted data.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-008',
            'title': 'Receive Notification Email',
            'objective': 'Verify user receives email notifications for status changes',
            'preconditions': 'User email is configured in profile',
            'steps': [
                'Submit a workflow request',
                'Wait for approver to take action (approve/reject)',
                'Check email inbox for notification',
                'Verify email contains correct details'
            ],
            'expected': 'Email notification received with correct request details and status update.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-009',
            'title': 'Conditional Field Display',
            'objective': 'Verify form fields appear/hide based on selections',
            'preconditions': 'Workflow has conditional visibility configured',
            'steps': [
                'Open workflow form',
                'Select different values in dropdown/checkbox fields',
                'Observe additional fields appearing or hiding',
                'Verify mandatory indicators appear for conditional required fields'
            ],
            'expected': 'Fields show/hide correctly based on conditions. Required indicators update dynamically.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-EU-010',
            'title': 'Cancel Pending Submission',
            'objective': 'Verify user can cancel a pending submission',
            'preconditions': 'User has a submission in Pending status',
            'steps': [
                'Navigate to My Submissions',
                'Open a pending submission',
                'Click Cancel button',
                'Enter cancellation reason',
                'Confirm cancellation',
                'Verify status changes to Cancelled'
            ],
            'expected': 'Submission cancelled successfully. Status shows "Cancelled" with timestamp.',
            'actual': '',
            'status': ''
        },
    ]

    create_uat_scenario_tables(doc, end_user_scenarios)

    doc.add_page_break()

    # 9. UAT Scenarios - Approver
    doc.add_heading('9. UAT Scenarios - Approver', 1)
    doc.add_paragraph('These scenarios are to be executed by users who approve/reject workflow requests.')

    approver_scenarios = [
        {
            'id': 'UAT-AP-001',
            'title': 'View Pending Approvals',
            'objective': 'Verify approver can see list of pending requests',
            'preconditions': 'Logged in as approver with pending items',
            'steps': [
                'Login as uat_approver1',
                'Navigate to Approvals section',
                'View list of pending approvals',
                'Verify requests assigned to you are visible'
            ],
            'expected': 'List shows all pending requests requiring your approval with key details visible.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-002',
            'title': 'Review Request Details',
            'objective': 'Verify approver can view complete request details',
            'preconditions': 'Pending approval exists',
            'steps': [
                'Click on a pending request',
                'Review all submitted information',
                'View any attached documents',
                'Check submitter information',
                'Review approval history if multi-level'
            ],
            'expected': 'Complete request details displayed including all form data, attachments, and approval history.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-003',
            'title': 'Approve Request',
            'objective': 'Verify approver can approve a request',
            'preconditions': 'Pending request available for approval',
            'steps': [
                'Open a pending request',
                'Review the details',
                'Click Approve button',
                'Enter approval comments (optional)',
                'Confirm approval',
                'Verify request moves to next level or completes'
            ],
            'expected': 'Request approved successfully. Submitter notified. Status updated.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-004',
            'title': 'Reject Request',
            'objective': 'Verify approver can reject a request with reason',
            'preconditions': 'Pending request available',
            'steps': [
                'Open a pending request',
                'Click Reject button',
                'Enter rejection reason (required)',
                'Confirm rejection',
                'Verify submitter is notified'
            ],
            'expected': 'Request rejected. Status shows "Rejected". Submitter receives notification with reason.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-005',
            'title': 'Request More Information',
            'objective': 'Verify approver can return request for clarification',
            'preconditions': 'Pending request available',
            'steps': [
                'Open a pending request',
                'Click "Request Info" or "Return" button',
                'Enter questions/clarification needed',
                'Submit',
                'Verify request returned to submitter'
            ],
            'expected': 'Request returned to submitter for additional information. Submitter can respond and resubmit.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-006',
            'title': 'Approval via Email Link',
            'objective': 'Verify approver can approve directly from email',
            'preconditions': 'Approval request email received',
            'steps': [
                'Open approval notification email',
                'Click the Approve link in email',
                'If prompted, enter credentials',
                'Add comments if needed',
                'Confirm approval'
            ],
            'expected': 'Request approved via email link. Confirmation displayed.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-007',
            'title': 'Rejection via Email Link',
            'objective': 'Verify approver can reject directly from email',
            'preconditions': 'Approval request email received',
            'steps': [
                'Open approval notification email',
                'Click the Reject link in email',
                'Enter rejection reason',
                'Confirm rejection'
            ],
            'expected': 'Request rejected via email link. Submitter notified.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-008',
            'title': 'Multi-Level Approval Flow',
            'objective': 'Verify request flows through multiple approval levels',
            'preconditions': 'Workflow configured with multiple approval levels',
            'steps': [
                'Submit a request requiring 2+ approvals',
                'Level 1 approver approves',
                'Verify request moves to Level 2',
                'Level 2 approver approves',
                'Verify request shows as fully approved'
            ],
            'expected': 'Request flows through all approval levels correctly. Final status shows "Approved".',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-009',
            'title': 'Approval with Delegation',
            'objective': 'Verify approver can delegate to another user',
            'preconditions': 'Delegation feature enabled',
            'steps': [
                'Open a pending request',
                'Click Delegate button',
                'Search and select delegate',
                'Confirm delegation',
                'Verify request appears in delegate\'s queue'
            ],
            'expected': 'Request delegated successfully. Delegate can see and action the request.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AP-010',
            'title': 'View Approval History',
            'objective': 'Verify complete approval audit trail',
            'preconditions': 'Request with approval history exists',
            'steps': [
                'Open a completed or in-progress request',
                'Navigate to approval history section',
                'Review all approval actions',
                'Verify timestamps and comments are accurate'
            ],
            'expected': 'Complete approval history displayed with all actions, timestamps, and comments.',
            'actual': '',
            'status': ''
        },
    ]

    create_uat_scenario_tables(doc, approver_scenarios)

    doc.add_page_break()

    # 10. UAT Scenarios - Administrator
    doc.add_heading('10. UAT Scenarios - Administrator', 1)
    doc.add_paragraph('These scenarios are to be executed by system administrators.')

    admin_scenarios = [
        {
            'id': 'UAT-AD-001',
            'title': 'Create New User Account',
            'objective': 'Verify admin can create user accounts',
            'preconditions': 'Logged in as administrator',
            'steps': [
                'Navigate to User Management',
                'Click Add New User',
                'Enter user details (name, email, username)',
                'Select role and department',
                'Save the user',
                'Verify user can login with new credentials'
            ],
            'expected': 'User created successfully. New user can login and access system based on role.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-002',
            'title': 'Modify User Permissions',
            'objective': 'Verify admin can change user roles',
            'preconditions': 'User account exists',
            'steps': [
                'Navigate to User Management',
                'Find and edit existing user',
                'Change user role',
                'Save changes',
                'Verify user\'s access reflects new role'
            ],
            'expected': 'User role updated. Access permissions change immediately.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-003',
            'title': 'Deactivate User Account',
            'objective': 'Verify admin can disable user access',
            'preconditions': 'Active user account exists',
            'steps': [
                'Navigate to User Management',
                'Find user to deactivate',
                'Toggle active status to inactive',
                'Save',
                'Attempt login with deactivated account'
            ],
            'expected': 'User deactivated. Login attempt shows "Account inactive" message.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-004',
            'title': 'Configure Workflow Approvers',
            'objective': 'Verify admin can set up approval hierarchy',
            'preconditions': 'Workflow exists',
            'steps': [
                'Navigate to Workflows',
                'Edit workflow configuration',
                'Go to Approvers section',
                'Configure approval levels and approvers',
                'Save configuration',
                'Test by submitting a request'
            ],
            'expected': 'Approval hierarchy configured. Requests route to correct approvers.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-005',
            'title': 'View Audit Logs',
            'objective': 'Verify admin can access system audit trail',
            'preconditions': 'Logged in as administrator',
            'steps': [
                'Navigate to Audit Logs',
                'Review recent system activities',
                'Filter by user or action type',
                'Export audit log',
                'Verify data accuracy'
            ],
            'expected': 'Audit logs show all system activities with accurate timestamps and details.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-006',
            'title': 'Configure Email Settings',
            'objective': 'Verify admin can set up email notifications',
            'preconditions': 'Admin access to settings',
            'steps': [
                'Navigate to Settings > Email',
                'Configure SMTP settings',
                'Test connection',
                'Send test email',
                'Verify test email received'
            ],
            'expected': 'Email settings saved. Test email received successfully.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-007',
            'title': 'Manage Organization Structure',
            'objective': 'Verify admin can configure organization hierarchy',
            'preconditions': 'Admin access',
            'steps': [
                'Navigate to Organization Management',
                'Create/Edit SBU, Corporate, Department structure',
                'Assign users to organization units',
                'Verify hierarchy displays correctly'
            ],
            'expected': 'Organization structure configured correctly. Users assigned to appropriate units.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-AD-008',
            'title': 'System Backup and Restore',
            'objective': 'Verify admin can backup system configuration',
            'preconditions': 'Admin access to settings',
            'steps': [
                'Navigate to Settings > Backup',
                'Export system configuration',
                'Save backup file',
                'Make a configuration change',
                'Restore from backup',
                'Verify configuration reverted'
            ],
            'expected': 'Backup created successfully. Restore reverts configuration changes.',
            'actual': '',
            'status': ''
        },
    ]

    create_uat_scenario_tables(doc, admin_scenarios)

    doc.add_page_break()

    # 11. UAT Scenarios - Report User
    doc.add_heading('11. UAT Scenarios - Report User', 1)
    doc.add_paragraph('These scenarios are for users who generate and view reports.')

    report_scenarios = [
        {
            'id': 'UAT-RP-001',
            'title': 'View Dashboard Statistics',
            'objective': 'Verify dashboard displays accurate metrics',
            'preconditions': 'User with report access logged in',
            'steps': [
                'Login to system',
                'View dashboard',
                'Check submission counts',
                'Verify approval status breakdown',
                'Cross-check numbers with actual data'
            ],
            'expected': 'Dashboard shows accurate statistics matching actual data.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-RP-002',
            'title': 'Generate Workflow Summary Report',
            'objective': 'Verify workflow report generation',
            'preconditions': 'Submission data exists',
            'steps': [
                'Navigate to Reports',
                'Select Workflow Summary report',
                'Choose date range',
                'Select workflow type',
                'Generate report',
                'Review results'
            ],
            'expected': 'Report generated with accurate submission data for selected criteria.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-RP-003',
            'title': 'Export Report to Excel',
            'objective': 'Verify Excel export functionality',
            'preconditions': 'Report generated',
            'steps': [
                'Generate a report',
                'Click Export to Excel button',
                'Open downloaded file',
                'Verify data matches on-screen report'
            ],
            'expected': 'Excel file downloads successfully with all report data intact.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-RP-004',
            'title': 'Filter Report by Department',
            'objective': 'Verify department-based filtering',
            'preconditions': 'Data exists across departments',
            'steps': [
                'Open report interface',
                'Select specific department filter',
                'Generate report',
                'Verify only selected department data shown'
            ],
            'expected': 'Report shows only data from selected department.',
            'actual': '',
            'status': ''
        },
        {
            'id': 'UAT-RP-005',
            'title': 'Approval Turnaround Report',
            'objective': 'Verify approval time analysis report',
            'preconditions': 'Completed approvals exist',
            'steps': [
                'Navigate to Reports',
                'Select Approval Turnaround report',
                'Set parameters',
                'Generate report',
                'Review average approval times'
            ],
            'expected': 'Report shows accurate approval duration metrics by workflow/department.',
            'actual': '',
            'status': ''
        },
    ]

    create_uat_scenario_tables(doc, report_scenarios)

    doc.add_page_break()

    # 12. Defect Management
    doc.add_heading('12. Defect Management', 1)

    doc.add_heading('12.1 Defect Severity Definitions', 2)
    severity_table = doc.add_table(rows=5, cols=3)
    severity_table.style = 'Table Grid'
    severity_data = [
        ('Severity', 'Definition', 'Response Time'),
        ('Critical', 'System unusable, data loss, security breach', '4 hours'),
        ('High', 'Major feature broken, no workaround available', '1 business day'),
        ('Medium', 'Feature impaired but workaround exists', '3 business days'),
        ('Low', 'Minor issue, cosmetic, enhancement request', '5 business days'),
    ]
    for i, row_data in enumerate(severity_data):
        for j, cell_data in enumerate(row_data):
            severity_table.rows[i].cells[j].text = cell_data
        if i == 0:
            for cell in severity_table.rows[i].cells:
                set_cell_shading(cell, '1976D2')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
        elif i == 1:
            set_cell_shading(severity_table.rows[i].cells[0], 'FFCDD2')
        elif i == 2:
            set_cell_shading(severity_table.rows[i].cells[0], 'FFE0B2')
        elif i == 3:
            set_cell_shading(severity_table.rows[i].cells[0], 'FFF9C4')
        elif i == 4:
            set_cell_shading(severity_table.rows[i].cells[0], 'C8E6C9')

    doc.add_heading('12.2 Defect Reporting Template', 2)
    doc.add_paragraph('Use the following format when reporting defects:')

    defect_template = doc.add_table(rows=9, cols=2)
    defect_template.style = 'Table Grid'
    defect_fields = [
        ('Defect ID', '[Auto-generated]'),
        ('Title', '[Brief description of the issue]'),
        ('Severity', '[Critical/High/Medium/Low]'),
        ('Scenario ID', '[UAT scenario where issue found]'),
        ('Steps to Reproduce', '[Numbered steps]'),
        ('Expected Result', '[What should happen]'),
        ('Actual Result', '[What actually happened]'),
        ('Screenshots/Evidence', '[Attach screenshots]'),
        ('Reported By', '[Your name and date]'),
    ]
    for i, (field, value) in enumerate(defect_fields):
        defect_template.rows[i].cells[0].text = field
        defect_template.rows[i].cells[1].text = value
        set_cell_shading(defect_template.rows[i].cells[0], 'E3F2FD')

    doc.add_page_break()

    # 13. UAT Sign-Off
    doc.add_heading('13. UAT Sign-Off', 1)

    doc.add_heading('13.1 Test Execution Summary', 2)
    exec_table = doc.add_table(rows=6, cols=5)
    exec_table.style = 'Table Grid'
    exec_headers = ['Category', 'Total', 'Passed', 'Failed', 'Blocked']
    for i, header in enumerate(exec_headers):
        exec_table.rows[0].cells[i].text = header
        set_cell_shading(exec_table.rows[0].cells[i], '1976D2')
        for paragraph in exec_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    categories = ['End User Scenarios', 'Approver Scenarios', 'Administrator Scenarios', 'Report Scenarios', 'TOTAL']
    for i, cat in enumerate(categories, 1):
        exec_table.rows[i].cells[0].text = cat
        for j in range(1, 5):
            exec_table.rows[i].cells[j].text = ''
        if cat == 'TOTAL':
            for cell in exec_table.rows[i].cells:
                set_cell_shading(cell, 'E3F2FD')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_heading('13.2 Defect Summary', 2)
    defect_summary = doc.add_table(rows=5, cols=4)
    defect_summary.style = 'Table Grid'
    defect_headers = ['Severity', 'Found', 'Fixed', 'Open']
    for i, header in enumerate(defect_headers):
        defect_summary.rows[0].cells[i].text = header
        set_cell_shading(defect_summary.rows[0].cells[i], '1976D2')
        for paragraph in defect_summary.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    severities = ['Critical', 'High', 'Medium', 'Low']
    for i, sev in enumerate(severities, 1):
        defect_summary.rows[i].cells[0].text = sev
        for j in range(1, 4):
            defect_summary.rows[i].cells[j].text = ''

    doc.add_heading('13.3 UAT Acceptance Sign-Off', 2)
    doc.add_paragraph(
        'By signing below, we confirm that User Acceptance Testing has been completed satisfactorily '
        'and the Sonar Workflow System is approved for production deployment.'
    )

    signoff_table = doc.add_table(rows=5, cols=4)
    signoff_table.style = 'Table Grid'
    signoff_headers = ['Role', 'Name', 'Signature', 'Date']
    for i, header in enumerate(signoff_headers):
        signoff_table.rows[0].cells[i].text = header
        set_cell_shading(signoff_table.rows[0].cells[i], 'E3F2FD')
        for paragraph in signoff_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    roles = ['UAT Lead', 'Business Owner', 'IT Manager', 'Project Sponsor']
    for i, role in enumerate(roles, 1):
        signoff_table.rows[i].cells[0].text = role

    doc.add_page_break()

    # 14. Appendix
    doc.add_heading('14. Appendix', 1)

    doc.add_heading('14.1 Glossary', 2)
    glossary = [
        ('UAT', 'User Acceptance Testing - Final testing phase by business users'),
        ('SBU', 'Strategic Business Unit - Top level organizational unit'),
        ('Workflow', 'A defined business process with approval steps'),
        ('Instance', 'A single submission of a workflow form'),
        ('Approver', 'User with authority to approve/reject requests'),
    ]
    glossary_table = doc.add_table(rows=len(glossary)+1, cols=2)
    glossary_table.style = 'Table Grid'
    glossary_table.rows[0].cells[0].text = 'Term'
    glossary_table.rows[0].cells[1].text = 'Definition'
    for cell in glossary_table.rows[0].cells:
        set_cell_shading(cell, 'E3F2FD')
    for i, (term, definition) in enumerate(glossary, 1):
        glossary_table.rows[i].cells[0].text = term
        glossary_table.rows[i].cells[1].text = definition

    doc.add_heading('14.2 Contact Information', 2)
    contact_table = doc.add_table(rows=4, cols=3)
    contact_table.style = 'Table Grid'
    contact_data = [
        ('Role', 'Name', 'Contact'),
        ('UAT Lead', '[Name]', '[Email/Phone]'),
        ('Technical Support', '[Name]', '[Email/Phone]'),
        ('Project Manager', '[Name]', '[Email/Phone]'),
    ]
    for i, row_data in enumerate(contact_data):
        for j, cell_data in enumerate(row_data):
            contact_table.rows[i].cells[j].text = cell_data
        if i == 0:
            for cell in contact_table.rows[i].cells:
                set_cell_shading(cell, 'E3F2FD')

    # Save document
    doc.save('tests/Sonar_Workflow_UAT_Script.docx')
    print('UAT Script Document created successfully!')


def create_uat_scenario_tables(doc, scenarios):
    """Create formatted UAT scenario tables"""
    for scenario in scenarios:
        # Scenario header
        header = doc.add_heading(f"{scenario['id']}: {scenario['title']}", 2)

        # Scenario details table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        # Objective
        table.rows[0].cells[0].text = 'Objective'
        table.rows[0].cells[1].text = scenario['objective']
        set_cell_shading(table.rows[0].cells[0], 'E3F2FD')

        # Preconditions
        table.rows[1].cells[0].text = 'Preconditions'
        table.rows[1].cells[1].text = scenario['preconditions']
        set_cell_shading(table.rows[1].cells[0], 'E3F2FD')

        # Steps
        table.rows[2].cells[0].text = 'Test Steps'
        steps_text = '\n'.join([f"{i+1}. {step}" for i, step in enumerate(scenario['steps'])])
        table.rows[2].cells[1].text = steps_text
        set_cell_shading(table.rows[2].cells[0], 'E3F2FD')

        # Expected Result
        table.rows[3].cells[0].text = 'Expected Result'
        table.rows[3].cells[1].text = scenario['expected']
        set_cell_shading(table.rows[3].cells[0], 'E3F2FD')

        # Actual Result
        table.rows[4].cells[0].text = 'Actual Result'
        table.rows[4].cells[1].text = '[To be filled during testing]'
        set_cell_shading(table.rows[4].cells[0], 'E3F2FD')

        # Status
        table.rows[5].cells[0].text = 'Status'
        table.rows[5].cells[1].text = '[ ] Pass  [ ] Fail  [ ] Blocked'
        set_cell_shading(table.rows[5].cells[0], 'E3F2FD')

        # Tester
        table.rows[6].cells[0].text = 'Tested By / Date'
        table.rows[6].cells[1].text = ''
        set_cell_shading(table.rows[6].cells[0], 'E3F2FD')

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)

        doc.add_paragraph()  # Space between scenarios


if __name__ == '__main__':
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    # Create tests directory if not exists
    os.makedirs('tests', exist_ok=True)

    print('Creating Quality Test Document...')
    create_quality_test_document()

    print('Creating UAT Script Document...')
    create_uat_document()

    print('\nBoth documents created successfully in docs/tests/ folder!')
