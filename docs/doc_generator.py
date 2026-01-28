"""
Sonar Workflow System - User Manual Document Generator
Utility functions for creating consistent, professional .docx documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_document(title, subtitle=""):
    """Create a new document with standard styling"""
    doc = Document()

    # Set up styles
    setup_styles(doc)

    # Add title page
    add_title_page(doc, title, subtitle)

    return doc


def setup_styles(doc):
    """Set up document styles"""
    # Modify Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Modify Heading 1
    if 'Heading 1' in doc.styles:
        h1 = doc.styles['Heading 1']
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)

    # Modify Heading 2
    if 'Heading 2' in doc.styles:
        h2 = doc.styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 76, 153)

    # Modify Heading 3
    if 'Heading 3' in doc.styles:
        h3 = doc.styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 102, 204)


def add_title_page(doc, title, subtitle=""):
    """Add a title page to the document"""
    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Add subtitle if provided
    if subtitle:
        doc.add_paragraph()
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Add system name
    system_para = doc.add_paragraph()
    system_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    system_run = system_para.add_run("SONARWORKS WORKFLOW SYSTEM")
    system_run.font.size = Pt(14)
    system_run.font.color.rgb = RGBColor(0, 102, 204)

    # Add version/date info
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run("Version 1.1")
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    # Page break
    doc.add_page_break()


def add_table_of_contents(doc, sections):
    """Add a table of contents"""
    doc.add_heading('Table of Contents', level=1)

    for i, section in enumerate(sections, 1):
        para = doc.add_paragraph()
        para.add_run(f"{i}. {section}")

    doc.add_page_break()


def add_section(doc, title, level=1):
    """Add a section heading"""
    doc.add_heading(title, level=level)


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def add_bullet_list(doc, items, level=0):
    """Add a bulleted list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.25 * level)


def add_numbered_list(doc, items):
    """Add a numbered list"""
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_table(doc, headers, rows, col_widths=None):
    """Add a table with headers and data rows"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        # Set header background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        header_cells[i]._tc.get_or_add_tcPr().append(shading)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # Add spacing after table
    return table


def add_note(doc, text, note_type="NOTE"):
    """Add a note/tip/warning box"""
    para = doc.add_paragraph()

    # Add note type label
    label_run = para.add_run(f"{note_type}: ")
    label_run.bold = True
    if note_type == "WARNING":
        label_run.font.color.rgb = RGBColor(204, 0, 0)
    elif note_type == "TIP":
        label_run.font.color.rgb = RGBColor(0, 153, 0)
    else:
        label_run.font.color.rgb = RGBColor(0, 102, 204)

    # Add note text
    para.add_run(text)

    # Add indentation
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)


def add_example(doc, title, content):
    """Add an example box"""
    doc.add_paragraph()
    example_para = doc.add_paragraph()
    example_para.paragraph_format.left_indent = Inches(0.5)

    title_run = example_para.add_run(f"Example: {title}\n")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 51)

    content_run = example_para.add_run(content)
    content_run.font.size = Pt(10)
    content_run.italic = True


def add_code_block(doc, code, language=""):
    """Add a code block with monospace formatting"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)

    if language:
        lang_run = para.add_run(f"[{language}]\n")
        lang_run.font.size = Pt(9)
        lang_run.font.color.rgb = RGBColor(128, 128, 128)

    code_run = para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)


def add_step_by_step(doc, steps, title="Steps:"):
    """Add numbered step-by-step instructions"""
    if title:
        add_paragraph(doc, title, bold=True)

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph()

        # Step number
        num_run = para.add_run(f"Step {i}: ")
        num_run.bold = True
        num_run.font.color.rgb = RGBColor(0, 102, 204)

        # Step content
        para.add_run(step)


def add_image_placeholder(doc, caption, width=4):
    """Add a placeholder for an image with caption"""
    # Add placeholder text
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[Screenshot: {caption}]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    # Add caption
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(f"Figure: {caption}")
    caption_run.font.size = Pt(10)
    caption_run.italic = True


def add_field_description(doc, field_name, description, field_type="", required=False):
    """Add a field description entry"""
    para = doc.add_paragraph()

    # Field name
    name_run = para.add_run(f"{field_name}")
    name_run.bold = True

    # Required indicator
    if required:
        req_run = para.add_run(" *")
        req_run.font.color.rgb = RGBColor(204, 0, 0)

    # Field type
    if field_type:
        type_run = para.add_run(f" ({field_type})")
        type_run.font.color.rgb = RGBColor(102, 102, 102)
        type_run.font.size = Pt(10)

    para.add_run(f": {description}")


def save_document(doc, filepath):
    """Save the document to file"""
    doc.save(filepath)
    print(f"Document saved: {filepath}")


# Convenience function to add horizontal line
def add_horizontal_line(doc):
    """Add a horizontal line separator"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run("_" * 80)
    run.font.color.rgb = RGBColor(192, 192, 192)
