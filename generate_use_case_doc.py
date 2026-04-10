from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Colors matching template --
TEAL = RGBColor(0x1b, 0x7f, 0xb4)       # #1b7fb4 - accent
DARK_TEAL = RGBColor(0x17, 0x64, 0x87)   # #176487 - dark accent
NEAR_BLACK = RGBColor(0x23, 0x1f, 0x20)  # #231f20 - body text
WHITE = RGBColor(0xff, 0xff, 0xff)
LIGHT_GRAY = RGBColor(0x96, 0x96, 0x96)
MID_GRAY = RGBColor(0x60, 0x60, 0x60)

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Default style --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = NEAR_BLACK

def add_colored_heading(text, level=1, color=DARK_TEAL):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_accent_bar():
    """Add a thin teal horizontal line"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B7FB4"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def add_industry_tag(text):
    p = doc.add_paragraph()
    run = p.add_run('\u25B8 Industries: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = TEAL
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = MID_GRAY
    p.paragraph_format.space_after = Pt(14)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

# "SONAR" branding line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SONAR')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = TEAL
run.font.name = 'Calibri'

# Main title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WORKFLOW')
run.bold = True
run.font.size = Pt(48)
run.font.color.rgb = DARK_TEAL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('USE CASE GUIDE')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = TEAL

doc.add_paragraph()
add_accent_bar()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A comprehensive guide to industry-specific use cases\nfor the Sonar Workflow Management System')
run.font.size = Pt(12)
run.font.color.rgb = MID_GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.5.0')
run.font.size = Pt(10)
run.font.color.rgb = LIGHT_GRAY

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_colored_heading('Table of Contents', level=1)
add_accent_bar()

toc_items = [
    '1.  Introduction',
    '2.  System Modules Overview',
    '3.  Use Cases',
    '4.  Industry Applicability Matrix',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = NEAR_BLACK

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_colored_heading('1.  Introduction', level=1)
add_accent_bar()

doc.add_paragraph(
    'Sonar Workflow is an enterprise-grade workflow management, approval, and process automation platform. '
    'It provides configurable multi-level approval chains, leave management, project tracking, critical deadline '
    'monitoring, comprehensive reporting, and full organizational structure management.'
)
doc.add_paragraph(
    'This guide presents real-world use cases across more than 20 industries, demonstrating the versatility '
    'and depth of the platform. Each use case identifies the business process, describes how Sonar Workflow '
    'supports it, and lists the industries where it is most commonly applied.'
)

doc.add_paragraph()

# ============================================================
# 2. SYSTEM MODULES OVERVIEW
# ============================================================
add_colored_heading('2.  System Modules Overview', level=1)
add_accent_bar()

modules = [
    ('Workflow Engine', 'Configurable multi-screen forms with 40+ field types, multi-level approvals, email notifications, reminders, escalation, digital approval stamps, child workflows, financial and non-financial categories, API integrations, and SQL data tables.'),
    ('Leave Management', 'Multi-level leave approval chains per department, configurable leave types and policies, balance tracking with carry-over, public holidays, team calendars, approval reminders, and automatic escalation.'),
    ('Project Management', 'Full project lifecycle with Gantt charts, task management, budgets with variance tracking, risk and issue registers, milestone tracking, team management, and document templates.'),
    ('Critical Deadlines', 'Deadline items with recurrence patterns (monthly, quarterly, annual), priority levels, owner assignment, automated email reminders, overdue tracking, and compliance monitoring.'),
    ('Reports & Analytics', '23 built-in reports across all modules with configurable filters, organizational scope filtering, and export to Excel, CSV, and PDF.'),
    ('Administration', 'User management with RBAC, corporate/SBU/branch/department hierarchy, audit logging, email configuration supporting 12+ providers, data import/export, and per-module enable/disable toggles.'),
]

for name, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.color.rgb = TEAL
    p.add_run('\n' + desc)

doc.add_page_break()

# ============================================================
# 3. USE CASES
# ============================================================
add_colored_heading('3.  Use Cases', level=1)
add_accent_bar()
doc.add_paragraph()

use_cases = [
    {
        'title': 'Expense Claims & Reimbursements',
        'lines': [
            'Employees submit expense claims with receipts and amounts through a structured multi-screen form.',
            'Financial workflow category enables amount-based approval limits, with automatic escalation for higher-value claims to senior management.',
        ],
        'industries': 'Banking, Insurance, Manufacturing, Telecommunications, NGOs, Accounting & Audit Firms, Government, Universities, Real Estate'
    },
    {
        'title': 'Purchase Orders & Procurement',
        'lines': [
            'Procurement requests flow through department heads, finance review, and executive approval before vendor engagement.',
            'Child workflows link purchase requisitions to purchase orders and goods received notes, creating a complete procurement chain with full audit trail.',
        ],
        'industries': 'Manufacturing, Wholesale & Retail, Transport & Logistics, Health, Government, Education, Universities, Local Authorities & City Councils'
    },
    {
        'title': 'Leave Application & Approval',
        'lines': [
            'Staff apply for leave through a multi-level approval chain configured per department, with automatic balance deduction and team calendar visibility.',
            'Configurable leave types, carry-over rules, half-day support, and automated reminders ensure approvals are processed before requested dates.',
        ],
        'industries': 'All Industries'
    },
    {
        'title': 'Insurance Claims Processing',
        'lines': [
            'Policyholders or agents submit claims with supporting documentation, routed through assessors, adjusters, and approval committees based on claim type and value.',
            'SQL Table fields pull live policy data from the database to display alongside the claim form, reducing manual lookups and data entry errors.',
        ],
        'industries': 'Insurance, Life Assurance, Funeral Assurance, Health'
    },
    {
        'title': 'Loan & Credit Applications',
        'lines': [
            'Loan applications move through credit assessment, risk review, and committee approval with configurable limits per approver level.',
            'Amount-based escalation automatically routes large facility requests to senior credit committees without manual intervention.',
        ],
        'industries': 'Banking, Insurance, Life Assurance, Real Estate'
    },
    {
        'title': 'Regulatory Compliance Filings',
        'lines': [
            'Statutory submission deadlines are tracked with automated reminders, ensuring filings are completed before due dates.',
            'The Critical Deadlines module supports monthly, quarterly, and annual recurrence with priority-based alerts and compliance rate reporting.',
        ],
        'industries': 'Banking, Insurance, Telecommunications, Regulatory Bodies, Accounting & Audit Firms, Life Assurance, Funeral Assurance'
    },
    {
        'title': 'Contract Review & Legal Approvals',
        'lines': [
            'Contracts flow through legal review, compliance check, and management approval with version tracking.',
            'Digital approval seals authenticate signed-off documents, and file attachments maintain the complete contract lifecycle in one place.',
        ],
        'industries': 'Legal Practitioners, Banking, Insurance, Real Estate, Government, Telecommunications'
    },
    {
        'title': 'Staff Onboarding & HR Processes',
        'lines': [
            'New hire onboarding checklists route through HR, IT, facilities, and department heads to ensure all setup tasks are completed.',
            'Multi-screen forms capture personal details, bank information, equipment requests, and access permissions in a single workflow.',
        ],
        'industries': 'All Industries'
    },
    {
        'title': 'Budget Approval & Financial Planning',
        'lines': [
            'Annual budgets are submitted by departments, reviewed by finance, and approved by executive management.',
            'Project module budget features track allocated vs actual spending with variance analysis across cost categories.',
        ],
        'industries': 'Government, NGOs, Churches, Universities, Local Authorities & City Councils, Manufacturing, Banking'
    },
    {
        'title': 'Patient Referral & Treatment Authorization',
        'lines': [
            'Medical referrals and treatment pre-authorizations flow through clinical review and administrative approval before procedures are scheduled.',
            'Escalation rules ensure urgent cases are auto-escalated if approvals are not processed within the configured timeout.',
        ],
        'industries': 'Health, Insurance, Life Assurance'
    },
    {
        'title': 'Grant & Donor Reporting',
        'lines': [
            'Program teams submit grant utilization reports with expenditure breakdowns for review by program managers and finance before donor submission.',
            'Export capabilities generate donor-ready reports in Excel or PDF format directly from the reporting module.',
        ],
        'industries': 'NGOs, Churches, Government, Education'
    },
    {
        'title': 'Vehicle & Fleet Management',
        'lines': [
            'Vehicle requests, maintenance scheduling, and fuel authorizations are managed through configurable approval workflows.',
            'Critical Deadlines track vehicle service dates and license renewals with automated reminder emails to fleet managers.',
        ],
        'industries': 'Transport & Logistics, Government, Manufacturing, NGOs, Local Authorities & City Councils'
    },
    {
        'title': 'Quality Control & Inspection Sign-off',
        'lines': [
            'Production batches require quality inspection approval before release, with inspectors logging findings and approving or rejecting batches.',
            'Table fields capture detailed inspection measurements per item, and rejection triggers rework workflows through child workflow linking.',
        ],
        'industries': 'Manufacturing, Food & Restaurants, Health, Wholesale & Retail'
    },
    {
        'title': 'Student Enrollment & Academic Approvals',
        'lines': [
            'Student applications, course registrations, transcript requests, and academic appeals flow through departmental and faculty approval chains.',
            'Screen-level notifications alert the registrar, bursary, and department heads at each stage of the enrollment process.',
        ],
        'industries': 'Education, Universities'
    },
    {
        'title': 'Membership & Congregation Management',
        'lines': [
            'Membership applications, tithe recording, event approvals, and ministry resource requests are processed through structured approval flows.',
            'Department-based access restricts visibility to relevant church ministries while leave management handles pastoral leave scheduling.',
        ],
        'industries': 'Churches, NGOs'
    },
    {
        'title': 'Policy Underwriting & Endorsements',
        'lines': [
            'New policy applications and endorsement requests flow through underwriting assessment, risk evaluation, and management authorization.',
            'Financial category with approval limits ensures policies above threshold values are reviewed by senior underwriters automatically.',
        ],
        'industries': 'Insurance, Life Assurance, Funeral Assurance'
    },
    {
        'title': 'Shipping & Logistics Clearance',
        'lines': [
            'Shipment documentation, customs clearance forms, and delivery approvals are managed through sequential approval workflows with document attachments.',
            'API field types integrate with external tracking systems while Critical Deadlines ensure customs submission windows are not missed.',
        ],
        'industries': 'Transport & Logistics, Manufacturing, Wholesale & Retail'
    },
    {
        'title': 'Audit Findings & Remediation Tracking',
        'lines': [
            'Audit observations are logged with severity ratings and assigned to responsible parties, with remediation plans requiring management sign-off.',
            'Deadline tracking ensures remediation actions are completed within agreed timescales, with compliance reports providing auditors full visibility.',
        ],
        'industries': 'Accounting & Audit Firms, Banking, Insurance, Government, Regulatory Bodies'
    },
    {
        'title': 'Tender & Bid Management',
        'lines': [
            'Tender submissions are prepared collaboratively, reviewed by technical and commercial teams, and approved before submission.',
            'Deadline module ensures bid submission windows are met with automated reminders to all stakeholders.',
        ],
        'industries': 'Manufacturing, Government, Real Estate, Local Authorities & City Councils, Transport & Logistics'
    },
    {
        'title': 'Property & Lease Management',
        'lines': [
            'Lease applications, rent reviews, maintenance requests, and tenant approvals flow through property management approval chains.',
            'Recurring deadlines track lease renewal dates while financial workflows manage rent payment authorizations with approval limits.',
        ],
        'industries': 'Real Estate, Local Authorities & City Councils, Universities, Government'
    },
    {
        'title': 'Service Activation & Provisioning',
        'lines': [
            'Customer service requests are routed through technical validation, credit checks, and activation approval before provisioning.',
            'Escalation rules prevent SLA breaches by auto-escalating pending activations that exceed configured response timeouts.',
        ],
        'industries': 'Telecommunications, Banking, Insurance'
    },
    {
        'title': 'Food Safety & Hygiene Compliance',
        'lines': [
            'Menu changes require nutritional review and management approval, while food safety inspections are tracked with mandatory sign-off schedules.',
            'Critical Deadlines with daily or weekly recurrence ensure hygiene inspections are never missed, with table fields capturing temperature logs.',
        ],
        'industries': 'Food & Restaurants, Health, Education'
    },
    {
        'title': 'License & Permit Applications',
        'lines': [
            'License applications are submitted by citizens or businesses, reviewed by inspectors, and approved by authorized officers.',
            'Digital approval seals authenticate issued permits, while multi-level routing ensures the correct officials review each application type.',
        ],
        'industries': 'Regulatory Bodies, Local Authorities & City Councils, Government'
    },
    {
        'title': 'IT Change Management & Access Requests',
        'lines': [
            'IT change requests and system access permissions flow through technical review, security assessment, and management approval.',
            'Mandatory comments on rejection ensure requestors understand why changes were declined, with audit logs providing complete change history.',
        ],
        'industries': 'Banking, Telecommunications, Insurance, Government, Universities, Manufacturing'
    },
    {
        'title': 'Funeral Claims & Bereavement Processing',
        'lines': [
            'Death notifications trigger claims workflows with document upload requirements, flowing through verification, assessment, and payout approval.',
            'Conditional fields show different requirements based on claim type, while amount-based escalation ensures large payouts receive senior authorization.',
        ],
        'industries': 'Funeral Assurance, Life Assurance, Insurance'
    },
    {
        'title': 'Research Ethics & Proposal Approvals',
        'lines': [
            'Research proposals require ethics committee review and faculty approval before studies commence, with document attachments for protocols.',
            'Multi-screen forms capture methodology, budget justification, and ethical considerations on separate screens with sequential review.',
        ],
        'industries': 'Universities, Health, Government, NGOs'
    },
    {
        'title': 'Supplier Onboarding & Vendor Management',
        'lines': [
            'New supplier applications flow through procurement, compliance, and finance verification before activation in the vendor database.',
            'File uploads capture tax certificates, insurance documents, and references, while SQL Object fields check for duplicate vendors.',
        ],
        'industries': 'Manufacturing, Wholesale & Retail, Government, Banking, Local Authorities & City Councils'
    },
    {
        'title': 'Inventory Write-off & Stock Adjustments',
        'lines': [
            'Stock discrepancies and write-off requests require warehouse, finance, and management approval with supporting documentation.',
            'Financial workflow category tracks write-off values with approval limits ensuring high-value adjustments receive executive authorization.',
        ],
        'industries': 'Wholesale & Retail, Manufacturing, Food & Restaurants, Health'
    },
    {
        'title': 'Council Resolutions & Municipal Approvals',
        'lines': [
            'Council agenda items, infrastructure projects, and community service requests follow multi-level committee and full council approval processes.',
            'Project module tracks long-term municipal projects with budgets, milestones, and risk registers for public accountability.',
        ],
        'industries': 'Local Authorities & City Councils, Government'
    },
    {
        'title': 'Tax Return Preparation & Review',
        'lines': [
            'Client tax returns are prepared by associates, reviewed by seniors, and signed off by partners with approval tracking at each stage.',
            'Deadline tracking with statutory filing dates ensures no penalties from late submissions, while approval seals authenticate partner sign-off.',
        ],
        'industries': 'Accounting & Audit Firms, Legal Practitioners'
    },
    {
        'title': 'Incident & Complaint Management',
        'lines': [
            'Customer complaints and operational incidents are logged, investigated, and resolved through structured workflows with SLA tracking.',
            'Escalation rules auto-escalate unresolved incidents after configurable timeouts, with reminder emails keeping assigned staff accountable.',
        ],
        'industries': 'Telecommunications, Banking, Insurance, Health, Government, Transport & Logistics'
    },
    {
        'title': 'Payroll Adjustments & Overtime Authorization',
        'lines': [
            'Salary adjustments, bonus payments, and overtime hours require HR and management approval before payroll processing.',
            'Financial workflow limits ensure large adjustments are reviewed by senior management, with reports tracking authorization patterns.',
        ],
        'industries': 'All Industries'
    },
    {
        'title': 'Board Meeting Minutes & Resolution Tracking',
        'lines': [
            'Meeting minutes are drafted, reviewed by attendees, and approved by the chairperson, with action items tracked to completion.',
            'Rich text fields capture formatted minutes while deadline items track action due dates with owner assignment and reminders.',
        ],
        'industries': 'Banking, Insurance, NGOs, Churches, Government, Regulatory Bodies, Universities'
    },
    {
        'title': 'Travel & Accommodation Requests',
        'lines': [
            'Business travel requests with itinerary, accommodation, and per diem estimates flow through line management and finance approval.',
            'Child workflows link travel requests to subsequent expense claims, creating a complete trip-to-reimbursement audit trail.',
        ],
        'industries': 'NGOs, Government, Banking, Manufacturing, Accounting & Audit Firms, Universities'
    },
    {
        'title': 'Asset Disposal & Write-down Approval',
        'lines': [
            'Requests to dispose of, donate, or write down fixed assets require technical assessment, finance review, and executive approval.',
            'Table fields list multiple assets per request while financial limits ensure high-value disposals receive board-level authorization.',
        ],
        'industries': 'Government, Manufacturing, Banking, Universities, Local Authorities & City Councils, NGOs'
    },
    {
        'title': 'Project Change Requests & Scope Management',
        'lines': [
            'Scope changes, timeline extensions, and budget amendments require stakeholder and project board approval before implementation.',
            'Risk and issue tracking captures change impact while budget variance reports show the financial effect of approved changes.',
        ],
        'industries': 'Manufacturing, Telecommunications, Government, Real Estate, NGOs, Universities'
    },
]

for idx, uc in enumerate(use_cases, 1):
    # Use case number and title
    h = doc.add_heading(f'3.{idx}  {uc["title"]}', level=2)
    for run in h.runs:
        run.font.color.rgb = DARK_TEAL

    # Supporting statements
    for line in uc['lines']:
        p = doc.add_paragraph(line, style='List Bullet')

    # Industry tags
    add_industry_tag(uc['industries'])

doc.add_page_break()

# ============================================================
# 4. INDUSTRY APPLICABILITY MATRIX
# ============================================================
add_colored_heading('4.  Industry Applicability Matrix', level=1)
add_accent_bar()

doc.add_paragraph(
    'The following table maps each industry to the Sonar Workflow modules most relevant to their operations. '
    'All modules are available to every industry; this matrix highlights the primary modules typically deployed first.'
)
doc.add_paragraph()

industries_matrix = [
    ('Banking', 'Workflow, Leave, Deadlines, Projects, Reports'),
    ('Insurance', 'Workflow, Leave, Deadlines, Reports'),
    ('Life Assurance', 'Workflow, Leave, Deadlines, Reports'),
    ('Funeral Assurance', 'Workflow, Leave, Deadlines, Reports'),
    ('Manufacturing', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Transport & Logistics', 'Workflow, Leave, Deadlines, Projects, Reports'),
    ('Telecommunications', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('NGOs', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Churches', 'Workflow, Leave, Deadlines, Reports'),
    ('Wholesale & Retail', 'Workflow, Leave, Deadlines, Reports'),
    ('Legal Practitioners', 'Workflow, Leave, Deadlines, Reports'),
    ('Health', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Government', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Education', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Universities', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Accounting & Audit Firms', 'Workflow, Leave, Deadlines, Reports'),
    ('Food & Restaurants', 'Workflow, Leave, Deadlines, Reports'),
    ('Regulatory Bodies', 'Workflow, Leave, Deadlines, Reports'),
    ('Local Authorities & City Councils', 'Workflow, Leave, Projects, Deadlines, Reports'),
    ('Real Estate', 'Workflow, Leave, Projects, Deadlines, Reports'),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Style header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
hdr_cells[1].text = ''

# Header formatting
for idx, (cell, text) in enumerate(zip(hdr_cells, ['Industry', 'Primary Modules'])):
    p = cell.paragraphs[0]
    p.text = ''
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    # Set cell background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="176487"/>')
    cell._element.get_or_add_tcPr().append(shading)

# Data rows
for i, (industry, modules) in enumerate(industries_matrix):
    row_cells = table.add_row().cells
    row_cells[0].text = industry
    row_cells[1].text = modules
    # Alternate row shading
    if i % 2 == 0:
        for cell in row_cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
            cell._element.get_or_add_tcPr().append(shading)
    for cell in row_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(7)
    row.cells[1].width = Cm(9)

# Add table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="1B7FB4"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="1B7FB4"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="1B7FB4"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="1B7FB4"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D6E9F5"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D6E9F5"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

doc.add_paragraph()
doc.add_paragraph()

# Footer
add_accent_bar()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sonar Workflow v1.5.0  \u2014  Built for Enterprise Process Automation')
run.font.size = Pt(9)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

# Save
output_path = r'C:\Users\Codebreaker\CODE\Sonar workflow\Sonar_Workflow_Use_Case_Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Use cases: {len(use_cases)}')
