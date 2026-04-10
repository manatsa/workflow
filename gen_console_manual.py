from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def code(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x0f, 0xd0, 0x0f)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E1E1E')
    shd.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

def hdr_row(t, cells):
    row = t.add_row()
    for i, txt in enumerate(cells):
        c = row.cells[i]
        r = c.paragraphs[0].add_run(txt)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1976D2')
        shd.set(qn('w:val'), 'clear')
        c._element.get_or_add_tcPr().append(shd)

def trow(t, cells):
    row = t.add_row()
    for i, txt in enumerate(cells):
        r = row.cells[i].paragraphs[0].add_run(txt)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

def note(doc, text, color=RGBColor(0xE6, 0x51, 0x00)):
    p = doc.add_paragraph()
    r = p.add_run('Note: ')
    r.bold = True
    r.font.color.rgb = color
    p.add_run(text)

# ====== COVER ======
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Sona Workflow System')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
r.bold = True

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Command Console Manual')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('Ctrl + Alt + P')
r.font.size = Pt(16)
r.font.name = 'Consolas'
r.font.color.rgb = RGBColor(0x0f, 0xd0, 0x0f)

for _ in range(3):
    doc.add_paragraph()
v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('Version 1.5.0\n' + datetime.date.today().strftime('%B %d, %Y'))
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Acad Arch Solutions Pvt. Ltd.')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ====== TOC ======
doc.add_heading('Table of Contents', level=1)
toc = [
    ('1.', 'Overview'), ('2.', 'Accessing the Console'), ('3.', 'Console Interface'),
    ('4.', 'Command Reference'), ('  4.1', 'User Management'), ('  4.2', 'System Management'),
    ('  4.3', 'Backup & Restore'), ('  4.4', 'Import & Export'), ('  4.5', 'Templates'),
    ('  4.6', 'Audit'), ('  4.7', 'Help'),
    ('5.', 'Entity Names'), ('6.', 'Usage Examples'), ('  6.1', 'User Lockout Scenario'),
    ('  6.2', 'Full Backup & Restore'), ('  6.3', 'Data Migration'),
    ('  6.4', 'Maintenance Mode'), ('7.', 'Security'), ('8.', 'Troubleshooting'),
]
for num, item in toc:
    p = doc.add_paragraph()
    r = p.add_run(num + '  ' + item)
    r.font.size = Pt(11)
    if not num.startswith(' '):
        r.bold = True

doc.add_page_break()

# ====== 1. OVERVIEW ======
doc.add_heading('1. Overview', level=1)
doc.add_paragraph('The Command Console is a built-in administrative tool that provides direct access to system operations through a terminal-style interface. It allows the super user to perform maintenance tasks such as backing up the database, managing user locks, importing/exporting data, and controlling system state.')

p = doc.add_paragraph()
p.add_run('Key characteristics:').bold = True
for b in [
    'Available exclusively to the "super" user account',
    'Accessed via keyboard shortcut Ctrl + Alt + P from anywhere in the application',
    'Terminal-style dark interface with green text',
    'Commands are case-insensitive',
    'Supports quoted arguments for paths with spaces',
    'All actions are logged in the audit trail'
]:
    doc.add_paragraph(b, style='List Bullet')

# ====== 2. ACCESSING ======
doc.add_heading('2. Accessing the Console', level=1)

p = doc.add_paragraph()
p.add_run('Keyboard Shortcut: ').bold = True
r = p.add_run('Ctrl + Alt + P')
r.font.name = 'Consolas'
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x0f, 0xd0, 0x0f)

doc.add_paragraph()
doc.add_paragraph('Prerequisites:')
for b in [
    'You must be logged in as the "super" user',
    'The console will not open for any other user account, including admin',
    'Press Ctrl + Alt + P again to close, or click the X button'
]:
    doc.add_paragraph(b, style='List Bullet')

note(doc, 'The keyboard shortcut works from any page in the application. The console overlays the current page without navigating away.')

# ====== 3. INTERFACE ======
doc.add_heading('3. Console Interface', level=1)
doc.add_paragraph('The console opens as a centered overlay with a dark terminal theme:')

for item in [
    'Header bar with "Command Console" title and close (X) button',
    'Scrollable output area showing command results with color-coded text',
    'Input line at the bottom with a green ">" prompt',
    'Auto-focus on the input field when opened'
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Color coding:').bold = True

t = doc.add_table(rows=0, cols=2)
t.style = 'Table Grid'
hdr_row(t, ['Color', 'Meaning'])
trow(t, ['Green', 'Successful output / user input'])
trow(t, ['Blue', 'Informational messages'])
trow(t, ['Red', 'Error messages'])
trow(t, ['Bright Green', 'Success confirmations'])

# ====== 4. COMMAND REFERENCE ======
doc.add_page_break()
doc.add_heading('4. Command Reference', level=1)
doc.add_paragraph('All commands are case-insensitive. Arguments in <angle brackets> are required, [square brackets] are optional.')

# 4.1 User Management
doc.add_heading('4.1 User Management', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Lock-user <username>', 'Lock a user account, preventing login', 'Lock-user john'])
trow(t, ['Unlock-user <username>', 'Unlock a locked user account', 'Unlock-user john'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Lock-user').font.name = 'Consolas'
doc.add_paragraph('Immediately locks the specified user account. The user will be unable to login until unlocked. The lock reason is recorded as "Locked via command console".')
code(doc, '> Lock-user john.smith\nUser \'john.smith\' has been locked')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Unlock-user').font.name = 'Consolas'
doc.add_paragraph('Removes the lock from a user account. Also clears failed login attempt counters.')
code(doc, '> Unlock-user john.smith\nUser \'john.smith\' has been unlocked')

# 4.2 System Management
doc.add_heading('4.2 System Management', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Lock-system', 'Lock entire system (only super can login)', 'Lock-system'])
trow(t, ['Unlock-system', 'Unlock the system for all users', 'Unlock-system'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Lock-system').font.name = 'Consolas'
doc.add_paragraph('Puts the entire system into maintenance mode. Only the "super" user can login while the system is locked. All other users will see a "System is currently locked" message at the login screen. Use this before major maintenance operations.')
code(doc, '> Lock-system\nSystem has been locked. Only \'super\' user can login.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Unlock-system').font.name = 'Consolas'
doc.add_paragraph('Restores normal operation. All users can login again.')
code(doc, '> Unlock-system\nSystem has been unlocked. All users can now login.')

# 4.3 Backup & Restore
doc.add_heading('4.3 Backup & Restore', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Backup', 'Create a full database backup', 'Backup'])
trow(t, ['Restore <file>', 'Restore full database from backup', 'Restore backup_20260401.sql'])
trow(t, ['Restore <entity> <file>', 'Restore specific entity from file', 'Restore users users_backup.xlsx'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Backup').font.name = 'Consolas'
doc.add_paragraph('Creates a PostgreSQL database dump using pg_dump. The backup is saved to the configured backups directory (default: C:\\Sonar Docs\\backups\\) with a timestamped filename.')
code(doc, '> Backup\nBackup created successfully: C:\\Sonar Docs\\backups\\workflow_backup_20260401_143022.sql')

note(doc, 'PostgreSQL pg_dump must be available on the system PATH for this command to work.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Restore <file>').font.name = 'Consolas'
doc.add_paragraph('Restores the database from a SQL backup file. The file can be specified as a full path or just a filename (searched in the backups directory).')
code(doc, '> Restore workflow_backup_20260401_143022.sql\nDatabase restored successfully from: C:\\Sonar Docs\\backups\\workflow_backup_20260401_143022.sql')

p = doc.add_paragraph()
r = p.add_run('WARNING: ')
r.bold = True
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
p.add_run('Restoring a database backup will overwrite all current data. Always create a fresh backup before restoring. Lock the system first to prevent user access during restore.')

# 4.4 Import & Export
doc.add_heading('4.4 Import & Export', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Import-all [folder]', 'Import all entities from a folder', 'Import-all "C:\\data"'])
trow(t, ['Import -entities <list> [folder]', 'Import specific entities', 'Import -entities users,roles'])
trow(t, ['Export-all [folder]', 'Export all entities to folder', 'Export-all "D:\\backup"'])
trow(t, ['Export -entities <list> [folder]', 'Export specific entities', 'Export -entities users,sbus'])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Import-all').font.name = 'Consolas'
doc.add_paragraph('Imports all supported entity types from Excel files in the specified folder. If no folder is specified, uses the default imports folder (C:\\Sonar Docs\\imports\\). Files are matched by name (e.g., a file containing "user" is treated as a user import).')
code(doc, '> Import-all\nImported 156 records from default import folder\n\n> Import-all "C:\\Migration Data"\nImported 89 records from C:\\Migration Data')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Import -entities').font.name = 'Consolas'
doc.add_paragraph('Imports only the specified entities. Provide a comma-separated list (no spaces). Optionally specify a source folder.')
code(doc, '> Import -entities users,roles,sbus\nImported 45 records for entities: users, roles, sbus\n\n> Import -entities users "C:\\HR Data"\nImported 23 records for entities: users')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Export-all / Export -entities').font.name = 'Consolas'
doc.add_paragraph('Exports data to Excel files. Each entity is exported as a separate .xlsx file with a timestamp. If no folder is specified, uses the default exports folder.')
code(doc, '> Export-all\nExported all entities to: C:\\Sonar Docs\\exports\n\n> Export -entities users,roles "D:\\Backup"\nExported entities to: D:\\Backup')

# 4.5 Templates
doc.add_heading('4.5 Templates', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Templates', 'Create import templates for all entities', 'Templates'])
trow(t, ['Template -entity <name>', 'Create template for a specific entity', 'Template -entity users'])

doc.add_paragraph()
doc.add_paragraph('Templates are pre-formatted Excel files with the correct column headers for each entity type. Use these as a starting point for preparing import data.')
code(doc, '> Templates\nTemplates created in: C:\\Sonar Docs\\templates\n\n> Template -entity users\nTemplate created for entity: users')

# 4.6 Audit
doc.add_heading('4.6 Audit', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Audit [folder]', 'Extract audit trail report', 'Audit "C:\\Reports"'])

doc.add_paragraph()
doc.add_paragraph('Exports the system audit log as an Excel report. Captures all system actions including logins, data changes, approvals, and command console usage.')
code(doc, '> Audit\nAudit report extracted to: C:\\Sonar Docs\\audits\\Audit_20260401_143022.xlsx')

# 4.7 Help
doc.add_heading('4.7 Help', level=2)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Command', 'Description', 'Example'])
trow(t, ['Help', 'Show all available commands', 'Help'])
trow(t, ['Help -backup', 'Show backup/restore command help', 'Help -backup'])
trow(t, ['Help -export', 'Show import/export command help', 'Help -export'])
trow(t, ['Help -lock', 'Show lock/unlock command help', 'Help -lock'])

code(doc, '> Help\nAvailable Commands:\n\nBackup & Restore:\n  Backup                     - Create a database backup\n  Restore <file>             - Restore database from file\n  Restore <entity> <file>    - Restore specific entity\n\nImport & Export:\n  Import-all [folder]                    - Import all from folder\n  Import -entities <list> [folder]       - Import specific entities\n  Export-all [folder]                    - Export all to folder\n  Export -entities <list> [folder]       - Export specific entities\n  Templates                              - Create all import templates\n  Template -entity <name>                - Create entity template\n\nAudit:\n  Audit [folder]             - Extract audit report\n\nHelp:\n  Help -backup               - Show backup/restore help\n  Help -export               - Show import/export help')

# ====== 5. ENTITY NAMES ======
doc.add_page_break()
doc.add_heading('5. Entity Names', level=1)
doc.add_paragraph('The following entity names can be used with Import, Export, and Template commands:')

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr_row(t, ['Entity Name', 'Aliases', 'Description'])
trow(t, ['users', 'user', 'User accounts'])
trow(t, ['roles', 'role', 'User roles and permissions'])
trow(t, ['privileges', 'privilege', 'Role privileges'])
trow(t, ['sbus', 'sbu', 'Strategic Business Units'])
trow(t, ['corporates', 'corporate', 'Corporate entities'])
trow(t, ['branches', 'branch', 'Branch offices'])
trow(t, ['departments', 'department', 'Departments'])
trow(t, ['categories', 'category', 'Workflow categories'])
trow(t, ['settings', 'setting', 'System settings'])

doc.add_paragraph()
doc.add_paragraph('Multiple entities are separated by commas (no spaces):')
code(doc, '> Export -entities users,roles,sbus,branches')

# ====== 6. EXAMPLES ======
doc.add_heading('6. Usage Examples', level=1)

doc.add_heading('6.1 User Lockout Scenario', level=2)
doc.add_paragraph('A user reports suspicious activity on their account:')
code(doc, '> Lock-user suspicious.user\nUser \'suspicious.user\' has been locked\n\n> # After investigation, unlock the account\n> Unlock-user suspicious.user\nUser \'suspicious.user\' has been unlocked')

doc.add_heading('6.2 Full Backup & Restore', level=2)
doc.add_paragraph('Performing a complete database backup before a major update:')
code(doc, '> Lock-system\nSystem has been locked. Only \'super\' user can login.\n\n> Backup\nBackup created successfully: C:\\Sonar Docs\\backups\\workflow_backup_20260401_143022.sql\n\n> # Perform updates...\n\n> Unlock-system\nSystem has been unlocked. All users can now login.')

doc.add_paragraph()
doc.add_paragraph('If the update fails, restore from backup:')
code(doc, '> Lock-system\nSystem has been locked.\n\n> Restore workflow_backup_20260401_143022.sql\nDatabase restored successfully.\n\n> Unlock-system\nSystem has been unlocked.')

doc.add_heading('6.3 Data Migration', level=2)
doc.add_paragraph('Migrating data from one environment to another:')
code(doc, '> # On source system: export all data\n> Export-all "C:\\Migration"\nExported all entities to: C:\\Migration\n\n> # Copy files to target system, then:\n> Import-all "C:\\Migration"\nImported 523 records from C:\\Migration')

doc.add_paragraph()
doc.add_paragraph('Or migrate specific entities only:')
code(doc, '> Export -entities users,roles,sbus "C:\\Migration"\nExported entities to: C:\\Migration\n\n> # On target:\n> Import -entities users,roles,sbus "C:\\Migration"\nImported 87 records for entities: users, roles, sbus')

doc.add_heading('6.4 Maintenance Mode', level=2)
doc.add_paragraph('Putting the system in maintenance mode while performing administrative tasks:')
code(doc, '> Lock-system\nSystem has been locked. Only \'super\' user can login.\n\n> # Perform maintenance: export, backup, import, etc.\n> Backup\nBackup created successfully.\n\n> Templates\nTemplates created.\n\n> Audit "C:\\Monthly Reports"\nAudit report extracted.\n\n> Unlock-system\nSystem has been unlocked. All users can now login.')

# ====== 7. SECURITY ======
doc.add_heading('7. Security', level=1)
for s in [
    'The Command Console is restricted exclusively to the "super" user account',
    'No other user, including admin, can access or see the console',
    'The keyboard shortcut (Ctrl+Alt+P) is ignored for non-super users',
    'All commands executed are logged in the system audit trail',
    'The audit log records who executed the command, what command, and when',
    'System lock/unlock events are specifically logged for compliance',
    'The console communicates with the backend over the authenticated API',
    'The backend endpoint (/api/console/execute) verifies super user status before executing',
    'Database operations (backup/restore) require PostgreSQL tools on the server PATH'
]:
    doc.add_paragraph(s, style='List Bullet')

# ====== 8. TROUBLESHOOTING ======
doc.add_heading('8. Troubleshooting', level=1)

issues = [
    ('Ctrl+Alt+P does nothing', 'You must be logged in as the "super" user. The shortcut is ignored for all other accounts. Make sure you are on a Sona Workflow page (not a different browser tab).'),
    ('Backup command fails', 'Ensure pg_dump is installed and on the system PATH. On Windows, add the PostgreSQL bin directory to your PATH environment variable (e.g., C:\\Program Files\\PostgreSQL\\18\\bin).'),
    ('Restore command fails', 'Ensure psql is installed and on the system PATH. Verify the backup file exists and is accessible. The database user must have sufficient privileges.'),
    ('Import shows 0 records', 'Check that the import files are in the correct folder and have the expected naming convention (e.g., files containing "user" for user imports). Files must be .xlsx format.'),
    ('"Access denied" error', 'Only the super user can execute commands. If you see this error, you are logged in with a different account. Log out and log in as "super".'),
    ('Command not recognized', 'Commands are case-insensitive but must be spelled correctly. Type "help" to see all available commands. Check for extra spaces or typos.'),
    ('Export folder not created', 'The application needs write permissions to the specified folder. Use an absolute path and ensure the parent directory exists.'),
    ('System remains locked after unlock', 'Clear your browser cache and refresh the page. If the issue persists, check the system_state table in the database.'),
]
for title, fix in issues:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(fix)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\u00A9 2026 Acad Arch Solutions Pvt. Ltd. All rights reserved.')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('Sona Command Console Manual.docx')
print('Created: Sona Command Console Manual.docx')
