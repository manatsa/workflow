from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
import os

WNS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

def shd(cell, hx):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {WNS} w:fill="{hx}"/>'))

def cmg(cell):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcMar {WNS}><w:top w:w="50" w:type="dxa"/><w:start w:w="80" w:type="dxa"/><w:bottom w:w="50" w:type="dxa"/><w:end w:w="80" w:type="dxa"/></w:tcMar>'))

def tbl(doc, hdr, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(hdr))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i,h in enumerate(hdr):
        c=t.rows[0].cells[i]; c.text=""; p=c.paragraphs[0]; r=p.add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(255,255,255)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; shd(c,"2B579A"); cmg(c)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=""; p=c.paragraphs[0]; r=p.add_run(str(v)); r.font.size=Pt(9.5); cmg(c)
            if ri%2==1: shd(c,"E8EEF4")
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width=Inches(w)

def bul(doc,txt,bp=None,lv=0):
    p=doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent=Inches(0.25+lv*0.25); p.paragraph_format.space_after=Pt(2)
    if bp:
        r=p.add_run(bp); r.bold=True; r.font.size=Pt(10)
        r2=p.add_run(txt); r2.font.size=Pt(10)
    else:
        r=p.add_run(txt); r.font.size=Pt(10)

def bod(doc,txt):
    p=doc.add_paragraph(txt); p.paragraph_format.space_after=Pt(6)
    for r in p.runs: r.font.size=Pt(10)
    return p

def stp(doc, items):
    for i,s in enumerate(items,1):
        bul(doc, " "+s, bp="Step %d."%i)

doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10)
for lv in range(1,4): doc.styles["Heading %d"%lv].font.color.rgb=RGBColor(0x1A,0x3C,0x6E)

# COVER
for _ in range(6): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Sonar Workflow System"); r.bold=True; r.font.size=Pt(32); r.font.color.rgb=RGBColor(0x1A,0x3C,0x6E)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("User Manual"); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor(0x2B,0x57,0x9A)
doc.add_paragraph()
for txt in ["Version 1.0.0","February 2026"]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(txt); r.font.size=Pt(14)
doc.add_paragraph(); doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Prepared By"); r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Acad Arch Solutions Pvt. Ltd."); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x1A,0x3C,0x6E)
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Classification: Internal Use"); r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x99,0x33,0x33); r.bold=True
print("Cover done")

# TOC
doc.add_page_break(); doc.add_heading("Table of Contents",level=1)
toc=[("1","Introduction",["1.1  About Sonar Workflow System","1.2  Key Features","1.3  System Requirements","1.4  Accessing the System"]),("2","Getting Started",["2.1  Logging In","2.2  Forgot Password","2.3  Dashboard Overview","2.4  Navigation Menu","2.5  User Profile"]),("3","Workflow Submissions",["3.1  Viewing Available Workflows","3.2  Creating a New Submission","3.3  Saving as Draft","3.4  Submitting for Approval","3.5  Tracking Submissions","3.6  Recalling a Submission","3.7  Resubmitting After Rejection","3.8  Cloning a Submission"]),("4","Approvals",["4.1  Viewing Pending Approvals","4.2  Reviewing a Request","4.3  Approving a Request","4.4  Rejecting a Request","4.5  Escalating a Request","4.6  Email Approvals"]),("5","Project Management",["5.1  Projects Dashboard","5.2  Creating a Project","5.3  Project Detail View","5.4  Managing Team Members","5.5  Managing Tasks","5.6  Milestones","5.7  Risk Management","5.8  Issue Tracking","5.9  Budget Management","5.10  Checklists","5.11  Documents","5.12  Project Reports","5.13  Project Approval"]),("6","Reports & Analytics",["6.1  Accessing Reports","6.2  Report Categories","6.3  Generating a Report","6.4  Exporting Reports"]),("7","Settings & Personalization",["7.1  Profile Settings","7.2  Theme Customization (Admin)"]),("8","Common Tasks Quick Reference",[]),("9","Troubleshooting",["9.1  Cannot Login","9.2  Form Validation Errors","9.3  File Upload Issues","9.4  Email Approval Link Not Working","9.5  Missing Menu Items","9.6  Page Not Loading"]),("","Glossary",[])]
for num,title,subs in toc:
    label="Chapter %s: %s"%(num,title) if num else title
    p=doc.add_paragraph(); r=p.add_run(label); r.bold=True; r.font.size=Pt(11); p.paragraph_format.space_after=Pt(2)
    for s in subs:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.4); p.paragraph_format.space_after=Pt(1); r=p.add_run(s); r.font.size=Pt(10)
print("TOC done")

# CH1
doc.add_page_break(); doc.add_heading("Chapter 1: Introduction",level=1)
doc.add_heading("1.1  About Sonar Workflow System",level=2)
bod(doc,"The Sonar Workflow System is an enterprise workflow management and project management platform that enables organizations to digitize approval processes, manage projects, track tasks, and generate comprehensive reports. It provides a centralized platform for streamlining business operations, eliminating paper-based processes, and ensuring accountability at every stage of a workflow.")
bod(doc,"Whether you need to route a purchase request through multiple levels of approval, manage a complex project with dozens of tasks and milestones, or generate executive reports across your entire organization, the Sonar Workflow System has been designed to meet these needs with a modern, intuitive interface.")
doc.add_heading("1.2  Key Features",level=2)
for bp,rest in [("Workflow Automation \u2013 ","Design and deploy multi-step approval workflows with configurable screens, fields, and approval levels."),("Multi-Level Approvals \u2013 ","Support for sequential and parallel approval chains with escalation capabilities."),("Project Management \u2013 ","Full-featured project management including tasks, milestones, risks, issues, budgets, and team management."),("Reporting & Analytics \u2013 ","Over 60 built-in reports across 10 categories with charts, tables, and Excel export."),("Role-Based Access Control \u2013 ","Granular privilege system ensuring users see only what they are authorized to access."),("Email Approvals \u2013 ","Approve or reject requests directly from email without logging into the system."),("Audit Trail \u2013 ","Complete history of every action taken on every submission and project."),("Theme Customization \u2013 ","Administrators can customize the look and feel of the system including colors, fonts, and branding."),("Import / Export \u2013 ","Import workflow definitions and export report data to Excel.")]:
    bul(doc,rest,bp=bp)
doc.add_heading("1.3  System Requirements",level=2)
bod(doc,"The Sonar Workflow System is a web-based application. To use it you need:")
for r in ["A modern web browser: Google Chrome (recommended), Microsoft Edge, Mozilla Firefox, or Apple Safari.","A stable internet connection.","A screen resolution of 1280 \u00d7 720 or higher (1920 \u00d7 1080 recommended).","JavaScript enabled in your browser (enabled by default in all modern browsers)."]:
    bul(doc,r)
doc.add_heading("1.4  Accessing the System",level=2)
bod(doc,"Your system administrator will provide you with the URL of the Sonar Workflow System as well as your login credentials (username and password). Bookmark the URL in your browser for quick access. If you have not received your credentials, please contact your IT administrator or the system administrator.")
print("CH1 done")

# CH2
doc.add_page_break(); doc.add_heading("Chapter 2: Getting Started",level=1)
doc.add_heading("2.1  Logging In",level=2)
bod(doc,"To log in to the Sonar Workflow System:")
stp(doc,["Open your web browser and navigate to the system URL provided by your administrator.","On the login page, enter your username in the Username field.","Enter your password in the Password field.","Click the Login button.","If this is your first time logging in, you may be prompted to change your password."])
bod(doc,"After successful authentication you will be redirected to the Dashboard, which serves as the home page of the application.")
doc.add_heading("2.2  Forgot Password",level=2)
bod(doc,"If you have forgotten your password:")
stp(doc,["Click the Forgot Password link on the login page.","Enter your registered email address or username.","Check your email inbox for a password reset link.","Click the link and set a new password following the complexity requirements shown on screen.","Return to the login page and log in with your new password."])
doc.add_heading("2.3  Dashboard Overview",level=2)
bod(doc,"The Dashboard is the first page you see after logging in. It provides a high-level overview of your workflow and project activity.")
doc.add_heading("Dashboard Components",level=3)
for bp,rest in [("Welcome Banner \u2013 ","Displays your name and the current date."),("Workflow Statistics \u2013 ","Four summary cards showing Pending Approvals, My Submissions, Approved count, and Active Workflows."),("Project Statistics \u2013 ","Four summary cards showing Total Projects, Active Projects, Completed Projects, and Overdue Projects."),("Submissions Donut Chart \u2013 ","A visual breakdown of submission statuses (Pending, Approved, Rejected, etc.)."),("Projects by Status Chart \u2013 ","A chart showing the distribution of projects across statuses."),("Quick Actions \u2013 ","Shortcut buttons to frequently used actions such as creating a new submission or project."),("Pending Approvals List \u2013 ","A list of the most recent items awaiting your approval, with direct links to review them."),("Project Metrics \u2013 ","Key metrics including overall completion percentage and budget utilization."),("Recent Projects \u2013 ","A list of recently updated projects for quick navigation.")]:
    bul(doc,rest,bp=bp)
doc.add_heading("2.4  Navigation Menu",level=2)
bod(doc,"The main navigation is located in the left sidebar. The menu items available to you depend on your assigned role and privileges.")
tbl(doc,["Menu Item","Description","Access"],[("Home / Dashboard","Main dashboard with statistics and quick actions","All users"),("System Admin","User management, roles, privileges, organization settings","Administrators only"),("Workflow Admin","Workflow builder, screen designer, approval configuration","Workflow builders only"),("Workflows","Dynamic list of published workflows available for submission","Based on privileges"),("Projects","Project list, creation, and management","Based on privileges"),("Pending Approvals","Items awaiting your approval (badge shows count)","Approvers"),("My Submissions","All submissions you have created (badge shows count)","All users"),("Reports","Categorized report menu with 10 report categories","Based on privileges")],cw=[1.8,3.5,1.5])
doc.add_heading("2.5  User Profile",level=2)
bod(doc,"To access your profile, click the avatar or account icon located in the top-right corner of the screen. From the profile menu you can:")
for item in ["View your profile information (name, email, role, SBU).","Edit your profile details where permitted.","Change your password.","Log out of the system."]:
    bul(doc,item)
print("CH2 done")

# CH3
doc.add_page_break(); doc.add_heading("Chapter 3: Workflow Submissions",level=1)
bod(doc,"Workflows are the core of the Sonar Workflow System. A workflow defines a structured process\u2014such as a leave request, purchase order, or expense claim\u2014with configurable forms (screens), fields, and approval levels.")
doc.add_heading("3.1  Viewing Available Workflows",level=2)
bod(doc,"Navigate to the Workflows section in the left sidebar. You will see a list of all published workflows that you have permission to access. Click on a workflow name to begin a new submission.")
doc.add_heading("3.2  Creating a New Submission",level=2)
bod(doc,"To create a new workflow submission:")
stp(doc,["Click on the desired workflow name from the sidebar, or use a Quick Action shortcut on the Dashboard.","The system displays the first screen (form) of the workflow.","Fill in the fields as required. Supported field types include:"])
for ft in ["Text fields (single-line and multi-line)","Number fields (with optional min/max validation)","Date and date-time pickers","Dropdown / select lists","Radio buttons and checkboxes","File upload fields (for attachments)","Signature fields (draw or type)","Star rating fields","Table / grid fields (for tabular data entry)","Rich text editors (formatted text with bold, italic, lists, etc.)","SQL-driven dynamic dropdowns"]:
    bul(doc,ft,lv=1)
bod(doc,"Mandatory fields are marked with a red asterisk (*). The system will display validation messages if you attempt to proceed with invalid or missing data. If the workflow has multiple screens, use the Next and Previous buttons to navigate between them.")
doc.add_heading("3.3  Saving as Draft",level=2)
bod(doc,"At any point during form entry, you can click the Save as Draft button. Your submission will be saved with a status of DRAFT and will appear in your My Submissions list. You can return to a draft at any time to continue and eventually submit.")
doc.add_heading("3.4  Submitting for Approval",level=2)
bod(doc,"When you have completed all required screens and fields:")
stp(doc,["Review all entered data on the summary screen.","Click the Submit button.","The system generates a unique reference number (e.g., WF-2026-00042).","Your submission enters the approval queue and its status changes to Pending.","An email notification is sent to the designated approver(s)."])
doc.add_heading("3.5  Tracking Submissions",level=2)
bod(doc,"Navigate to My Submissions from the sidebar to view all your submissions. The list shows reference number, workflow name, status, and submission date.")
bod(doc,"You can filter submissions by status:")
for s in ["Draft","Pending","Approved","Rejected","Escalated","Cancelled","On Hold","Recalled"]:
    bul(doc,s)
bod(doc,"Click on any submission to view full details including field values, attachments, and approval history.")
doc.add_heading("3.6  Recalling a Submission",level=2)
bod(doc,"If you need to withdraw a submission that is currently pending approval:")
stp(doc,["Open the submission from My Submissions.","Click the Recall button (available only for Pending submissions).","Provide a reason for the recall.","The submission is returned with a Recalled status.","You can then edit and resubmit if needed."])
doc.add_heading("3.7  Resubmitting After Rejection",level=2)
bod(doc,"If your submission has been rejected:")
stp(doc,["Open the rejected submission from My Submissions.","Review the rejection comments.","Make the necessary corrections.","Click Resubmit to send it back into the approval queue."])
doc.add_heading("3.8  Cloning a Submission",level=2)
bod(doc,"To create a new submission based on an existing one, open any previous submission and click the Clone button. The system creates a new draft pre-populated with the same field values. You can then modify any fields before submitting.")
print("CH3 done")

# CH4
doc.add_page_break(); doc.add_heading("Chapter 4: Approvals",level=1)
bod(doc,"If you have been designated as an approver for one or more workflows, you will receive requests for review and approval. This chapter describes the approval process.")
doc.add_heading("4.1  Viewing Pending Approvals",level=2)
bod(doc,"There are multiple ways to see items awaiting your approval:")
for bp,rest in [("Dashboard \u2013 ","The Pending Approvals count is shown on a summary card, and a list of recent pending items is displayed."),("Sidebar Badge \u2013 ","The Pending Approvals menu item displays a badge with the count of items awaiting your action."),("Pending Approvals Page \u2013 ","Click Pending Approvals in the sidebar to see the full list with filtering and sorting options.")]:
    bul(doc,rest,bp=bp)
doc.add_heading("4.2  Reviewing a Request",level=2)
bod(doc,"To review a pending request:")
stp(doc,["Click on the request from the Pending Approvals list.","The system displays all submitted data organized by screen.","Review all field values, attachments, and comments.","View the approval history to see actions taken at previous levels."])
doc.add_heading("4.3  Approving a Request",level=2)
bod(doc,"To approve a request:")
stp(doc,["Click the Approve button.","Optionally add comments to provide feedback or notes.","Confirm the action.","The request moves to the next approval level or is marked as Approved if this is the final level.","The submitter and relevant stakeholders are notified."])
doc.add_heading("4.4  Rejecting a Request",level=2)
bod(doc,"To reject a request:")
stp(doc,["Click the Reject button.","Provide a reason for rejection (mandatory). Be specific so the submitter can correct and resubmit.","Confirm the action.","The submission status changes to Rejected and the submitter is notified with your comments."])
doc.add_heading("4.5  Escalating a Request",level=2)
bod(doc,"In some workflows, approvers have the option to escalate a request to a higher authority rather than approving or rejecting it.")
stp(doc,["Click the Escalate button (visible only if escalation is enabled for the workflow).","Optionally select the escalation target (higher-level approver).","Provide a reason for escalation.","The request is forwarded to the escalation target and the submitter is notified."])
doc.add_heading("4.6  Email Approvals",level=2)
bod(doc,"The Sonar Workflow System supports email-based approvals, allowing approvers to take action directly from their email inbox without logging into the system.")
bod(doc,"How email approvals work:")
stp(doc,["You receive an email with a summary of the request and action links (Approve, Reject, Review).","Click the desired action link in the email.","You are taken to a secure email approval page in your browser.","Review the submission details, optionally add comments, and confirm your action.","No full login is required \u2013 the action is authenticated via a secure, one-time-use token."])
p=bod(doc,"")
r=p.add_run("Note: "); r.bold=True; r.font.size=Pt(10)
r=p.add_run("Email approval tokens have an expiry period configured by your administrator. If the link has expired, you will need to log in to the system to take action. Each link can only be used once."); r.font.size=Pt(10)
print("CH4 done")

# Chapter 5
doc.add_page_break()
doc.add_heading("Chapter 5: Project Management", level=1)
bod(doc, "The Sonar Workflow System includes a comprehensive project management module that allows organizations to plan, execute, and track projects from initiation to closure.")

doc.add_heading("5.1  Projects Dashboard", level=2)
bod(doc, "The Projects Dashboard provides an overview of all projects:")
for item in ["Summary cards showing Total, Active, Completed, and Overdue project counts.",
             "A chart displaying the distribution of projects by status.",
             "A link to navigate to the All Projects list."]:
    bul(doc, item)

doc.add_heading("5.2  Creating a New Project", level=2)
bod(doc, "To create a new project:")
stp(doc, [
    "Navigate to Projects from the sidebar menu.",
    "Click New Project.",
    "Fill in the project details: Name, Description, Start Date, End Date, Priority, and Assigned Team.",
    "Click Save to create the project.",
])

doc.add_heading("5.3  Project Details", level=2)
bod(doc, "The project details view includes:")
for item in ["Project name, description, and status.",
             "Timeline showing start and end dates with progress indicators.",
             "List of associated tasks and their statuses.",
             "Team members assigned to the project.",
             "Activity log showing all project-related events."]:
    bul(doc, item)

doc.add_heading("5.4  Task Management", level=2)
bod(doc, "Tasks are the building blocks of projects. Each task includes:")
tbl(doc, ["Field", "Description"], [
    ["Task Name", "A short descriptive title for the task."],
    ["Assignee", "The user responsible for completing the task."],
    ["Due Date", "The deadline for task completion."],
    ["Priority", "Low, Medium, High, or Critical."],
    ["Status", "To Do, In Progress, Review, or Done."],
    ["Description", "Detailed information about the task requirements."],
], cw=[1.5, 4.5])

doc.add_heading("5.5  Project Statuses", level=2)
bod(doc, "Projects transition through several statuses during their lifecycle:")
tbl(doc, ["Status", "Description"], [
    ["Draft", "Project is being planned and is not yet active."],
    ["Active", "Project is currently in progress."],
    ["On Hold", "Project has been temporarily paused."],
    ["Completed", "All tasks have been finished and project is closed."],
    ["Cancelled", "Project has been cancelled before completion."],
], cw=[1.5, 4.5])

doc.add_heading("5.6  Milestones", level=2)
bod(doc, "Milestones mark important checkpoints within a project. They help track overall progress and ensure the project stays on schedule. Each milestone has a name, target date, and completion status.")

doc.add_heading("5.7  Project Timeline", level=2)
bod(doc, "The timeline view provides a Gantt-chart-style visualization of tasks and milestones, making it easy to see task dependencies, overlapping work, and critical-path items at a glance.")

doc.add_heading("5.8  Team Assignment", level=2)
bod(doc, "Project managers can assign team members to a project. Each member can have a specific role such as Contributor, Reviewer, or Manager. Team members receive notifications about tasks assigned to them.")

doc.add_heading("5.9  Project Activity Log", level=2)
bod(doc, "Every action taken on a project is recorded in the activity log, including task creation, status changes, comment additions, and file uploads. This provides a complete audit trail for project governance.")

doc.add_heading("5.10 Attachments", level=2)
bod(doc, "Files can be attached to projects and individual tasks. Supported file types include documents, spreadsheets, images, and PDFs. The system tracks file versions and upload history.")

doc.add_heading("5.11 Project Filters and Search", level=2)
bod(doc, "The project list supports filtering by status, priority, date range, and assigned team. A search bar allows quick lookup by project name or description keywords.")

doc.add_heading("5.12 Archiving Projects", level=2)
bod(doc, "Completed or cancelled projects can be archived to keep the active project list clean. Archived projects remain accessible for reference and reporting but do not appear in the main dashboard.")

doc.add_heading("5.13 Project Reports", level=2)
bod(doc, "Project-level reports include task completion rates, overdue task summaries, team workload distribution, and milestone achievement timelines. These can be exported for stakeholder presentations.")

print("CH5 done")
# Chapter 6
doc.add_page_break()
doc.add_heading("Chapter 6: Reports & Analytics", level=1)
bod(doc, "The Reports module provides actionable insights into workflow performance, user activity, and system usage. Reports help administrators identify bottlenecks and optimize processes.")

doc.add_heading("6.1  Workflow Summary Report", level=2)
bod(doc, "This report provides a high-level overview of all workflows:")
for item in ["Total submissions over a selected date range.",
             "Breakdown by status: Pending, Approved, Rejected, Returned.",
             "Average processing time from submission to final action.",
             "Top workflows by volume."]:
    bul(doc, item)

doc.add_heading("6.2  User Activity Report", level=2)
bod(doc, "The User Activity Report tracks individual user actions within the system:")
for item in ["Number of submissions per user.",
             "Number of approvals and rejections per approver.",
             "Average response time for each approver.",
             "Login frequency and session duration."]:
    bul(doc, item)

doc.add_heading("6.3  Approval Turnaround Report", level=2)
bod(doc, "This report focuses on the efficiency of the approval process:")
tbl(doc, ["Metric", "Description"], [
    ["Avg. Turnaround", "Average time from submission to final approval."],
    ["Bottleneck Steps", "Steps where submissions spend the most time waiting."],
    ["Escalation Rate", "Percentage of submissions that required escalation."],
    ["Rejection Rate", "Percentage of submissions rejected at each step."],
], cw=[2.0, 4.0])

doc.add_heading("6.4  Exporting Reports", level=2)
bod(doc, "All reports can be exported in multiple formats:")
for item in ["PDF format for sharing and printing.",
             "Excel (XLSX) format for further analysis.",
             "CSV format for data import into other systems."]:
    bul(doc, item)
bod(doc, "To export a report, select the desired date range and filters, then click the Export button and choose the output format.")

print("CH6 done")

# Chapter 7
doc.add_page_break()
doc.add_heading("Chapter 7: Settings & Personalization", level=1)
bod(doc, "The Settings area allows administrators and users to personalize their experience and configure system-level options.")

doc.add_heading("7.1  User Profile Settings", level=2)
bod(doc, "Each user can update their profile from the Settings page:")
for item in ["Display Name and Contact Information.",
             "Password change (current password required for verification).",
             "Notification Preferences: enable or disable email and in-app notifications.",
             "Default Landing Page: choose which page appears after login.",
             "Timezone and Date Format preferences."]:
    bul(doc, item)

doc.add_heading("7.2  System Administration Settings", level=2)
bod(doc, "Administrators have access to additional configuration options:")
tbl(doc, ["Setting", "Description"], [
    ["Organization Profile", "Company name, logo, and branding settings."],
    ["Email Configuration", "SMTP server settings for outbound notifications."],
    ["Approval Token Expiry", "Duration before email approval links expire."],
    ["Session Timeout", "Idle time before users are automatically logged out."],
    ["Audit Log Retention", "How long audit records are kept before archival."],
    ["Backup Schedule", "Automated database backup frequency and retention."],
], cw=[2.0, 4.0])

print("CH7 done")
# Chapter 8
doc.add_page_break()
doc.add_heading("Chapter 8: Common Tasks Quick Reference", level=1)
bod(doc, "Use this quick-reference table to find step-by-step instructions for the most frequently performed tasks in the Sonar Workflow System.")

tbl(doc, ["Task", "Steps"], [
    ["Submit a workflow", "Sidebar > Submissions > New > Fill form > Attach files > Submit"],
    ["Approve a request", "Sidebar > Approvals > Select item > Review > Approve / Reject"],
    ["Create a user", "Sidebar > User Management > Add User > Fill details > Save"],
    ["Assign a role", "User Management > Select user > Edit > Set Role > Save"],
    ["Create a workflow", "Sidebar > Workflow Config > New > Define screens and steps > Publish"],
    ["Track a submission", "Sidebar > Submissions > Click row > View timeline and status"],
    ["Run a report", "Sidebar > Reports > Select report type > Set filters > Generate"],
    ["Change password", "Top-right avatar > Settings > Security > Change Password"],
    ["Export data", "Reports > Generate > Click Export > Choose format (PDF/Excel/CSV)"],
    ["Archive a project", "Projects > Select project > Actions > Archive"],
], cw=[1.8, 4.2])

print("CH8 done")

# Chapter 9
doc.add_page_break()
doc.add_heading("Chapter 9: Troubleshooting", level=1)
bod(doc, "This chapter addresses common issues users may encounter and provides recommended solutions.")

doc.add_heading("9.1  Login Issues", level=2)
tbl(doc, ["Problem", "Solution"], [
    ["Incorrect password error", "Ensure Caps Lock is off. Use the Forgot Password link to reset."],
    ["Account locked", "Contact your system administrator to unlock the account."],
    ["Session expired unexpectedly", "Check the session timeout setting. Re-login and save work frequently."],
], cw=[2.5, 3.5])

doc.add_heading("9.2  Workflow Submission Errors", level=2)
tbl(doc, ["Problem", "Solution"], [
    ["Required field missing", "Check for red-highlighted fields and fill in all mandatory data."],
    ["File upload fails", "Ensure the file does not exceed the maximum size limit (default 10 MB)."],
    ["Form does not save", "Check your internet connection. Try refreshing the page and re-entering data."],
], cw=[2.5, 3.5])

doc.add_heading("9.3  Approval Issues", level=2)
tbl(doc, ["Problem", "Solution"], [
    ["Submission stuck in pending", "Verify the approver has not been deactivated. Check escalation rules."],
    ["Email approval link expired", "Login to the system and approve directly from the Approvals screen."],
    ["Wrong approver assigned", "An administrator can reassign the step or return the submission."],
], cw=[2.5, 3.5])

doc.add_heading("9.4  Report Issues", level=2)
tbl(doc, ["Problem", "Solution"], [
    ["Report shows no data", "Verify the date range and filters. Ensure you have permission to view the data."],
    ["Export fails", "Try a different format. Check that popup blockers are disabled for the site."],
], cw=[2.5, 3.5])

doc.add_heading("9.5  Performance Issues", level=2)
for item in ["Clear browser cache and cookies, then reload the application.",
             "Ensure you are using a supported browser (Chrome, Edge, Firefox).",
             "If slowness persists, contact your administrator to check server resources."]:
    bul(doc, item)

doc.add_heading("9.6  Contacting Support", level=2)
bod(doc, "If the troubleshooting steps above do not resolve your issue, contact the support team:")
for item in ["Email: support@acadarch.com",
             "Phone: +91-XXX-XXX-XXXX (Mon-Fri, 9 AM - 6 PM IST)",
             "In-app: Use the Help icon in the top navigation bar to submit a support ticket."]:
    bul(doc, item)

print("CH9 done")
# Glossary
doc.add_page_break()
doc.add_heading("Glossary", level=1)
bod(doc, "Key terms used throughout this manual:")

terms = [
    ["Workflow", "A defined sequence of steps that a business process follows from initiation to completion."],
    ["Submission", "An instance of a workflow initiated by a user, containing form data and optional attachments."],
    ["Approval Step", "A stage in the workflow where a designated approver must review and take action."],
    ["Escalation", "Automatic reassignment of a pending approval to another user after a configured timeout period."],
    ["Role", "A named set of permissions assigned to users that controls access to system features."],
    ["Privilege", "A specific permission (e.g., create workflow, approve submission) that can be granted via roles."],
    ["Screen", "A configured form layout within a workflow that defines which fields are displayed."],
    ["Field", "An individual data-entry element on a screen, such as a text box, dropdown, or date picker."],
    ["Audit Trail", "A chronological record of all actions performed on a submission for accountability."],
    ["Token Expiry", "The duration after which an email-based approval link becomes invalid."],
    ["Child Workflow", "A sub-workflow triggered automatically as part of a parent workflow process."],
]
tbl(doc, ["Term", "Definition"], terms, cw=[1.8, 4.2])

print("Glossary done")

# End matter
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run("End of Document")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Sonar Workflow System User Manual v1.0.0")
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(chr(169) + " 2026 Acad Arch Solutions Pvt. Ltd. All rights reserved.")
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# Global font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

for i in range(1, 4):
    hs = doc.styles["Heading %d" % i]
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.name = "Calibri"

# Save
out = r"C:\Users\Codebreaker\CODE\Sonar workflow\Sonar Workflow System User Manual.docx"
doc.save(out)
print("Saved to:", out)