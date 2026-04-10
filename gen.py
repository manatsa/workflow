import os,sys
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
O=r"C:\Users\Codebreaker\CODE\Sonar workflow\Sonar Workflow System Technical Manual.docx"
d=Document()
s=d.styles["Normal"];s.font.name="Calibri";s.font.size=Pt(10.5);s.paragraph_format.space_after=Pt(4)
def sh(c,x):c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{x}"/>'))
def T(h,r,w=None):
    t=d.add_table(rows=1,cols=len(h));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,x in enumerate(h):
        c=t.rows[0].cells[i];c.text='';rn=c.paragraphs[0].add_run(x);rn.bold=True;rn.font.size=Pt(9);rn.font.color.rgb=RGBColor(255,255,255);sh(c,'2E4057')
    for ri,rd in enumerate(r):
        ro=t.add_row()
        for ci,v in enumerate(rd):
            c=ro.cells[ci];c.text=str(v)
            for p in c.paragraphs:
                for rn in p.runs:rn.font.size=Pt(9)
            if ri%2==0:sh(c,'F0F4F8')
    if w:
        for ro in t.rows:
            for i,x in enumerate(w):
                if i<len(ro.cells):ro.cells[i].width=Inches(x)
def B(t):p=d.add_paragraph(style='List Bullet');p.paragraph_format.left_indent=Inches(0.25);p.add_run(t)
def C(t):p=d.add_paragraph();p.paragraph_format.left_indent=Inches(0.5);r=p.add_run(t);r.font.name='Consolas';r.font.size=Pt(9);r.font.color.rgb=RGBColor(40,40,40)
cm={'GET':RGBColor(0,128,0),'POST':RGBColor(0,80,180),'PUT':RGBColor(200,130,0),'DELETE':RGBColor(200,0,0)}
def E(title,el):
    d.add_heading(title,level=2)
    for e in el:
        p=d.add_paragraph();r=p.add_run(e[0]+' ');r.bold=True;r.font.name='Consolas';r.font.size=Pt(9.5);r.font.color.rgb=cm.get(e[0],RGBColor(0,0,0));r2=p.add_run(e[1]);r2.font.name='Consolas';r2.font.size=Pt(9.5)
        if len(e)>2:p.add_run('  \u2014  '+e[2])
def SE(e):
    p=d.add_paragraph();r=p.add_run(e[0]+' ');r.bold=True;r.font.name='Consolas';r.font.size=Pt(9.5);r2=p.add_run(e[1]);r2.font.name='Consolas';r2.font.size=Pt(9.5);p.add_run('  \u2014  '+e[2])
print('Building...')
# COVER
for _ in range(6):d.add_paragraph()
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('SONAR WORKFLOW SYSTEM');r.bold=True;r.font.size=Pt(32);r.font.color.rgb=RGBColor(30,60,90)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Technical Manual');r.bold=True;r.font.size=Pt(24);r.font.color.rgb=RGBColor(60,100,140)
d.add_paragraph()
hr=d.add_table(rows=1,cols=1);hr.alignment=WD_TABLE_ALIGNMENT.CENTER;c=hr.rows[0].cells[0];c.text='';sh(c,'2E4057');c.width=Inches(5)
d.add_paragraph()
for l,v in[('Version:','1.0.0'),('Date:','February 2026'),('Prepared By:','Acad Arch Solutions Pvt. Ltd.'),('Classification:','Confidential \u2014 Technical Staff Only')]:
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(l+' ');r.bold=True;r.font.size=Pt(12);r.font.color.rgb=RGBColor(80,80,80);p.add_run(v).font.size=Pt(12)
d.add_page_break()
# TOC
d.add_heading('Table of Contents',level=1);d.add_paragraph()
for n,t,ss in[('1','System Architecture',['1.1 Technology Stack','1.2 Project Structure','1.3 Authentication & Security','1.4 Database Schema']),('2','Backend API Reference',['2.1 Authentication','2.2 User Management','2.3 Roles & Privileges','2.4 Workflows','2.5 Workflow Instances','2.6 Email Approval','2.7 Projects','2.8 Tasks','2.9 Reports','2.10 Settings/Import','2.11 System']),('3','Database Entities',['3.1 Core Entities','3.2 Project Entities']),('4','Workflow Module Functions Table',[]),('5','Field Types Reference',[]),('6','Roles, Privileges & Access Control',['6.1 System Privileges','6.2 Default Roles','6.3 PrivilegeChecker']),('7','Enums & Status Reference',[]),('8','Data Initialization',['8.1 DataInitializer','8.2 Default Users','8.3 Default Settings']),('9','Configuration & Deployment',['9.1 Properties','9.2 Building','9.3 Running','9.4 Profiles','9.5 Frontend']),('10','Troubleshooting',['10.1 Common Issues','10.2 Log Files','10.3 DB Maintenance'])]:
    p=d.add_paragraph();r=p.add_run(f'Chapter {n}: {t}');r.bold=True;r.font.size=Pt(11)
    for x in ss:sp=d.add_paragraph();sp.paragraph_format.left_indent=Inches(0.5);sp.add_run(x).font.size=Pt(10)
d.add_page_break()
print('Cover+TOC')
# CH1
d.add_heading('Chapter 1: System Architecture',level=1)
d.add_heading('1.1 Technology Stack',level=2)
d.add_paragraph('The Sonar Workflow System is a full-stack web application built on modern enterprise technologies.')
d.add_heading('Backend',level=3)
for i in['Java 21 \u2014 Primary language with modern features (records, sealed classes, pattern matching)','Spring Boot 3.x \u2014 Application framework with auto-configuration and embedded server','Spring Security \u2014 Authentication and authorization with JWT token-based security','Spring Data JPA \u2014 Data access layer with Hibernate ORM','PostgreSQL \u2014 Primary relational database','Lombok \u2014 Boilerplate code reduction (getters, setters, constructors, builders)','Maven \u2014 Build tool and dependency management']:B(i)
d.add_heading('Frontend',level=3)
for i in['Angular 19 \u2014 Frontend framework with standalone components architecture','Angular Material \u2014 Material Design UI component library','TypeScript \u2014 Typed superset of JavaScript','RxJS \u2014 Reactive programming for asynchronous data streams']:B(i)
d.add_heading('Deployment',level=3)
for i in['Embedded Tomcat \u2014 No external application server required','Single JAR deployment \u2014 Backend and frontend packaged together','Static files served from Spring Boot \u2014 Angular dist copied to resources/static/']:B(i)
d.add_heading('1.2 Project Structure',level=2)
d.add_heading('Backend Package Structure',level=3)
C('com.sonar.workflow/\n\u251c\u2500\u2500 config/          - Configuration (Security, CORS, DataInitializer)\n\u251c\u2500\u2500 controller/      - REST API controllers\n\u251c\u2500\u2500 dto/             - Data Transfer Objects\n\u251c\u2500\u2500 entity/          - JPA entities mapped to tables\n\u251c\u2500\u2500 repository/      - Spring Data JPA repositories\n\u251c\u2500\u2500 service/         - Business logic services\n\u251c\u2500\u2500 security/        - JWT, PrivilegeChecker, SecurityConfig\n\u2514\u2500\u2500 projects/        - Project management sub-module')
d.add_heading('Frontend Structure',level=3)
C('src/app/\n\u251c\u2500\u2500 core/       - Services, guards, interceptors, models\n\u251c\u2500\u2500 features/   - dashboard, projects, workflows, users, roles, settings, reports\n\u251c\u2500\u2500 layouts/    - Main layout, sidebar, header\n\u2514\u2500\u2500 shared/     - Shared components, pipes, directives')
d.add_heading('1.3 Authentication & Security',level=2)
d.add_paragraph('The system uses a stateless JWT-based authentication model.')
for i in['JWT Token-Based Auth \u2014 Login returns a signed JWT; Bearer token in all requests.','AuthInterceptor \u2014 Angular HTTP interceptor automatically attaches JWT to outgoing requests.','401 Handling \u2014 Any 401 triggers automatic logout and redirect to login.','Role-Based Access Control (RBAC) \u2014 Users assigned roles; roles contain privileges.','Privilege-Based Authorization \u2014 @PreAuthorize("@priv.has(\'PRIVILEGE\')") on endpoints.','PrivilegeChecker Bean \u2014 Custom @priv bean. ADMIN/ROLE_ADMIN bypass all checks.','BCrypt Password Encryption \u2014 All passwords hashed with BCrypt.']:B(i)
d.add_heading('1.4 Database Schema',level=2)
for i in['PostgreSQL \u2014 Production-grade RDBMS.','JPA/Hibernate Auto-DDL \u2014 Schema from annotations (ddl-auto=update).','UUID Primary Keys \u2014 Global uniqueness.','Audit Fields \u2014 createdAt, updatedAt, createdBy on entities.','Soft Delete \u2014 Active/inactive flags where applicable.','Cascade Operations \u2014 JPA cascade for parent-child.']:B(i)
d.add_page_break()
print('Ch1')

# CH2
d.add_heading('Chapter 2: Backend API Reference',level=1)
d.add_paragraph('All endpoints are served under the base URL of the Spring Boot application (default: http://localhost:8080). Authentication is required for all endpoints except /api/auth/login.')

d.add_heading('2.1 Authentication',level=2)
E('Authentication Endpoints',[
 ('POST','/api/auth/login','Authenticate user, returns JWT token and user details'),
 ('POST','/api/auth/change-password','Change current user password'),
 ('GET','/api/auth/me','Get current authenticated user profile'),
])

d.add_heading('2.2 User Management',level=2)
E('User Endpoints',[
 ('GET','/api/users','List all users (optionally filter by active status)'),
 ('GET','/api/users/{id}','Get user by ID'),
 ('POST','/api/users','Create new user'),
 ('PUT','/api/users/{id}','Update user'),
 ('DELETE','/api/users/{id}','Delete user (soft delete)'),
 ('PUT','/api/users/{id}/toggle-active','Toggle user active/inactive status'),
 ('PUT','/api/users/{id}/reset-password','Reset user password (admin)'),
 ('GET','/api/users/by-role/{roleId}','Get users by role'),
 ('GET','/api/users/search','Search users by name or username'),
])

d.add_heading('2.3 Roles & Privileges',level=2)
E('Role Endpoints',[
 ('GET','/api/roles','List all roles'),
 ('GET','/api/roles/{id}','Get role by ID with privileges'),
 ('POST','/api/roles','Create new role'),
 ('PUT','/api/roles/{id}','Update role and its privileges'),
 ('DELETE','/api/roles/{id}','Delete role'),
])
E('Privilege Endpoints',[
 ('GET','/api/privileges','List all system privileges'),
])

d.add_heading('2.4 Workflows',level=2)
E('Workflow Definition Endpoints',[
 ('GET','/api/workflows','List all workflow definitions'),
 ('GET','/api/workflows/{id}','Get workflow by ID with full config'),
 ('POST','/api/workflows','Create new workflow definition'),
 ('PUT','/api/workflows/{id}','Update workflow definition'),
 ('DELETE','/api/workflows/{id}','Delete workflow definition'),
 ('GET','/api/workflows/{id}/export','Export workflow as JSON'),
 ('POST','/api/workflows/import','Import workflow from JSON'),
 ('POST','/api/workflows/{id}/duplicate','Duplicate an existing workflow'),
])

d.add_heading('2.5 Workflow Instances',level=2)
E('Workflow Instance Endpoints',[
 ('GET','/api/workflow-instances','List instances (filterable by status, workflow, user)'),
 ('GET','/api/workflow-instances/{id}','Get instance detail with history'),
 ('POST','/api/workflow-instances','Submit/create new workflow instance'),
 ('PUT','/api/workflow-instances/{id}','Update instance (resubmit/edit)'),
 ('POST','/api/workflow-instances/{id}/approve','Approve current step'),
 ('POST','/api/workflow-instances/{id}/reject','Reject current step'),
 ('POST','/api/workflow-instances/{id}/return','Return to previous step or initiator'),
 ('POST','/api/workflow-instances/{id}/cancel','Cancel instance'),
 ('GET','/api/workflow-instances/my-submissions','Get current user submissions'),
 ('GET','/api/workflow-instances/my-approvals','Get pending approvals for current user'),
 ('GET','/api/workflow-instances/{id}/history','Get approval history for instance'),
 ('GET','/api/workflow-instances/{id}/attachments','Get attachments for instance'),
 ('POST','/api/workflow-instances/{id}/attachments','Upload attachment to instance'),
 ('DELETE','/api/workflow-instances/{id}/attachments/{attachId}','Delete attachment'),
])

d.add_heading('2.6 Email Approval',level=2)
E('Email Approval Endpoints',[
 ('GET','/api/email-approval/approve','Approve via email token link'),
 ('GET','/api/email-approval/reject','Reject via email token link'),
])
d.add_paragraph('Email approval tokens are generated when a workflow step is pending and email notifications are enabled. Tokens expire based on the system setting email.approval.token.expiry.hours (default 72 hours).')

d.add_heading('2.7 Projects',level=2)
E('Project Endpoints',[
 ('GET','/api/projects','List all projects'),
 ('GET','/api/projects/{id}','Get project by ID'),
 ('POST','/api/projects','Create new project'),
 ('PUT','/api/projects/{id}','Update project'),
 ('DELETE','/api/projects/{id}','Delete project'),
 ('GET','/api/projects/{id}/members','Get project members'),
 ('POST','/api/projects/{id}/members','Add member to project'),
 ('DELETE','/api/projects/{id}/members/{memberId}','Remove member from project'),
])

d.add_heading('2.8 Tasks',level=2)
E('Task Endpoints',[
 ('GET','/api/tasks','List tasks (filter by project, assignee, status)'),
 ('GET','/api/tasks/{id}','Get task by ID'),
 ('POST','/api/tasks','Create new task'),
 ('PUT','/api/tasks/{id}','Update task'),
 ('DELETE','/api/tasks/{id}','Delete task'),
 ('PUT','/api/tasks/{id}/status','Update task status'),
 ('GET','/api/tasks/my-tasks','Get tasks assigned to current user'),
])

d.add_heading('2.9 Reports',level=2)
E('Report Endpoints',[
 ('GET','/api/reports/workflow-summary','Workflow summary statistics'),
 ('GET','/api/reports/approval-metrics','Approval turnaround metrics'),
 ('GET','/api/reports/user-activity','User activity report'),
 ('GET','/api/reports/project-summary','Project status summary'),
])

d.add_heading('2.10 Settings / Import',level=2)
E('Settings Endpoints',[
 ('GET','/api/settings','Get all system settings'),
 ('GET','/api/settings/{key}','Get setting by key'),
 ('PUT','/api/settings/{key}','Update setting value'),
 ('POST','/api/settings','Create new setting'),
])

d.add_heading('2.11 System',level=2)
E('System Endpoints',[
 ('GET','/api/system/health','System health check'),
 ('GET','/api/system/info','System info and version'),
])
d.add_page_break()
print('Ch2')

# CH3
d.add_heading('Chapter 3: Database Entities',level=1)
d.add_heading('3.1 Core Entities',level=2)
d.add_paragraph('The following tables list the core JPA entities mapped to PostgreSQL tables.')

T(['Entity','Table','Key Fields','Description'],[
 ('User','users','id, username, email, password, fullName, active, role','System user accounts'),
 ('Role','roles','id, name, description, privileges','User roles for RBAC'),
 ('Privilege','privileges','id, name, description','Granular permissions'),
 ('Workflow','workflows','id, name, description, screens, approvalChain, active','Workflow definitions with screen/field config'),
 ('WorkflowInstance','workflow_instances','id, workflow, initiator, currentStep, status, data','Running workflow instances'),
 ('ApprovalHistory','approval_history','id, instance, step, approver, action, comments, timestamp','Audit trail of approval actions'),
 ('Attachment','attachments','id, instance, fileName, fileType, filePath, uploadedBy','File attachments on instances'),
 ('EmailApprovalToken','email_approval_tokens','id, instance, token, approver, action, used, expiresAt','Tokens for email-based approval/rejection'),
 ('Setting','settings','id, key, value, description','System configuration key-value pairs'),
 ('Notification','notifications','id, user, title, message, read, type, referenceId','User notifications'),
 ('AuditLog','audit_logs','id, user, action, entityType, entityId, details, timestamp','System audit log'),
],[2,2,3.5,3])

d.add_heading('3.2 Project Entities',level=2)
T(['Entity','Table','Key Fields','Description'],[
 ('Project','projects','id, name, description, status, startDate, endDate, owner','Project definitions'),
 ('ProjectMember','project_members','id, project, user, role','Project team membership'),
 ('Task','tasks','id, project, title, description, assignee, status, priority, dueDate','Project tasks'),
 ('TaskComment','task_comments','id, task, author, content, timestamp','Comments on tasks'),
],[2,2,3.5,3])
d.add_page_break()
print('Ch3')

