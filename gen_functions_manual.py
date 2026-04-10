from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def code(doc, text, sz=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor(0x0f, 0xc0, 0x0f)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1E1E1E')
    shd.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

def hdr(t, cells):
    row = t.add_row()
    for i, txt in enumerate(cells):
        c = row.cells[i]
        r = c.paragraphs[0].add_run(txt)
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1976D2')
        shd.set(qn('w:val'), 'clear')
        c._element.get_or_add_tcPr().append(shd)

def row(t, cells):
    r = t.add_row()
    for i, txt in enumerate(cells):
        run = r.cells[i].paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

def func_entry(doc, name, syntax, desc, default_ex, validation_ex):
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
    doc.add_paragraph(desc)
    code(doc, 'Syntax: ' + syntax)
    if default_ex:
        p = doc.add_paragraph()
        r = p.add_run('Default Value: ')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        code(doc, default_ex)
    if validation_ex:
        p = doc.add_paragraph()
        r = p.add_run('Validation: ')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
        code(doc, validation_ex)

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Sona Workflow System')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
r.bold = True
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Functions Library Manual')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
doc.add_paragraph()
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('Default Values | Validation | Transformation')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
for _ in range(3): doc.add_paragraph()
v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('Version 1.5.0\n' + datetime.date.today().strftime('%B %d, %Y'))
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run('Acad Arch Solutions Pvt. Ltd.')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

# ====== TOC ======
doc.add_heading('Table of Contents', level=1)
toc = [
    ('1.', 'Introduction'), ('2.', 'Where Functions Are Used'),
    ('  2.1', 'Default Value Expressions'), ('  2.2', 'Validation Expressions'),
    ('  2.3', 'Transformation Expressions'), ('  2.4', 'Visibility Expressions'),
    ('3.', 'Syntax Basics'), ('4.', 'String Functions'), ('5.', 'Number Functions'),
    ('6.', 'Date & Time Functions'), ('7.', 'Boolean & Logic Functions'),
    ('8.', 'Validation Functions'), ('9.', 'Transformation Functions'),
    ('10.', 'Table Functions'), ('11.', 'Utility Functions'),
    ('12.', 'User & Context Functions'), ('13.', 'Combining Functions'),
    ('14.', 'Quick Reference'),
]
for num, item in toc:
    p = doc.add_paragraph()
    r = p.add_run(num + '  ' + item)
    r.font.size = Pt(11)
    if not num.startswith(' '): r.bold = True
doc.add_page_break()

# ====== 1. INTRO ======
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('The Sona Workflow System includes a comprehensive Functions Library with over 200 functions that can be used in form field configurations. Functions enable dynamic default values, complex validation rules, data transformations, and conditional field visibility.')
doc.add_paragraph('All functions work consistently across all four usage contexts: Default Values, Validation, Transformation, and Visibility expressions.')

# ====== 2. WHERE USED ======
doc.add_heading('2. Where Functions Are Used', level=1)

doc.add_heading('2.1 Default Value Expressions', level=2)
doc.add_paragraph('Set in the field\'s "Default Value" property in the Workflow Builder. Computes a value when the form loads.')
code(doc, 'Examples:\n  TODAY()                          -> Current date\n  CONCAT(@{firstName}, " ", @{lastName})  -> Full name\n  UPPER(@{code})                   -> Uppercase code\n  IF(@{amount} > 1000, "High", "Low")     -> Conditional')

doc.add_heading('2.2 Validation Expressions', level=2)
doc.add_paragraph('Set in the field\'s "Validation Expression" in the Validation & Transformation section. Functions that return false trigger an error. Evaluated on Next and Submit.')
code(doc, 'Examples:\n  Required()                       -> Field must not be empty\n  Required() AND MinLength(5)      -> Required + min 5 chars\n  IS_PAST(@{dateOfLoss})           -> Must be past date\n  ValidWhen(BETWEEN(@{age}, 18, 65), "Age must be 18-65")\n  IN(@{currency}, "USD", "EUR", "GBP")  -> Must be in list')

doc.add_heading('2.3 Transformation Expressions', level=2)
doc.add_paragraph('Set in the field\'s "Transform Expression" in the Validation & Transformation section. Transforms the value before submission.')
code(doc, 'Examples:\n  UPPER()                          -> Convert to uppercase\n  TRIM() AND UPPER()               -> Trim then uppercase\n  ROUND(2)                         -> Round to 2 decimals\n  CAPITALIZE()                     -> Capitalize each word')

doc.add_heading('2.4 Visibility Expressions', level=2)
doc.add_paragraph('Set in the field\'s "Visibility Expression". Controls when a field is shown/hidden. Use @{fieldName} to reference other fields.')
code(doc, 'Examples:\n  @{country} == \'US\'               -> Show only for US\n  @{amount} > 1000                 -> Show for large amounts\n  @{type} == \'Other\'               -> Show for "Other" type')

# ====== 3. SYNTAX ======
doc.add_heading('3. Syntax Basics', level=1)
for item in [
    'Functions are case-insensitive: TODAY() = today() = Today()',
    'Reference other fields with @{fieldName}: @{firstName}, @{amount}',
    'String arguments in double quotes: CONTAINS(@{name}, "John")',
    'Numbers without quotes: Min(0), Range(1, 100)',
    'Combine validations with AND: Required() AND Email()',
    'Custom error messages as last quoted argument: Required("Please fill this in")',
    'Nested functions: UPPER(TRIM(@{name}))',
    'All functions work in all four contexts (default, validation, transform, visibility)'
]:
    doc.add_paragraph(item, style='List Bullet')

# ====== 4. STRING FUNCTIONS ======
doc.add_page_break()
doc.add_heading('4. String Functions', level=1)

funcs = [
    ('UPPER', 'UPPER(text)', 'Converts text to uppercase', 'UPPER(@{code})', 'IS_UPPERCASE()'),
    ('LOWER', 'LOWER(text)', 'Converts text to lowercase', 'LOWER(@{email})', 'IS_LOWERCASE()'),
    ('TRIM', 'TRIM(text)', 'Removes leading/trailing whitespace', 'TRIM(@{name})', 'TRIM()'),
    ('CONCAT', 'CONCAT(a, b, ...)', 'Joins text strings together', 'CONCAT(@{first}, " ", @{last})', None),
    ('CONCAT_WS', 'CONCAT_WS(sep, a, b, ...)', 'Joins with separator, skipping empty', 'CONCAT_WS(", ", @{city}, @{state})', None),
    ('LENGTH', 'LENGTH(text)', 'Returns character count', 'LENGTH(@{code})', 'ValidWhen(LENGTH(@{code}) >= 5, "Code too short")'),
    ('LEFT', 'LEFT(text, n)', 'First n characters', 'LEFT(@{fullName}, 1)', None),
    ('RIGHT', 'RIGHT(text, n)', 'Last n characters', 'RIGHT(@{phone}, 4)', None),
    ('SUBSTRING', 'SUBSTRING(text, start, length)', 'Extract portion of text', 'SUBSTRING(@{code}, 0, 3)', 'SUBSTRING(start, end)  [transform]'),
    ('REPLACE', 'REPLACE(text, old, new)', 'Replace all occurrences', 'REPLACE(@{phone}, "-", "")', 'REPLACE("-", "")  [transform]'),
    ('CONTAINS', 'CONTAINS(text, search)', 'True if text contains search', 'IF(CONTAINS(@{email}, "@"), "Valid", "Invalid")', 'Contains("@", "Must contain @")'),
    ('STARTS_WITH', 'STARTS_WITH(text, prefix)', 'True if text starts with prefix', None, 'StartsWith("PRJ-", "Must start with PRJ-")'),
    ('ENDS_WITH', 'ENDS_WITH(text, suffix)', 'True if text ends with suffix', None, 'EndsWith(".pdf", "Must be PDF")'),
    ('CAPITALIZE', 'CAPITALIZE(text)', 'Capitalize first letter of each word', 'CAPITALIZE(@{name})', 'CAPITALIZE()  [transform]'),
    ('TITLE_CASE', 'TITLE_CASE(text)', 'Title Case conversion', 'TITLE_CASE(@{title})', None),
    ('REVERSE', 'REVERSE(text)', 'Reverse the string', 'REVERSE(@{code})', None),
    ('REPEAT', 'REPEAT(text, n)', 'Repeat text n times', 'REPEAT("*", 5)', None),
    ('PAD_LEFT', 'PAD_LEFT(text, len, char)', 'Pad left to length', 'PAD_LEFT(@{id}, 8, "0")', 'PAD_LEFT(8, "0")  [transform]'),
    ('PAD_RIGHT', 'PAD_RIGHT(text, len, char)', 'Pad right to length', None, 'PAD_RIGHT(10, " ")  [transform]'),
    ('WORD_COUNT', 'WORD_COUNT(text)', 'Count words in text', 'WORD_COUNT(@{description})', None),
    ('INDEX_OF', 'INDEX_OF(text, search)', 'Position of first occurrence (-1 if not found)', None, None),
    ('SPLIT', 'SPLIT(text, separator)', 'Split into array', 'SPLIT(@{tags}, ",")', None),
    ('JOIN', 'JOIN(array, separator)', 'Join array into text', None, None),
    ('INITIALS', 'INITIALS(text)', 'First letter of each word', 'INITIALS(@{fullName})', None),
    ('SLUG', 'SLUG(text)', 'URL-friendly slug', 'SLUG(@{title})', 'SLUG()  [transform]'),
    ('REMOVE_SPACES', 'REMOVE_SPACES(text)', 'Remove all whitespace', None, 'REMOVE_SPACES()  [transform]'),
    ('EXTRACT_NUMBERS', 'EXTRACT_NUMBERS(text)', 'Keep only digits', 'EXTRACT_NUMBERS(@{phone})', None),
    ('EXTRACT_LETTERS', 'EXTRACT_LETTERS(text)', 'Keep only letters', None, None),
    ('MASK', 'MASK(text, start, end, char)', 'Mask characters', 'MASK(@{ssn}, 0, 5, "*")', None),
    ('MASK_EMAIL', 'MASK_EMAIL(email)', 'Mask email: j***@domain.com', 'MASK_EMAIL(@{email})', None),
    ('TRUNCATE', 'TRUNCATE(text, len, suffix)', 'Truncate with suffix', 'TRUNCATE(@{desc}, 50, "...")', None),
    ('REGEX_MATCH', 'REGEX_MATCH(text, pattern)', 'True if regex matches', None, 'REGEX_MATCH(@{code}, "^[A-Z]{3}-\\d{4}$")'),
    ('REGEX_REPLACE', 'REGEX_REPLACE(text, pattern, repl)', 'Regex-based replace', 'REGEX_REPLACE(@{phone}, "[^0-9]", "")', None),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 5. NUMBER FUNCTIONS ======
doc.add_page_break()
doc.add_heading('5. Number Functions', level=1)

funcs = [
    ('SUM', 'SUM(a, b, ...)', 'Add numbers together', 'SUM(@{price}, @{tax}, @{shipping})', None),
    ('SUBTRACT', 'SUBTRACT(a, b)', 'Subtract b from a', 'SUBTRACT(@{total}, @{discount})', None),
    ('MULTIPLY', 'MULTIPLY(a, b)', 'Multiply two numbers', 'MULTIPLY(@{qty}, @{price})', None),
    ('DIVIDE', 'DIVIDE(a, b)', 'Divide a by b (returns 0 if b=0)', 'DIVIDE(@{total}, @{count})', None),
    ('ROUND', 'ROUND(number, decimals)', 'Round to n decimal places', 'ROUND(DIVIDE(@{a}, @{b}), 2)', 'ROUND(2)  [transform]'),
    ('ROUND_UP', 'ROUND_UP(number, decimals)', 'Round up (ceiling)', None, 'ROUND_UP(0)  [transform]'),
    ('ROUND_DOWN', 'ROUND_DOWN(number, decimals)', 'Round down (floor)', None, 'ROUND_DOWN(2)  [transform]'),
    ('FLOOR', 'FLOOR(number)', 'Round down to integer', 'FLOOR(@{amount})', None),
    ('CEIL', 'CEIL(number)', 'Round up to integer', 'CEIL(@{amount})', None),
    ('ABS', 'ABS(number)', 'Absolute value', 'ABS(@{difference})', None),
    ('MIN', 'MIN(a, b, ...)', 'Smallest value', 'MIN(@{bid1}, @{bid2}, @{bid3})', None),
    ('MAX', 'MAX(a, b, ...)', 'Largest value', 'MAX(@{score1}, @{score2})', None),
    ('AVERAGE', 'AVERAGE(a, b, ...)', 'Average of values', 'AVERAGE(@{q1}, @{q2}, @{q3}, @{q4})', None),
    ('MEDIAN', 'MEDIAN(a, b, ...)', 'Middle value when sorted', None, None),
    ('PERCENTAGE', 'PERCENTAGE(value, total)', 'Calculate percentage', 'PERCENTAGE(@{completed}, @{total})', None),
    ('MOD', 'MOD(a, b)', 'Remainder after division', 'MOD(@{number}, 2)', None),
    ('POWER', 'POWER(base, exp)', 'Raise to power', 'POWER(@{base}, 2)', None),
    ('SQRT', 'SQRT(number)', 'Square root', 'SQRT(@{variance})', None),
    ('CLAMP', 'CLAMP(value, min, max)', 'Constrain value to range', 'CLAMP(@{rating}, 1, 5)', None),
    ('FORMAT_NUMBER', 'FORMAT_NUMBER(number, decimals)', 'Format with commas', 'FORMAT_NUMBER(@{amount}, 2)', None),
    ('FORMAT_CURRENCY', 'FORMAT_CURRENCY(number, currency, decimals)', 'Format as currency', 'FORMAT_CURRENCY(@{amount}, "$", 2)', None),
    ('RANDOM_INT', 'RANDOM_INT(min, max)', 'Random integer in range', 'RANDOM_INT(1000, 9999)', None),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 6. DATE FUNCTIONS ======
doc.add_page_break()
doc.add_heading('6. Date & Time Functions', level=1)

funcs = [
    ('TODAY', 'TODAY()', 'Current date', 'TODAY()', None),
    ('NOW', 'NOW()', 'Current date and time', 'NOW()', None),
    ('DATE', 'DATE(year, month, day)', 'Create a date', 'DATE(2026, 1, 1)', None),
    ('DATE_ADD', 'DATE_ADD(date, n, unit)', 'Add time to date (years/months/weeks/days/hours/minutes)', 'DATE_ADD(TODAY(), 30, "days")', None),
    ('DATE_SUBTRACT', 'DATE_SUBTRACT(date, n, unit)', 'Subtract time from date', 'DATE_SUBTRACT(TODAY(), 1, "years")', None),
    ('DATE_DIFF', 'DATE_DIFF(date1, date2, unit)', 'Difference between dates', 'DATE_DIFF(@{start}, @{end}, "days")', 'ValidWhen(DATE_DIFF(@{start}, @{end}, "days") > 0, "End must be after start")'),
    ('DATE_FORMAT', 'DATE_FORMAT(date, pattern)', 'Format date as text', 'DATE_FORMAT(TODAY(), "dd/MM/yyyy")', None),
    ('AGE', 'AGE(birthDate)', 'Calculate age in years', 'AGE(@{dateOfBirth})', 'ValidWhen(AGE(@{dob}) >= 18, "Must be 18+")'),
    ('YEAR', 'YEAR(date)', 'Extract year', 'YEAR(TODAY())', None),
    ('MONTH', 'MONTH(date)', 'Extract month (1-12)', 'MONTH(TODAY())', None),
    ('DAY', 'DAY(date)', 'Extract day (1-31)', 'DAY(TODAY())', None),
    ('WEEKDAY', 'WEEKDAY(date)', 'Day of week (1=Mon, 7=Sun)', 'WEEKDAY(TODAY())', None),
    ('WEEKDAY_NAME', 'WEEKDAY_NAME(date)', 'Day name (Monday, etc.)', 'WEEKDAY_NAME(TODAY())', None),
    ('QUARTER', 'QUARTER(date)', 'Quarter (1-4)', 'QUARTER(TODAY())', None),
    ('IS_PAST', 'IS_PAST(date)', 'True if date is before today', None, 'IS_PAST()  or  IS_PAST(@{dateOfLoss})'),
    ('IS_FUTURE', 'IS_FUTURE(date)', 'True if date is after today', None, 'IS_FUTURE()  or  IS_FUTURE(@{expiryDate})'),
    ('IS_WEEKEND', 'IS_WEEKEND(date)', 'True if Saturday or Sunday', None, 'InvalidWhen(IS_WEEKEND(@{date}), "Cannot be weekend")'),
    ('IS_WORKDAY', 'IS_WORKDAY(date)', 'True if Monday-Friday', None, 'ValidWhen(IS_WORKDAY(@{startDate}), "Must be a workday")'),
    ('IS_TODAY', 'IS_TODAY(date)', 'True if date is today', None, None),
    ('IS_BEFORE', 'IS_BEFORE(date1, date2)', 'True if date1 is before date2', None, 'ValidWhen(IS_BEFORE(@{start}, @{end}), "Start must be before end")'),
    ('IS_AFTER', 'IS_AFTER(date1, date2)', 'True if date1 is after date2', None, None),
    ('BUSINESS_DAYS', 'BUSINESS_DAYS(start, end)', 'Count business days between dates', 'BUSINESS_DAYS(@{start}, @{end})', None),
    ('ADD_BUSINESS_DAYS', 'ADD_BUSINESS_DAYS(date, n)', 'Add n business days', 'ADD_BUSINESS_DAYS(TODAY(), 5)', None),
    ('START_OF_MONTH', 'START_OF_MONTH(date)', 'First day of month', 'START_OF_MONTH(TODAY())', None),
    ('END_OF_MONTH', 'END_OF_MONTH(date)', 'Last day of month', 'END_OF_MONTH(TODAY())', None),
    ('RELATIVE_TIME', 'RELATIVE_TIME(date)', 'Human readable (e.g., "3 days ago")', 'RELATIVE_TIME(@{createdAt})', None),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 7. BOOLEAN FUNCTIONS ======
doc.add_page_break()
doc.add_heading('7. Boolean & Logic Functions', level=1)

funcs = [
    ('IF', 'IF(condition, then, else)', 'Conditional value', 'IF(@{amount} > 1000, "High", "Low")', None),
    ('IFS', 'IFS(cond1, val1, cond2, val2, ..., default)', 'Multiple conditions', 'IFS(@{score} >= 90, "A", @{score} >= 80, "B", @{score} >= 70, "C", "F")', None),
    ('SWITCH', 'SWITCH(value, match1, result1, ..., default)', 'Switch/case', 'SWITCH(@{status}, "A", "Active", "I", "Inactive", "Unknown")', None),
    ('AND', 'AND(a, b, ...)', 'True if all true', None, None),
    ('OR', 'OR(a, b, ...)', 'True if any true', None, None),
    ('NOT', 'NOT(value)', 'Negate boolean', None, None),
    ('IS_EMPTY', 'IS_EMPTY(value)', 'True if empty/blank', 'IF(IS_EMPTY(@{middle}), @{first}, CONCAT(@{first}, " ", @{middle}))', None),
    ('IS_NOT_EMPTY', 'IS_NOT_EMPTY(value)', 'True if has value', None, 'IS_NOT_EMPTY(@{field})'),
    ('EQUALS', 'EQUALS(a, b)', 'True if equal', None, 'ValidWhen(EQUALS(@{password}, @{confirm}), "Passwords must match")'),
    ('NOT_EQUALS', 'NOT_EQUALS(a, b)', 'True if not equal', None, None),
    ('GREATER_THAN', 'GREATER_THAN(a, b)', 'True if a > b', None, 'ValidWhen(GREATER_THAN(@{end}, @{start}), "End must be after start")'),
    ('LESS_THAN', 'LESS_THAN(a, b)', 'True if a < b', None, None),
    ('BETWEEN', 'BETWEEN(value, min, max)', 'True if min <= value <= max', None, 'BETWEEN(@{age}, 18, 65, "Age must be 18-65")'),
    ('IN', 'IN(value, a, b, ...)', 'True if value is in list', None, 'IN(@{currency}, "USD", "EUR", "GBP")'),
    ('NOT_IN', 'NOT_IN(value, a, b, ...)', 'True if value not in list', None, 'NOT_IN(@{status}, "CANCELLED", "DELETED")'),
    ('IS_NUMBER', 'IS_NUMBER(value)', 'True if numeric', None, 'IS_NUMBER()'),
    ('IS_EMAIL', 'IS_EMAIL(value)', 'True if valid email', None, 'IS_EMAIL()'),
    ('IS_URL', 'IS_URL(value)', 'True if valid URL', None, 'IS_URL()'),
    ('IS_PHONE', 'IS_PHONE(value)', 'True if valid phone', None, 'IS_PHONE()'),
    ('IS_ALPHA', 'IS_ALPHA(value)', 'True if only letters', None, 'IS_ALPHA()'),
    ('IS_ALPHANUMERIC', 'IS_ALPHANUMERIC(value)', 'True if letters and digits only', None, 'IS_ALPHANUMERIC()'),
    ('IS_NUMERIC', 'IS_NUMERIC(value)', 'True if digits only', None, 'IS_NUMERIC()'),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 8. VALIDATION FUNCTIONS ======
doc.add_page_break()
doc.add_heading('8. Validation Functions', level=1)
doc.add_paragraph('These functions are specifically designed for the Validation Expression field. All accept an optional custom error message as the last argument.')

funcs = [
    ('Required', 'Required("message")', 'Field must not be empty', None, 'Required()\nRequired("Please enter your name")'),
    ('MinLength', 'MinLength(n, "message")', 'Minimum text length', None, 'MinLength(5)\nMinLength(3, "Code must be at least 3 characters")'),
    ('MaxLength', 'MaxLength(n, "message")', 'Maximum text length', None, 'MaxLength(100)\nMaxLength(50, "Description too long")'),
    ('Min', 'Min(n, "message")', 'Minimum numeric value', None, 'Min(0, "Cannot be negative")\nMin(1)'),
    ('Max', 'Max(n, "message")', 'Maximum numeric value', None, 'Max(999999)\nMax(100, "Maximum is 100")'),
    ('Range', 'Range(min, max, "message")', 'Numeric range', None, 'Range(1, 100)\nRange(0, 999, "Must be 0-999")'),
    ('Pattern', 'Pattern(/regex/, "message")', 'Regex pattern match', None, 'Pattern(/^[A-Z]{3}-\\d{4}$/, "Format: ABC-1234")\nPattern(/^\\d{10}$/, "Must be 10 digits")'),
    ('Email', 'Email("message")', 'Valid email format', None, 'Email()\nEmail("Please enter a valid email")'),
    ('Phone', 'Phone("message")', 'Valid phone format', None, 'Phone()\nPhone("Invalid phone number")'),
    ('URL', 'URL("message")', 'Valid URL format', None, 'URL("Please enter a valid URL")'),
    ('PastDate', 'PastDate("message")', 'Date must be before today. Aliases: IS_PAST', None, 'PastDate("Must be a past date")\nIS_PAST()'),
    ('FutureDate', 'FutureDate("message")', 'Date must be after today. Aliases: IS_FUTURE', None, 'FutureDate("Must be a future date")\nIS_FUTURE()'),
    ('DateBefore', 'DateBefore("date", "message")', 'Date before target ("today" supported)', None, 'DateBefore("today")\nDateBefore("2026-12-31", "Must be before year end")'),
    ('DateAfter', 'DateAfter("date", "message")', 'Date after target', None, 'DateAfter("2024-01-01", "Must be after Jan 1, 2024")'),
    ('ValidWhen', 'ValidWhen(expression, "message")', 'Valid when expression is true. Use any library function inside.', None, 'ValidWhen(@{end} > @{start}, "End must be after start")\nValidWhen(IS_PAST(@{dol}), "Must be past date")\nValidWhen(BETWEEN(@{age}, 18, 65), "Age 18-65")'),
    ('InvalidWhen', 'InvalidWhen(expression, "message")', 'Invalid when expression is true', None, 'InvalidWhen(IS_WEEKEND(@{date}), "Cannot be weekend")\nInvalidWhen(@{a} == @{b}, "Values must differ")'),
    ('MatchField', 'MatchField(fieldName, "message")', 'Must match another field', None, 'MatchField(password, "Passwords must match")'),
    ('Unique', 'Unique("message")', 'Must be unique across submissions (async)', None, 'Unique("This ID already exists")'),
    ('Alpha', 'Alpha("message")', 'Letters only', None, 'Alpha("Letters only please")'),
    ('AlphaNumeric', 'AlphaNumeric("message")', 'Letters and numbers only', None, 'AlphaNumeric()'),
    ('Digits', 'Digits("message")', 'Digits only (0-9)', None, 'Digits("Digits only")'),
    ('Positive', 'Positive("message")', 'Must be > 0', None, 'Positive("Must be positive")'),
    ('Integer', 'Integer("message")', 'Must be whole number', None, 'Integer("No decimals allowed")'),
    ('Decimal', 'Decimal(places, "message")', 'Max decimal places', None, 'Decimal(2, "Max 2 decimal places")'),
    ('IsTrue', 'IsTrue("message")', 'Checkbox must be checked', None, 'IsTrue("You must agree to the terms")'),
    ('MinItems', 'MinItems(n, "message")', 'Min selected items (multi-select)', None, 'MinItems(2, "Select at least 2")'),
    ('MaxItems', 'MaxItems(n, "message")', 'Max selected items', None, 'MaxItems(5)'),
    ('MinRows', 'MinRows(n, "message")', 'Min table rows', None, 'MinRows(1, "Add at least one row")'),
    ('MaxRows', 'MaxRows(n, "message")', 'Max table rows', None, 'MaxRows(20)'),
    ('CreditCard', 'CreditCard("message")', 'Valid credit card (Luhn)', None, 'CreditCard("Invalid card number")'),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 9. TRANSFORMATION ======
doc.add_page_break()
doc.add_heading('9. Transformation Functions', level=1)
doc.add_paragraph('Used in the "Transform Expression" field. Applied to the field value on Next/Submit. Combine with AND.')

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr(t, ['Function', 'Description', 'Example'])
for fn, desc, ex in [
    ('UPPER()', 'Convert to uppercase', 'UPPER()'),
    ('LOWER()', 'Convert to lowercase', 'LOWER()'),
    ('TRIM()', 'Remove leading/trailing spaces', 'TRIM()'),
    ('LTRIM()', 'Remove leading spaces', 'LTRIM()'),
    ('RTRIM()', 'Remove trailing spaces', 'RTRIM()'),
    ('CAPITALIZE()', 'Capitalize each word', 'CAPITALIZE()'),
    ('SLUG()', 'URL-friendly slug', 'SLUG()'),
    ('REMOVE_SPACES()', 'Remove all whitespace', 'REMOVE_SPACES()'),
    ('ROUND(n)', 'Round to n decimals', 'ROUND(2)'),
    ('ROUND_UP(n)', 'Round up to n decimals', 'ROUND_UP(0)'),
    ('ROUND_DOWN(n)', 'Round down to n decimals', 'ROUND_DOWN(2)'),
    ('PAD_LEFT(len, "char")', 'Pad left to length', 'PAD_LEFT(8, "0")'),
    ('PAD_RIGHT(len, "char")', 'Pad right to length', 'PAD_RIGHT(10, " ")'),
    ('SUBSTRING(start, end)', 'Extract substring', 'SUBSTRING(0, 5)'),
    ('REPLACE("old", "new")', 'Replace text', 'REPLACE("-", "")'),
]:
    row(t, [fn, desc, ex])

doc.add_paragraph()
doc.add_paragraph('Combining transformations:')
code(doc, 'TRIM() AND UPPER()              -> Trim then uppercase\nTRIM() AND CAPITALIZE()         -> Trim then capitalize\nREMOVE_SPACES() AND UPPER()     -> Remove spaces then uppercase\nROUND(2) AND PAD_LEFT(10, "0")  -> Round then pad')

# ====== 10. TABLE FUNCTIONS ======
doc.add_page_break()
doc.add_heading('10. Table Functions', level=1)
doc.add_paragraph('Functions for TABLE field type default values and computed columns.')

funcs = [
    ('ROW', 'ROW()', 'Current row number', 'ROW()', None),
    ('ROW_COUNT', 'ROW_COUNT(fieldName)', 'Number of rows in table', 'ROW_COUNT(@{items})', 'MinRows(1)'),
    ('SUM_COLUMN', 'SUM_COLUMN(fieldName, columnName)', 'Sum values in a column', 'SUM_COLUMN(@{items}, "amount")', None),
    ('AVG_COLUMN', 'AVG_COLUMN(fieldName, columnName)', 'Average of column values', 'AVG_COLUMN(@{items}, "score")', None),
    ('MIN_COLUMN', 'MIN_COLUMN(fieldName, columnName)', 'Minimum value in column', 'MIN_COLUMN(@{items}, "price")', None),
    ('MAX_COLUMN', 'MAX_COLUMN(fieldName, columnName)', 'Maximum value in column', 'MAX_COLUMN(@{items}, "price")', None),
    ('GET_CELL', 'GET_CELL(fieldName, row, col)', 'Get specific cell value', 'GET_CELL(@{items}, 0, "name")', None),
    ('FIND_ROW', 'FIND_ROW(fieldName, column, value)', 'Find row by column value', 'FIND_ROW(@{items}, "code", "ABC")', None),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 11. UTILITY FUNCTIONS ======
doc.add_page_break()
doc.add_heading('11. Utility Functions', level=1)

funcs = [
    ('UUID', 'UUID()', 'Generate unique ID', 'UUID()', None),
    ('SHORT_UUID', 'SHORT_UUID()', 'Short unique ID (8 chars)', 'SHORT_UUID()', None),
    ('TIMESTAMP', 'TIMESTAMP()', 'Current Unix timestamp', 'TIMESTAMP()', None),
    ('COALESCE', 'COALESCE(a, b, ...)', 'First non-empty value', 'COALESCE(@{nickname}, @{firstName}, "Unknown")', None),
    ('TO_NUMBER', 'TO_NUMBER(value)', 'Convert to number', 'TO_NUMBER(@{quantity})', None),
    ('TO_TEXT', 'TO_TEXT(value)', 'Convert to text', 'TO_TEXT(@{amount})', None),
    ('FORMAT_CURRENCY', 'FORMAT_CURRENCY(amount, symbol, dec)', 'Format as currency', 'FORMAT_CURRENCY(@{total}, "$", 2)', None),
    ('BASE64_ENCODE', 'BASE64_ENCODE(text)', 'Encode to Base64', 'BASE64_ENCODE(@{data})', None),
    ('URL_ENCODE', 'URL_ENCODE(text)', 'URL encode', 'URL_ENCODE(@{query})', None),
    ('TEMPLATE', 'TEMPLATE("Hello {0}, welcome!", arg)', 'Template string', 'TEMPLATE("Ref: {0}-{1}", @{code}, @{id})', None),
    ('TYPEOF', 'TYPEOF(value)', 'Get value type', None, None),
    ('RANDOM_STRING', 'RANDOM_STRING(length)', 'Random alphanumeric string', 'RANDOM_STRING(8)', None),
    ('SEQUENCE', 'SEQUENCE(prefix)', 'Auto-incrementing sequence', 'SEQUENCE("INV-")', None),
]
for name, syntax, desc, def_ex, val_ex in funcs:
    func_entry(doc, name, syntax, desc, def_ex, val_ex)

# ====== 12. USER FUNCTIONS ======
doc.add_heading('12. User & Context Functions', level=1)

t = doc.add_table(rows=0, cols=3)
t.style = 'Table Grid'
hdr(t, ['Function', 'Returns', 'Example'])
for fn, ret, ex in [
    ('CURRENT_USER', 'Full name of logged-in user', 'CURRENT_USER'),
    ('CURRENT_USER_EMAIL', 'Email of logged-in user', 'CURRENT_USER_EMAIL'),
    ('CURRENT_USER_ID', 'UUID of logged-in user', 'CURRENT_USER_ID'),
    ('CURRENT_USERNAME', 'Username', 'CURRENT_USERNAME'),
    ('CURRENT_USER_PHONE', 'Phone number', 'CURRENT_USER_PHONE'),
    ('CURRENT_USER_DEPARTMENT', 'Department', 'CURRENT_USER_DEPARTMENT'),
    ('CURRENT_USER_STAFFID', 'Staff ID', 'CURRENT_USER_STAFFID'),
    ('CURRENT_USER_ROLE', 'Primary role', 'CURRENT_USER_ROLE'),
    ('CURRENT_USER_SBU', 'SBU name', 'CURRENT_USER_SBU'),
    ('CURRENT_USER_BRANCH', 'Branch name', 'CURRENT_USER_BRANCH'),
    ('CURRENT_DATE', 'Today (YYYY-MM-DD)', 'CURRENT_DATE'),
    ('CURRENT_TIME', 'Current time (HH:MM:SS)', 'CURRENT_TIME'),
    ('CURRENT_DATETIME', 'Date + time', 'CURRENT_DATETIME'),
    ('CURRENT_YEAR', 'Current year', 'CURRENT_YEAR'),
    ('WORKFLOW_NAME', 'Current workflow name', 'WORKFLOW_NAME'),
    ('ENVIRONMENT', 'Environment (prod/dev)', 'ENVIRONMENT'),
]:
    row(t, [fn, ret, ex])

doc.add_paragraph()
doc.add_paragraph('These are typically used as Default Values (no parentheses needed):')
code(doc, 'Default Value: CURRENT_USER_EMAIL\nDefault Value: TODAY\nDefault Value: CONCAT("REF-", CURRENT_YEAR, "-", SHORT_UUID())')

# ====== 13. COMBINING ======
doc.add_page_break()
doc.add_heading('13. Combining Functions', level=1)

doc.add_heading('13.1 In Default Values', level=2)
code(doc, '# Full name from parts\nCONCAT(@{firstName}, " ", @{lastName})\n\n# Computed total\nMULTIPLY(@{quantity}, @{unitPrice})\n\n# Conditional default\nIF(@{type} == "Urgent", DATE_ADD(TODAY(), 1, "days"), DATE_ADD(TODAY(), 7, "days"))\n\n# Reference number\nCONCAT("INV-", CURRENT_YEAR, "-", PAD_LEFT(SEQUENCE("inv"), 6, "0"))\n\n# Nested functions\nUPPER(LEFT(@{lastName}, 1))')

doc.add_heading('13.2 In Validation Expressions', level=2)
code(doc, '# Multiple validations with AND\nRequired() AND MinLength(3) AND MaxLength(50)\n\n# Required email format\nRequired() AND Email("Please enter a valid email")\n\n# Numeric range\nRequired() AND Min(0) AND Max(999999)\n\n# Cross-field validation\nValidWhen(IS_BEFORE(@{startDate}, @{endDate}), "Start must be before end")\n\n# Conditional validation with library functions\nValidWhen(AGE(@{dateOfBirth}) >= 18, "Must be 18 or older")\n\n# List validation\nIN(@{currency}, "USD", "EUR", "GBP", "ZiG", "Only supported currencies")\n\n# Complex validation\nRequired() AND Pattern(/^[A-Z]{2}\\d{6}$/, "Format: XX000000")')

doc.add_heading('13.3 In Transform Expressions', level=2)
code(doc, '# Chain transforms\nTRIM() AND UPPER()\n\n# Clean and format\nTRIM() AND REMOVE_SPACES() AND UPPER()\n\n# Round currency\nROUND(2)')

doc.add_heading('13.4 In Visibility Expressions', level=2)
code(doc, '# Show when country is US\n@{country} == \'US\'\n\n# Show when amount exceeds threshold\n@{amount} > 10000\n\n# Show when type is "Other"\n@{type} == \'Other\'\n\n# Combine conditions\n@{country} == \'US\' && @{amount} > 1000')

# ====== 14. QUICK REFERENCE ======
doc.add_page_break()
doc.add_heading('14. Quick Reference', level=1)

doc.add_heading('Most Used Functions', level=2)
t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
hdr(t, ['Function', 'Category', 'Use In', 'Example'])
for fn, cat, use, ex in [
    ('TODAY()', 'Date', 'Default', 'TODAY()'),
    ('NOW()', 'Date', 'Default', 'NOW()'),
    ('CURRENT_USER', 'Context', 'Default', 'CURRENT_USER'),
    ('CONCAT()', 'String', 'Default', 'CONCAT(@{a}, @{b})'),
    ('IF()', 'Logic', 'Default', 'IF(@{x}>0, "Yes", "No")'),
    ('UPPER()', 'String', 'Transform', 'UPPER()'),
    ('TRIM()', 'String', 'Transform', 'TRIM()'),
    ('ROUND()', 'Number', 'Transform', 'ROUND(2)'),
    ('Required()', 'Validation', 'Validation', 'Required()'),
    ('Email()', 'Validation', 'Validation', 'Email()'),
    ('MinLength()', 'Validation', 'Validation', 'MinLength(5)'),
    ('IS_PAST()', 'Date', 'Validation', 'IS_PAST()'),
    ('ValidWhen()', 'Validation', 'Validation', 'ValidWhen(expr, "msg")'),
    ('IN()', 'Logic', 'Validation', 'IN(@{x}, "A", "B")'),
    ('BETWEEN()', 'Logic', 'Validation', 'BETWEEN(@{n}, 1, 100)'),
]:
    row(t, [fn, cat, use, ex])

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\u00A9 2026 Acad Arch Solutions Pvt. Ltd. All rights reserved.')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('Sonar Workflow Functions Manual.docx')
print('Created: Sonar Workflow Functions Manual.docx')
