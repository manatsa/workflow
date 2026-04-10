#!/usr/bin/env python3
"""
Script to generate a Non-Disclosure Agreement for Sonar Microsystems Pvt. Ltd.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_nda():
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('(Confidentiality and Non-Disclosure Agreement for Software Systems)')
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Agreement Date
    date_para = doc.add_paragraph()
    date_para.add_run('Agreement Number: ').bold = True
    date_para.add_run('SONAR-NDA-_____________')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para2 = doc.add_paragraph()
    date_para2.add_run('Date: ').bold = True
    date_para2.add_run('_______________, 20____')
    date_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Preamble
    doc.add_heading('PREAMBLE', level=1)

    preamble = doc.add_paragraph()
    preamble.add_run('This Non-Disclosure Agreement ("Agreement") is entered into as of the date last signed below (the "Effective Date") by and between:')

    doc.add_paragraph()

    # Disclosing Party
    p1 = doc.add_paragraph()
    p1.add_run('DISCLOSING PARTY:').bold = True

    p1a = doc.add_paragraph()
    p1a.add_run('SONAR MICROSYSTEMS PVT. LTD.').bold = True
    p1a.add_run(', a company incorporated under the laws of India, having its registered office at ___________________________________ (hereinafter referred to as "')
    p1a.add_run('Sonar').bold = True
    p1a.add_run('" or "')
    p1a.add_run('Disclosing Party').bold = True
    p1a.add_run('", which expression shall, unless repugnant to the context or meaning thereof, include its successors, affiliates, and permitted assigns);')

    doc.add_paragraph()

    # Receiving Party
    p2 = doc.add_paragraph()
    p2.add_run('RECEIVING PARTY:').bold = True

    p2a = doc.add_paragraph()
    p2a.add_run('___________________________________').bold = True
    p2a.add_run(', a company incorporated under the laws of _______________, having its registered office at ___________________________________ (hereinafter referred to as "')
    p2a.add_run('Recipient').bold = True
    p2a.add_run('" or "')
    p2a.add_run('Receiving Party').bold = True
    p2a.add_run('", which expression shall, unless repugnant to the context or meaning thereof, include its successors and permitted assigns);')

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.add_run('(Sonar and the Recipient are hereinafter individually referred to as a "')
    p3.add_run('Party').bold = True
    p3.add_run('" and collectively as the "')
    p3.add_run('Parties').bold = True
    p3.add_run('")')

    doc.add_paragraph()

    # Recitals
    doc.add_heading('RECITALS', level=1)

    recitals = [
        'WHEREAS, Sonar Microsystems Pvt. Ltd. has developed, owns, and operates proprietary software systems, including but not limited to the "Sonar Workflow System" (as defined herein), which comprises valuable trade secrets, proprietary technologies, methodologies, algorithms, source code, object code, documentation, and related intellectual property;',
        'WHEREAS, the Recipient desires to receive access to, evaluate, and/or utilize certain aspects of the Sonar Workflow System and related Confidential Information for the purposes set forth herein;',
        'WHEREAS, Sonar is willing to disclose such Confidential Information to the Recipient solely for the Permitted Purpose (as defined herein), subject to the terms and conditions of this Agreement;',
        'WHEREAS, both Parties recognize the sensitive and proprietary nature of the Confidential Information and acknowledge that unauthorized disclosure or use thereof would cause irreparable harm to Sonar;',
        'NOW, THEREFORE, in consideration of the mutual covenants, terms, and conditions set forth herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:'
    ]

    for i, recital in enumerate(recitals):
        para = doc.add_paragraph(recital, style='List Number')

    doc.add_paragraph()

    # Article 1 - Definitions
    doc.add_heading('ARTICLE 1: DEFINITIONS', level=1)

    definitions = [
        ('"Confidential Information"', '''means any and all information, data, materials, documents, trade secrets, know-how, and intellectual property of Sonar, whether disclosed orally, in writing, electronically, visually, or in any other form or medium, whether or not marked as "confidential," "proprietary," or with similar designation, including but not limited to:

(a) The Sonar Workflow System in its entirety, including all versions, updates, patches, modifications, and enhancements thereto;

(b) All source code, object code, algorithms, software architecture, system designs, data structures, APIs, interfaces, protocols, and technical specifications;

(c) All documentation, user manuals, technical manuals, installation guides, configuration guides, and training materials;

(d) All business information, including customer lists, pricing information, marketing strategies, business plans, financial data, and operational procedures;

(e) All proprietary methodologies, processes, techniques, formulas, compositions, and inventions;

(f) All security measures, encryption methods, authentication systems, access control mechanisms, and vulnerability information;

(g) All database schemas, data models, data dictionaries, and stored procedures;

(h) All user interface designs, user experience designs, graphics, icons, and visual elements;

(i) All testing procedures, test cases, test data, quality assurance processes, and bug reports;

(j) All deployment procedures, infrastructure configurations, and operational runbooks;

(k) All third-party software, libraries, and components integrated into or used with the Sonar Workflow System;

(l) Any information that a reasonable person would understand to be confidential given the nature of the information and circumstances of disclosure.'''),

        ('"Sonar Workflow System"', '''means Sonar's proprietary enterprise workflow management software system, including all associated modules, components, features, functionalities, and related technologies, whether currently existing or developed in the future, including but not limited to:

(a) The core workflow engine and all business process automation capabilities;
(b) The user management and authentication systems;
(c) The role-based access control and privilege management systems;
(d) The form builder and dynamic field configuration systems;
(e) The approval workflow and routing mechanisms;
(f) The audit logging and compliance tracking systems;
(g) The reporting and analytics modules;
(h) The notification and alert systems;
(i) The integration APIs and connectors;
(j) The administrative console and configuration interfaces;
(k) All frontend and backend components, databases, and infrastructure configurations.'''),

        ('"Permitted Purpose"', '''means the limited purpose for which Confidential Information is disclosed, specifically:

(a) Evaluation of the Sonar Workflow System for potential licensing, partnership, or acquisition;
(b) Integration, implementation, or deployment of the Sonar Workflow System within the Recipient's organization pursuant to a separate licensing agreement;
(c) Such other specific purposes as may be agreed upon in writing by both Parties.'''),

        ('"Representatives"', 'means the officers, directors, employees, agents, contractors, consultants, advisors, and other authorized representatives of the Receiving Party who have a legitimate need to know the Confidential Information for the Permitted Purpose and who are bound by confidentiality obligations at least as protective as those contained herein.'),

        ('"Derivative Works"', 'means any modifications, adaptations, translations, enhancements, improvements, or derivative works based upon or incorporating any Confidential Information.'),

        ('"Intellectual Property Rights"', 'means all patents, copyrights, trademarks, trade secrets, trade dress, moral rights, rights of publicity, and all other intellectual property rights, whether registered or unregistered, and all applications, renewals, extensions, and restorations thereof, now or hereafter in force and effect worldwide.'),
    ]

    for term, definition in definitions:
        para = doc.add_paragraph()
        para.add_run(f'1.{definitions.index((term, definition)) + 1} ').bold = True
        para.add_run(term).bold = True
        para.add_run(f' {definition}')

    doc.add_paragraph()

    # Article 2 - Obligations of Confidentiality
    doc.add_heading('ARTICLE 2: OBLIGATIONS OF CONFIDENTIALITY', level=1)

    obligations = [
        ('Protection of Confidential Information', '''The Receiving Party shall:

(a) Hold all Confidential Information in strict confidence and protect it with at least the same degree of care as the Receiving Party uses to protect its own most sensitive confidential information, but in no event less than a reasonable degree of care;

(b) Use Confidential Information solely for the Permitted Purpose and for no other purpose whatsoever;

(c) Not disclose, publish, disseminate, or otherwise make available any Confidential Information to any third party without the prior written consent of Sonar;

(d) Limit disclosure of Confidential Information to those Representatives who have a genuine need to know such information for the Permitted Purpose;

(e) Ensure that all Representatives who receive Confidential Information are informed of its confidential nature and are bound by written confidentiality obligations no less restrictive than those contained in this Agreement;

(f) Immediately notify Sonar in writing of any unauthorized access, use, disclosure, or loss of Confidential Information;

(g) Cooperate fully with Sonar in any investigation of unauthorized disclosure and in any efforts to protect Sonar's rights.'''),

        ('Prohibited Actions', '''The Receiving Party shall NOT:

(a) Copy, reproduce, duplicate, or replicate any Confidential Information except as strictly necessary for the Permitted Purpose and with Sonar's prior written approval;

(b) Reverse engineer, decompile, disassemble, or otherwise attempt to derive the source code, algorithms, structure, or design of the Sonar Workflow System or any component thereof;

(c) Create any Derivative Works based on the Confidential Information without Sonar's prior written consent;

(d) Remove, alter, or obscure any proprietary notices, labels, or markings on any Confidential Information;

(e) Use any Confidential Information to develop, design, create, manufacture, or market any product or service that is competitive with or similar to the Sonar Workflow System;

(f) Benchmark, test, or evaluate the Sonar Workflow System for the purpose of comparing it to competing products or publishing such comparisons;

(g) Transfer, assign, sublicense, or otherwise convey any rights in the Confidential Information to any third party;

(h) Use the Confidential Information to recruit, solicit, or hire any employees, contractors, or consultants of Sonar;

(i) Access or attempt to access any systems, networks, or data of Sonar beyond what is expressly authorized in writing.'''),

        ('Security Measures', '''The Receiving Party shall implement and maintain appropriate technical and organizational security measures to protect Confidential Information, including but not limited to:

(a) Physical security controls to prevent unauthorized access to facilities where Confidential Information is stored or processed;

(b) Logical access controls, including unique user identification, strong password requirements, and multi-factor authentication;

(c) Encryption of Confidential Information during transmission and at rest using industry-standard encryption algorithms;

(d) Network security measures including firewalls, intrusion detection systems, and regular security monitoring;

(e) Regular security assessments, vulnerability scans, and penetration testing;

(f) Secure disposal procedures for physical and electronic media containing Confidential Information;

(g) Employee security awareness training and background checks for personnel with access to Confidential Information;

(h) Incident response procedures to address security breaches promptly and effectively;

(i) Maintenance of detailed access logs and audit trails for all access to Confidential Information.'''),
    ]

    for i, (title, content) in enumerate(obligations):
        para = doc.add_paragraph()
        para.add_run(f'2.{i + 1} ').bold = True
        para.add_run(title).bold = True
        para.add_run(f' {content}')

    doc.add_paragraph()

    # Article 3 - Exclusions
    doc.add_heading('ARTICLE 3: EXCLUSIONS FROM CONFIDENTIAL INFORMATION', level=1)

    exclusions_intro = doc.add_paragraph()
    exclusions_intro.add_run('3.1 ').bold = True
    exclusions_intro.add_run('The obligations of confidentiality set forth in this Agreement shall not apply to information that the Receiving Party can demonstrate by clear and convincing documentary evidence:')

    exclusions = [
        'Was already in the public domain at the time of disclosure through no fault or action of the Receiving Party or its Representatives;',
        'Becomes publicly available after disclosure through no breach of this Agreement by the Receiving Party or its Representatives;',
        'Was already lawfully in the possession of the Receiving Party prior to disclosure by Sonar, as evidenced by contemporaneous written records, and was not subject to any confidentiality obligation;',
        'Is independently developed by the Receiving Party without reference to, use of, or access to any Confidential Information, as demonstrated by documented evidence of independent development;',
        'Is lawfully obtained by the Receiving Party from a third party who has the lawful right to disclose such information without any confidentiality restriction;',
        'Is required to be disclosed by applicable law, regulation, court order, or governmental authority, provided that the Receiving Party: (i) provides Sonar with prompt written notice of such requirement prior to disclosure to the extent legally permitted; (ii) cooperates with Sonar in seeking a protective order or other appropriate remedy; (iii) discloses only that portion of the Confidential Information that is legally required to be disclosed; and (iv) uses reasonable efforts to obtain confidential treatment for any Confidential Information so disclosed.'
    ]

    for i, exclusion in enumerate(exclusions):
        para = doc.add_paragraph(f'({chr(97 + i)}) {exclusion}')
        para.paragraph_format.left_indent = Inches(0.5)

    para_32 = doc.add_paragraph()
    para_32.add_run('3.2 ').bold = True
    para_32.add_run('The burden of proving that any exclusion applies shall rest solely with the Receiving Party. Any combination of features or information shall not be deemed to be within the foregoing exclusions merely because individual features or information are within such exclusions, but only if the combination itself is within such exclusions.')

    doc.add_paragraph()

    # Article 4 - Intellectual Property Rights
    doc.add_heading('ARTICLE 4: INTELLECTUAL PROPERTY RIGHTS', level=1)

    ip_sections = [
        ('Ownership', 'All Confidential Information, including the Sonar Workflow System and all related Intellectual Property Rights, are and shall remain the sole and exclusive property of Sonar. Nothing in this Agreement shall be construed as granting or conferring any rights, by license or otherwise, expressly, impliedly, or by estoppel, to any Confidential Information or any Intellectual Property Rights therein.'),

        ('No License', 'No license, express or implied, in the Confidential Information or any patent, copyright, trademark, trade secret, or other intellectual property right is granted to the Receiving Party under this Agreement. The Receiving Party acknowledges that any use of the Sonar Workflow System beyond the Permitted Purpose requires a separate written license agreement with Sonar.'),

        ('Derivative Works', 'Any Derivative Works created by the Receiving Party that are based upon or incorporate any Confidential Information shall be the sole and exclusive property of Sonar. The Receiving Party hereby assigns and agrees to assign to Sonar all right, title, and interest in and to any such Derivative Works, including all Intellectual Property Rights therein.'),

        ('Trademarks', 'The Receiving Party shall not use any trademarks, service marks, trade names, logos, or other brand identifiers of Sonar without Sonar\'s prior written consent. Nothing in this Agreement grants the Receiving Party any right to use Sonar\'s marks for any purpose.'),

        ('Feedback', 'If the Receiving Party provides any suggestions, ideas, feedback, or recommendations regarding the Confidential Information or the Sonar Workflow System ("Feedback"), Sonar shall own all right, title, and interest in such Feedback and may use, implement, and commercialize such Feedback without any obligation to the Receiving Party.'),
    ]

    for i, (title, content) in enumerate(ip_sections):
        para = doc.add_paragraph()
        para.add_run(f'4.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Article 5 - Term and Termination
    doc.add_heading('ARTICLE 5: TERM AND TERMINATION', level=1)

    term_sections = [
        ('Term', 'This Agreement shall commence on the Effective Date and shall continue in full force and effect for a period of five (5) years, unless earlier terminated in accordance with this Article 5 (the "Term"). Upon expiration, this Agreement may be renewed by mutual written agreement of the Parties.'),

        ('Survival of Obligations', 'Notwithstanding any termination or expiration of this Agreement, the obligations of confidentiality set forth herein shall survive and continue for a period of ten (10) years following the date of termination or expiration, or for so long as the Confidential Information remains a trade secret under applicable law, whichever is longer.'),

        ('Termination for Convenience', 'Either Party may terminate this Agreement at any time, with or without cause, upon thirty (30) days\' prior written notice to the other Party.'),

        ('Termination for Breach', 'Sonar may terminate this Agreement immediately upon written notice if: (a) the Receiving Party breaches any provision of this Agreement; (b) the Receiving Party becomes insolvent, files for bankruptcy, or is subject to any similar proceeding; or (c) Sonar reasonably believes that the Confidential Information is at risk of unauthorized disclosure.'),

        ('Effect of Termination', '''Upon termination or expiration of this Agreement for any reason, the Receiving Party shall immediately:

(a) Cease all use of the Confidential Information;
(b) Return to Sonar all originals and copies of Confidential Information in any form or medium;
(c) Permanently destroy all electronic copies of Confidential Information from all systems, storage media, and backup systems;
(d) Certify in writing, signed by an authorized officer, that all Confidential Information has been returned or destroyed;
(e) Delete all Confidential Information from any cloud storage, third-party systems, or other locations;
(f) Provide Sonar with access to verify the destruction of Confidential Information if requested.'''),

        ('Return of Materials', 'Notwithstanding the foregoing, the Receiving Party may retain one (1) archival copy of Confidential Information solely for legal compliance purposes, provided that such copy is stored securely and remains subject to the confidentiality obligations of this Agreement.'),
    ]

    for i, (title, content) in enumerate(term_sections):
        para = doc.add_paragraph()
        para.add_run(f'5.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Article 6 - Warranties and Disclaimers
    doc.add_heading('ARTICLE 6: WARRANTIES AND DISCLAIMERS', level=1)

    warranty_sections = [
        ('Recipient Warranties', '''The Receiving Party represents and warrants that:

(a) It has the full power and authority to enter into this Agreement and to perform its obligations hereunder;
(b) The execution and performance of this Agreement does not conflict with any other agreement or obligation to which the Receiving Party is bound;
(c) It will comply with all applicable laws, regulations, and industry standards in connection with its receipt and use of Confidential Information;
(d) It maintains appropriate security measures and internal controls to protect confidential information;
(e) It will not export or re-export any Confidential Information in violation of applicable export control laws and regulations.'''),

        ('Disclaimer', 'THE CONFIDENTIAL INFORMATION IS PROVIDED "AS IS" WITHOUT WARRANTY OF ANY KIND. SONAR MAKES NO WARRANTIES, EXPRESS, IMPLIED, STATUTORY, OR OTHERWISE, WITH RESPECT TO THE CONFIDENTIAL INFORMATION, INCLUDING WITHOUT LIMITATION ANY IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, NON-INFRINGEMENT, OR ACCURACY. SONAR DOES NOT WARRANT THAT THE CONFIDENTIAL INFORMATION WILL BE ERROR-FREE, COMPLETE, OR SUITABLE FOR ANY PARTICULAR PURPOSE.'),

        ('No Obligation', 'Nothing in this Agreement shall obligate Sonar to: (a) disclose any particular Confidential Information; (b) continue disclosing Confidential Information; (c) enter into any further agreement with the Receiving Party; or (d) proceed with any business relationship or transaction with the Receiving Party.'),
    ]

    for i, (title, content) in enumerate(warranty_sections):
        para = doc.add_paragraph()
        para.add_run(f'6.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Article 7 - Remedies
    doc.add_heading('ARTICLE 7: REMEDIES', level=1)

    remedy_sections = [
        ('Irreparable Harm', 'The Receiving Party acknowledges and agrees that any unauthorized use, disclosure, or breach of this Agreement would cause irreparable harm to Sonar for which monetary damages would be inadequate. Accordingly, Sonar shall be entitled to seek injunctive relief, specific performance, and other equitable remedies, without the necessity of proving actual damages or posting any bond or other security.'),

        ('Cumulative Remedies', 'The remedies provided in this Agreement are cumulative and not exclusive of any other remedies available at law or in equity. Sonar may pursue any and all remedies available, including but not limited to: (a) injunctive relief; (b) specific performance; (c) compensatory damages; (d) consequential damages; (e) punitive damages; (f) disgorgement of profits; and (g) recovery of attorneys\' fees and costs.'),

        ('Indemnification', 'The Receiving Party shall indemnify, defend, and hold harmless Sonar and its officers, directors, employees, agents, successors, and assigns from and against any and all claims, damages, losses, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) any breach of this Agreement by the Receiving Party or its Representatives; (b) any unauthorized use or disclosure of Confidential Information; (c) any violation of applicable laws or regulations by the Receiving Party.'),

        ('Liquidated Damages', 'In addition to all other remedies, the Receiving Party agrees to pay Sonar liquidated damages in the amount of INR 50,00,000 (Fifty Lakhs Indian Rupees) or USD 60,000 (Sixty Thousand US Dollars), whichever is greater, for each material breach of this Agreement. The Parties agree that this amount represents a reasonable estimate of the damages Sonar would suffer and is not a penalty.'),

        ('Audit Rights', 'Sonar shall have the right, upon reasonable notice and during normal business hours, to audit the Receiving Party\'s compliance with this Agreement. The Receiving Party shall cooperate fully with any such audit and provide access to relevant systems, facilities, and personnel. If any audit reveals a breach of this Agreement, the Receiving Party shall bear the costs of the audit.'),
    ]

    for i, (title, content) in enumerate(remedy_sections):
        para = doc.add_paragraph()
        para.add_run(f'7.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Article 8 - General Provisions
    doc.add_heading('ARTICLE 8: GENERAL PROVISIONS', level=1)

    general_sections = [
        ('Governing Law', 'This Agreement shall be governed by and construed in accordance with the laws of India, without regard to its conflict of laws principles. The Parties submit to the exclusive jurisdiction of the courts located in [City], India for any disputes arising out of or relating to this Agreement.'),

        ('Dispute Resolution', 'Any dispute arising out of or relating to this Agreement shall first be attempted to be resolved through good faith negotiations between senior representatives of the Parties. If not resolved within thirty (30) days, the dispute shall be submitted to binding arbitration in accordance with the Arbitration and Conciliation Act, 1996, as amended. The arbitration shall be conducted in English in [City], India, before a panel of three arbitrators.'),

        ('Entire Agreement', 'This Agreement constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, understandings, negotiations, and discussions, whether oral or written, between the Parties relating to the subject matter of this Agreement.'),

        ('Amendment', 'This Agreement may not be amended, modified, or supplemented except by a written instrument signed by authorized representatives of both Parties. No waiver of any provision of this Agreement shall be effective unless in writing and signed by the waiving Party.'),

        ('Waiver', 'The failure of either Party to enforce any provision of this Agreement shall not constitute a waiver of that Party\'s right to enforce that provision or any other provision. No waiver of any breach shall be deemed a waiver of any subsequent breach.'),

        ('Severability', 'If any provision of this Agreement is held to be invalid, illegal, or unenforceable, the remaining provisions shall continue in full force and effect. The invalid provision shall be modified to the minimum extent necessary to make it valid, legal, and enforceable while preserving its original intent.'),

        ('Assignment', 'The Receiving Party may not assign or transfer this Agreement or any rights or obligations hereunder without the prior written consent of Sonar. Any attempted assignment in violation of this provision shall be null and void. Sonar may assign this Agreement to any successor in interest or affiliate without consent.'),

        ('Notices', '''All notices, requests, and other communications under this Agreement shall be in writing and shall be deemed given when: (a) delivered personally; (b) sent by registered or certified mail, return receipt requested; (c) sent by reputable overnight courier; or (d) sent by email with confirmation of receipt. Notices shall be sent to the addresses set forth in this Agreement or such other addresses as may be designated in writing.'''),

        ('Independent Parties', 'The Parties are independent contractors. Nothing in this Agreement shall be construed as creating a partnership, joint venture, agency, employment, or fiduciary relationship between the Parties.'),

        ('No Third-Party Beneficiaries', 'This Agreement is for the sole benefit of the Parties and their permitted successors and assigns. Nothing in this Agreement shall confer any rights or remedies upon any third party.'),

        ('Counterparts', 'This Agreement may be executed in counterparts, each of which shall be deemed an original and all of which together shall constitute one and the same instrument. Electronic signatures shall be deemed valid and binding.'),

        ('Headings', 'The headings in this Agreement are for convenience only and shall not affect the interpretation of this Agreement.'),

        ('Language', 'This Agreement is executed in the English language. In the event of any translation, the English version shall prevail.'),

        ('Force Majeure', 'Neither Party shall be liable for any failure or delay in performance due to causes beyond its reasonable control, including but not limited to acts of God, natural disasters, war, terrorism, riots, embargoes, acts of government, or failures of third-party telecommunications or power supply.'),
    ]

    for i, (title, content) in enumerate(general_sections):
        para = doc.add_paragraph()
        para.add_run(f'8.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()

    # Article 9 - Special Provisions for Software Systems
    doc.add_heading('ARTICLE 9: SPECIAL PROVISIONS FOR SOFTWARE SYSTEMS', level=1)

    software_sections = [
        ('Source Code Protection', 'The Receiving Party acknowledges that the source code of the Sonar Workflow System constitutes the most sensitive and valuable trade secret of Sonar. Access to source code, if granted, shall be subject to additional security measures and restrictions as specified by Sonar, including but not limited to access only in a secure facility, prohibition on copying, and real-time monitoring.'),

        ('No Competitive Use', 'The Receiving Party agrees that for a period of three (3) years following termination of this Agreement, it shall not: (a) develop any software product that competes with the Sonar Workflow System; (b) assist any third party in developing such competing software; or (c) use any knowledge gained from the Confidential Information to develop similar functionality.'),

        ('Security Incidents', 'The Receiving Party shall immediately notify Sonar (within 24 hours) of any security incident, data breach, or unauthorized access involving Confidential Information. The Receiving Party shall cooperate with Sonar\'s incident response procedures and take all reasonable steps to mitigate any damage.'),

        ('Personnel Restrictions', 'The Receiving Party shall maintain a list of all personnel who have access to Confidential Information and shall provide this list to Sonar upon request. The Receiving Party shall immediately revoke access for any personnel who leave the organization or no longer require access for the Permitted Purpose.'),

        ('Export Controls', 'The Receiving Party acknowledges that the Confidential Information may be subject to export control laws and regulations. The Receiving Party agrees to comply with all applicable export control laws and not to export or re-export any Confidential Information without appropriate government authorization.'),
    ]

    for i, (title, content) in enumerate(software_sections):
        para = doc.add_paragraph()
        para.add_run(f'9.{i + 1} ').bold = True
        para.add_run(f'{title}. ').bold = True
        para.add_run(content)

    doc.add_paragraph()
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Signature Block
    doc.add_heading('SIGNATURES', level=1)

    sig_intro = doc.add_paragraph()
    sig_intro.add_run('IN WITNESS WHEREOF, the Parties have executed this Non-Disclosure Agreement as of the date last written below.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Create signature table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Disclosing Party Signature
    cell1 = table.rows[0].cells[0]
    p1 = cell1.add_paragraph()
    p1.add_run('DISCLOSING PARTY:').bold = True
    p1.add_run('\nSONAR MICROSYSTEMS PVT. LTD.')
    cell1.add_paragraph()
    cell1.add_paragraph()
    cell1.add_paragraph('_________________________________')
    cell1.add_paragraph('Authorized Signatory')
    cell1.add_paragraph()
    cell1.add_paragraph('Name: _________________________')
    cell1.add_paragraph('Title: __________________________')
    cell1.add_paragraph('Date: __________________________')
    cell1.add_paragraph()
    cell1.add_paragraph('_________________________________')
    cell1.add_paragraph('Witness Signature')
    cell1.add_paragraph('Name: _________________________')

    # Receiving Party Signature
    cell2 = table.rows[0].cells[1]
    p2 = cell2.add_paragraph()
    p2.add_run('RECEIVING PARTY:').bold = True
    p2.add_run('\n_________________________________')
    cell2.add_paragraph()
    cell2.add_paragraph()
    cell2.add_paragraph('_________________________________')
    cell2.add_paragraph('Authorized Signatory')
    cell2.add_paragraph()
    cell2.add_paragraph('Name: _________________________')
    cell2.add_paragraph('Title: __________________________')
    cell2.add_paragraph('Date: __________________________')
    cell2.add_paragraph()
    cell2.add_paragraph('_________________________________')
    cell2.add_paragraph('Witness Signature')
    cell2.add_paragraph('Name: _________________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Schedule A - Permitted Purpose
    doc.add_page_break()
    doc.add_heading('SCHEDULE A: PERMITTED PURPOSE', level=1)

    schedule_a = doc.add_paragraph()
    schedule_a.add_run('The Permitted Purpose for the disclosure of Confidential Information under this Agreement is as follows:')

    doc.add_paragraph()

    purposes = [
        'Evaluation: To evaluate the Sonar Workflow System for potential licensing, acquisition, or partnership opportunities.',
        'Implementation: To implement and deploy the Sonar Workflow System within the Receiving Party\'s organization pursuant to a valid license agreement.',
        'Integration: To integrate the Sonar Workflow System with the Receiving Party\'s existing systems and infrastructure.',
        'Training: To train authorized personnel on the use and administration of the Sonar Workflow System.',
        'Support: To receive technical support and maintenance services from Sonar.',
        'Other: [Specify any additional permitted purposes agreed upon by the Parties]'
    ]

    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')

    doc.add_paragraph()

    # Schedule B - List of Authorized Representatives
    doc.add_heading('SCHEDULE B: AUTHORIZED REPRESENTATIVES', level=1)

    schedule_b = doc.add_paragraph()
    schedule_b.add_run('The following individuals are authorized to receive Confidential Information on behalf of the Receiving Party:')

    doc.add_paragraph()

    # Create table for authorized personnel
    rep_table = doc.add_table(rows=6, cols=4)
    rep_table.style = 'Table Grid'

    headers = ['Name', 'Title/Position', 'Email', 'Signature']
    for i, header in enumerate(headers):
        cell = rep_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    for row in range(1, 6):
        for col in range(4):
            rep_table.rows[row].cells[col].text = ''

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = doc.add_paragraph()
    footer.add_run('SONAR MICROSYSTEMS PVT. LTD. - CONFIDENTIAL').bold = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer2 = doc.add_paragraph()
    footer2.add_run(f'Document Generated: {datetime.now().strftime("%B %d, %Y")}')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save('Sonar NDA Agreement.docx')
    print('NDA document created successfully: Sonar NDA Agreement.docx')

if __name__ == '__main__':
    create_nda()
